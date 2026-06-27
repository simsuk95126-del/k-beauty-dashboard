from __future__ import annotations

"""Country-specific customer report layer for K-Beauty v9.1.1.

This module keeps the proven r3 layout/localization utilities while changing the
customer-facing report model to the backend v6.2 schema:

- country/market official terminology and references first
- business screening interpretation separated from internal status codes
- audit-source traceability exposed in Excel and concise official-source context
  exposed in Word
- HALAL retained as an independent evidence-review workflow
"""

import io
import re
from datetime import datetime
from typing import Any, Dict, Iterable, List, Optional, Sequence

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

import report_builder_v911_r3 as base


REPORT_BUILD_ID = "v9.1.1-report-r4-country-specific"

# Public constants reused by dashboard.py
TRANSLATION_GLOSSARY_PROMPTS = base.TRANSLATION_GLOSSARY_PROMPTS
LANGUAGE_FILE_SUFFIXES = base.LANGUAGE_FILE_SUFFIXES
MARKET_LABELS = base.MARKET_LABELS
OUTPUT_LANGUAGES = base.OUTPUT_LANGUAGES
STATUS_LABELS = base.STATUS_LABELS

# Public utility aliases retained for compatibility.
clean_text = base.clean_text
safe_int = base.safe_int
normalize_status = base.normalize_status
safe_filename = base.safe_filename
output_file_name = base.output_file_name
output_zip_name = base.output_zip_name
status_counts = base.status_counts
dominant_status = base.dominant_status
halal_decision = base.halal_decision
halal_status_label = base.halal_status_label
required_evidence_for_detail = base.required_evidence_for_detail
_base_action_for_detail = base.action_for_detail


SCREENING_DECISION_LABELS: Dict[str, str] = {
    "IMMEDIATE_STOP_SIGNAL_IDENTIFIED": "우선 중단 요인 확인",
    "CONDITIONS_REQUIRE_VERIFICATION": "조건 확인 후 진행 판단",
    "ADDITIONAL_INFORMATION_REQUIRED": "추가 정보·원문 확인 후 판단",
    "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED": "현재 확인 범위에서 즉시 중단 신호 미확인",
}

SCREENING_DECISION_DESCRIPTIONS: Dict[str, str] = {
    "IMMEDIATE_STOP_SIGNAL_IDENTIFIED": (
        "탑재된 공식 규제자료에서 현재 처방의 진행을 우선 멈추고 동일물질 여부, "
        "적용범위와 예외를 확인해야 할 명시적 신호가 확인되었습니다."
    ),
    "CONDITIONS_REQUIRE_VERIFICATION": (
        "사용조건, 농도, 제품유형, 사용부위, 표시 또는 고지요건을 확인한 뒤 "
        "진행 여부를 판단해야 합니다."
    ),
    "ADDITIONAL_INFORMATION_REQUIRED": (
        "공식 원문, 성분 정체, 공급사 자료 또는 규정 적용범위에 관한 추가정보가 "
        "필요하여 현재 단계에서 결론을 확정하지 않습니다."
    ),
    "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED": (
        "현재 시스템에 탑재된 자료에서는 시장 검토를 즉시 중단해야 할 명시적 "
        "성분 일치 신호가 확인되지 않았습니다. 최종 적합성이나 판매승인을 뜻하지 않습니다."
    ),
}

SCREENING_DECISION_ORDER = [
    "IMMEDIATE_STOP_SIGNAL_IDENTIFIED",
    "CONDITIONS_REQUIRE_VERIFICATION",
    "ADDITIONAL_INFORMATION_REQUIRED",
    "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED",
]

FALLBACK_MARKET_PROFILES: Dict[str, dict] = {
    "US": {
        "display_name_ko": "미국",
        "screening_scope_ko": "현재 시스템에 탑재된 미국 화장품 성분 규제자료를 대조합니다.",
        "no_match_interpretation_ko": SCREENING_DECISION_DESCRIPTIONS[
            "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"
        ],
        "next_checks_ko": [
            "연방·주 규제, 색소첨가물 및 제품 분류",
            "라벨, 안전성 입증, MoCRA 및 수입·유통 요건",
        ],
    },
    "EU": {
        "display_name_ko": "유럽연합(EU)",
        "screening_scope_ko": "현재 시스템에 탑재된 EU 화장품 규정 부속서 기반 성분자료를 대조합니다.",
        "no_match_interpretation_ko": SCREENING_DECISION_DESCRIPTIONS[
            "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"
        ],
        "next_checks_ko": ["배합농도·제품유형별 부속서 조건", "CPSR·PIF·책임자·CPNP·라벨"],
    },
    "UK": {
        "display_name_ko": "영국(GB)",
        "screening_scope_ko": "현재 시스템에 탑재된 영국 화장품 규정 성분자료를 대조합니다.",
        "no_match_interpretation_ko": SCREENING_DECISION_DESCRIPTIONS[
            "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"
        ],
        "next_checks_ko": ["배합농도·제품유형별 조건", "영국 책임자·안전성 평가·SCPN·라벨"],
    },
    "CN": {
        "display_name_ko": "중국",
        "screening_scope_ko": "현재 시스템에 탑재된 NMPA 화장품 성분 규제자료를 대조합니다.",
        "no_match_interpretation_ko": SCREENING_DECISION_DESCRIPTIONS[
            "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"
        ],
        "next_checks_ko": ["기존·신규 원료 지위", "등록·비안·안전성 평가·중국어 라벨"],
    },
    "ASEAN": {
        "display_name_ko": "ASEAN",
        "screening_scope_ko": "현재 시스템에 탑재된 ASEAN Cosmetic Directive 부속서 성분자료를 대조합니다.",
        "no_match_interpretation_ko": SCREENING_DECISION_DESCRIPTIONS[
            "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"
        ],
        "next_checks_ko": ["ACD 조건", "판매 회원국별 신고·책임자·라벨·집행요건"],
    },
    "SFDA": {
        "display_name_ko": "사우디아라비아",
        "screening_scope_ko": "현재 시스템에 탑재된 SFDA 성분 규제목록을 대조합니다.",
        "no_match_interpretation_ko": SCREENING_DECISION_DESCRIPTIONS[
            "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"
        ],
        "next_checks_ko": ["배합농도·제품유형·사용부위 조건", "제품 등록·라벨·현지 책임요건"],
    },
    "EAC": {
        "display_name_ko": "유라시아경제연합(EAEU)",
        "screening_scope_ko": "현재 시스템에 탑재된 TR CU 009/2011 부속서 기반 성분자료를 대조합니다.",
        "no_match_interpretation_ko": SCREENING_DECISION_DESCRIPTIONS[
            "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"
        ],
        "next_checks_ko": ["배합농도·제품유형별 조건", "적합성 선언·국가등록·EAEU 표시요건"],
    },
    "HALAL": {
        "display_name_ko": "HALAL 추가검토",
        "screening_scope_ko": "성분명에서 원료 기원·구성·공정·인증자료 확인이 필요한 후보를 선별합니다.",
        "no_match_interpretation_ko": "성분명에서 명시적 위험 후보가 확인되지 않았다는 뜻이며 인증판정이 아닙니다.",
        "next_checks_ko": ["원료 기원·구성표", "제조공정·가공보조제·교차오염·인증자료"],
    },
}


def _dict(value: Any) -> dict:
    return value if isinstance(value, dict) else {}


def _list(value: Any) -> list:
    return value if isinstance(value, list) else []


def internal_status_for_detail(detail: dict) -> str:
    return normalize_status(
        detail.get("internal_status")
        or detail.get("compliance_status")
        or detail.get("restriction_type")
    )


def screening_decision_from_status(status: Any) -> str:
    status = normalize_status(status)
    if status == "BANNED":
        return "IMMEDIATE_STOP_SIGNAL_IDENTIFIED"
    if status in {"RESTRICTED", "WARNING_REQUIRED"}:
        return "CONDITIONS_REQUIRE_VERIFICATION"
    if status in {"REGULATED", "REVIEW_REQUIRED", "VERIFICATION_REQUIRED"}:
        return "ADDITIONAL_INFORMATION_REQUIRED"
    return "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"


def detail_screening_decision(detail: dict) -> str:
    decision = clean_text(detail.get("screening_decision")).upper()
    if decision in SCREENING_DECISION_LABELS:
        return decision
    return screening_decision_from_status(internal_status_for_detail(detail))


def detail_screening_label(detail: dict) -> str:
    explicit = clean_text(detail.get("screening_decision_label"))
    return explicit or SCREENING_DECISION_LABELS[detail_screening_decision(detail)]


def result_screening_decision(result_data: dict) -> str:
    decision = clean_text(result_data.get("screening_decision")).upper()
    if decision in SCREENING_DECISION_LABELS:
        return decision
    return screening_decision_from_status(
        result_data.get("internal_status")
        or result_data.get("compliance_status")
        or dominant_status(result_data)
    )


def result_screening_label(result_data: dict) -> str:
    explicit = clean_text(result_data.get("screening_decision_label"))
    return explicit or SCREENING_DECISION_LABELS[result_screening_decision(result_data)]


def screening_decision_counts(result_data: dict) -> Dict[str, int]:
    counts = {key: 0 for key in SCREENING_DECISION_ORDER}
    for detail in base.result_details(result_data):
        decision = detail_screening_decision(detail)
        counts[decision] = counts.get(decision, 0) + 1
    return counts


def market_profile_for_result(target: str, result_data: Optional[dict] = None) -> dict:
    profile = _dict((result_data or {}).get("market_profile"))
    fallback = FALLBACK_MARKET_PROFILES.get(target, {})
    merged = dict(fallback)
    merged.update({key: value for key, value in profile.items() if value not in (None, "", [])})
    return merged


def regulatory_reference(detail: dict) -> dict:
    return _dict(detail.get("regulatory_reference"))


def official_term_summary(detail: dict) -> str:
    reference = regulatory_reference(detail)
    term_ko = clean_text(reference.get("official_term_ko"))
    term = clean_text(reference.get("official_term"))
    if term_ko and term and term_ko.casefold() != term.casefold():
        return f"{term_ko} ({term})"
    return term_ko or term or "-"


def official_document_summary(detail: dict) -> str:
    return clean_text(regulatory_reference(detail).get("official_document"), "-")


def entry_reference_summary(detail: dict) -> str:
    return clean_text(regulatory_reference(detail).get("entry_reference"), "-")


def jurisdiction_summary(detail: dict) -> str:
    return clean_text(regulatory_reference(detail).get("jurisdiction"), "-")


def official_reference_text(detail: dict) -> str:
    reference = regulatory_reference(detail)
    explicit = clean_text(reference.get("official_reference_text"))
    if explicit:
        return explicit
    if clean_text(detail.get("match_source")) in {
        "no_database_match",
        "translation_verification",
        "halal_prescreen_verification",
    }:
        return "-"
    return clean_text(
        detail.get("regulation_reason") or detail.get("regulation_notice"),
        "-",
    )


def source_trace_files(detail: dict) -> str:
    values: List[str] = []
    for trace in _list(detail.get("source_trace")):
        if not isinstance(trace, dict):
            continue
        filename = clean_text(trace.get("source_file"))
        if filename and filename not in values:
            values.append(filename)
    return " / ".join(values) or "-"


def source_trace_locations(detail: dict) -> str:
    values: List[str] = []
    for trace in _list(detail.get("source_trace")):
        if not isinstance(trace, dict):
            continue
        sheet = clean_text(trace.get("source_sheet"))
        row = clean_text(trace.get("source_row"))
        location = " · ".join(value for value in [sheet, f"행 {row}" if row else ""] if value)
        if location and location not in values:
            values.append(location)
    return " / ".join(values) or "-"


def source_trace_summary(detail: dict) -> str:
    files = source_trace_files(detail)
    locations = source_trace_locations(detail)
    if files == "-" and locations == "-":
        return "-"
    if locations == "-":
        return files
    if files == "-":
        return locations
    return f"{files} | {locations}"


def official_sources_rows(result_data: dict) -> List[List[str]]:
    rows: List[List[str]] = []
    for source in _list(result_data.get("official_sources")):
        if not isinstance(source, dict):
            continue
        rows.append(
            [
                clean_text(source.get("authority"), "-"),
                clean_text(source.get("title"), "-"),
                clean_text(source.get("url"), "-"),
            ]
        )
    return rows


def official_sources_summary(result_data: dict) -> str:
    values = []
    for authority, title, url in official_sources_rows(result_data):
        value = " — ".join(part for part in [authority, title] if part and part != "-")
        if url and url != "-":
            value = f"{value} | {url}" if value else url
        if value:
            values.append(value)
    return " / ".join(values) or "-"


def normalized_reason(detail: dict, target: str = "") -> str:
    status = internal_status_for_detail(detail)
    source = clean_text(detail.get("match_source"))
    profile = market_profile_for_result(target, {})
    if status == "PASS" and source == "no_database_match":
        if target == "HALAL":
            return profile.get(
                "no_match_interpretation_ko",
                "현재 성분명 기준 HALAL 위험 후보가 확인되지 않았습니다.",
            )
        return profile.get(
            "no_match_interpretation_ko",
            SCREENING_DECISION_DESCRIPTIONS["NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"],
        )
    if status == "VERIFICATION_REQUIRED" and source == "translation_verification":
        return "공식 INCI 명칭 또는 CAS 번호가 확정되지 않아 정확한 규제 대조를 완료하지 못했습니다."
    if status == "VERIFICATION_REQUIRED" and target == "HALAL":
        return "원료 기원·구성·제조공정 또는 인증자료를 공급사 자료로 확인해야 합니다."
    return official_reference_text(detail)


def action_for_detail(detail: dict, target: str = "") -> str:
    provided = clean_text(detail.get("recommended_action"))
    if provided:
        return provided
    return _base_action_for_detail(detail, target)


def market_decision(result_data: dict) -> Dict[str, str]:
    decision = result_screening_decision(result_data)
    label = result_screening_label(result_data)
    profile = _dict(result_data.get("market_profile"))
    if decision == "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED":
        conclusion = clean_text(
            profile.get("no_match_interpretation_ko"),
            SCREENING_DECISION_DESCRIPTIONS[decision],
        )
    else:
        conclusion = SCREENING_DECISION_DESCRIPTIONS[decision]
    return {
        "overall": label,
        "possibility": label,
        "conclusion": conclusion,
        "screening_decision": decision,
    }


def market_impact_for_status(status: str) -> str:
    return SCREENING_DECISION_LABELS[screening_decision_from_status(status)]


def result_action_rows(result_data: dict, target: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for detail in base.attention_details(result_data):
        status = internal_status_for_detail(detail)
        label = halal_status_label(detail) if target == "HALAL" else detail_screening_label(detail)
        rows.append(
            [
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name"), "확인 필요"),
                label,
                action_for_detail(detail, target),
                required_evidence_for_detail(detail, status, target),
                base.completion_criterion(status, target),
            ]
        )
    return rows


def _attention_details(result_data: dict) -> List[dict]:
    return [
        item
        for item in base.result_details(result_data)
        if detail_screening_decision(item) != "NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"
    ]


def market_result_rows(
    result_data: dict,
    source_file: str,
    target: str = "",
) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for index, detail in enumerate(base.result_details(result_data), start=1):
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                detail_screening_label(detail),
                official_term_summary(detail),
                official_document_summary(detail),
                entry_reference_summary(detail),
                jurisdiction_summary(detail),
                normalized_reason(detail, target),
                base.applicable_condition(internal_status_for_detail(detail)),
                action_for_detail(detail, target),
                source_trace_files(detail),
                source_trace_locations(detail),
                base.match_source_label(detail.get("match_source")),
                source_file,
            ]
        )
    return rows


def technical_trace_rows(result_data: dict) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for index, detail in enumerate(base.result_details(result_data), start=1):
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                f"CODE_{internal_status_for_detail(detail)}",
                detail_screening_decision(detail),
                detail_screening_label(detail),
                clean_text(detail.get("match_source"), "-"),
                official_reference_text(detail),
                source_trace_files(detail),
                source_trace_locations(detail),
                clean_text(result_data.get("report_number"), "-"),
                clean_text(result_data.get("api_version"), "-"),
                clean_text(result_data.get("report_schema_version"), "-"),
                clean_text(result_data.get("database_version"), "-"),
                clean_text(result_data.get("database_fingerprint"), "-"),
            ]
        )
    return rows


def summary_required_action(result_data: dict) -> str:
    actions = base.prioritized_action_summaries(result_data, clean_text(result_data.get("target_market")))
    if actions:
        return " / ".join(actions[:6])
    profile = _dict(result_data.get("market_profile"))
    next_checks = _list(profile.get("next_checks_ko"))
    return " / ".join(clean_text(value) for value in next_checks if clean_text(value)) or "성분 변경을 요구하는 즉시 중단 신호는 확인되지 않았습니다."


def market_summary_rows(
    product_name: str,
    target: str,
    result_data: dict,
    halal_result: Optional[dict],
) -> List[List[Any]]:
    profile = market_profile_for_result(target, result_data)
    counts = screening_decision_counts(result_data)
    decision = market_decision(result_data)
    snapshot = _dict(result_data.get("database_snapshot"))
    rows: List[List[Any]] = [
        ["제품명", product_name],
        ["대상 시장", profile.get("display_name_ko") or MARKET_LABELS.get(target, target)],
        ["보고서 번호", clean_text(result_data.get("report_number"), "-")],
        ["검사 성분 수", safe_int(result_data.get("total_checked"))],
        ["종합 스크리닝 해석", decision["overall"]],
        ["해석", decision["conclusion"]],
        ["확인 범위", clean_text(profile.get("screening_scope_ko"), "-")],
        ["우선 중단 요인 확인", counts["IMMEDIATE_STOP_SIGNAL_IDENTIFIED"]],
        ["조건 확인 후 진행 판단", counts["CONDITIONS_REQUIRE_VERIFICATION"]],
        ["추가 정보·원문 확인 후 판단", counts["ADDITIONAL_INFORMATION_REQUIRED"]],
        ["즉시 중단 신호 미확인", counts["NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"]],
        ["필요한 조치", summary_required_action(result_data)],
        ["다음 확인사항", " / ".join(_list(profile.get("next_checks_ko"))) or "-"],
        ["공식 출처", official_sources_summary(result_data)],
        ["DB 버전", clean_text(result_data.get("database_version"), "-")],
        ["확정 DB", clean_text(snapshot.get("confirmed_file") or result_data.get("database_file"), "-")],
        ["확정 DB SHA256", clean_text(snapshot.get("confirmed_sha256"), "-")],
        ["감사자료", clean_text(snapshot.get("audit_file") or result_data.get("audit_database_file"), "-")],
        ["감사자료 SHA256", clean_text(snapshot.get("audit_sha256"), "-")],
        ["보고서 생성일", clean_text(result_data.get("report_generated_at"), datetime.now().isoformat(timespec="seconds"))],
        ["사용 범위", "국가·시장별 공식 규제자료와의 화장품 성분 사전 스크리닝"],
    ]
    if halal_result is not None:
        relation = base.halal_market_relation(target)
        hdecision = halal_decision(halal_result)
        rows.extend(
            [
                ["HALAL 추가검토", "별도 분석"],
                ["HALAL 관련성", relation["relevance"]],
                ["이번 HALAL 검토 결과", hdecision["overall"]],
            ]
        )
    return rows


def manual_rows(
    result_data: dict,
    target: str,
    halal_result: Optional[dict] = None,
) -> List[List[Any]]:
    rows: List[List[Any]] = []
    index = 1
    for detail in _attention_details(result_data):
        status = internal_status_for_detail(detail)
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                detail_screening_label(detail),
                official_term_summary(detail),
                normalized_reason(detail, target),
                base.applicable_condition(status),
                required_evidence_for_detail(detail, status, target),
                "",
                "",
                "",
                "",
            ]
        )
        index += 1
    return rows


def _write_generic_sheet(
    ws: Any,
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    widths: Sequence[float],
    highlight_col: Optional[int] = None,
) -> None:
    base._xlsx_title(ws, title, len(headers))
    base._xlsx_header(ws, 3, headers)
    materialized = list(rows)
    if not materialized:
        materialized = [["해당 항목 없음"] + [""] * (len(headers) - 1)]
    for row_index, row in enumerate(materialized, start=4):
        for col_index, value in enumerate(row, start=1):
            ws.cell(row_index, col_index, value)
    end_row = 3 + len(materialized)
    base._xlsx_body(ws, 4, end_row, len(headers))
    ws.freeze_panes = "A4"
    ws.auto_filter.ref = f"A3:{get_column_letter(len(headers))}{end_row}"
    base._set_widths(ws, widths)
    if highlight_col is not None:
        for row_index in range(4, end_row + 1):
            text = clean_text(ws.cell(row_index, highlight_col + 1).value)
            fill = None
            if "중단" in text:
                fill = base.RED
            elif any(term in text for term in ("조건", "추가 정보", "원문")):
                fill = base.YELLOW
            elif "미확인" in text:
                fill = base.GREEN
            if fill:
                cell = ws.cell(row_index, highlight_col + 1)
                cell.fill = PatternFill("solid", fgColor=fill)
                cell.font = Font(bold=True, size=9)


def _write_manual_sheet(ws: Any, title: str, rows: Sequence[Sequence[Any]]) -> None:
    headers = [
        "번호",
        "성분명",
        "INCI 명칭",
        "CAS 번호",
        "현재 스크리닝 해석",
        "공식 규제 항목",
        "확인 사유·공식 근거",
        "확인할 사항",
        "필요한 자료",
        "확인 결과",
        "최종 검토결론",
        "검토자",
        "검토일",
    ]
    _write_generic_sheet(
        ws,
        title,
        headers,
        rows,
        [7, 18, 26, 14, 25, 28, 42, 34, 30, 22, 22, 14, 14],
        4,
    )
    validation = base.DataValidation(
        type="list",
        formula1='"확인 중,진행 가능,조건부 진행,중단,재검토 필요"',
        allow_blank=True,
    )
    ws.add_data_validation(validation)
    validation.add(f"K4:K{max(ws.max_row, 100)}")


def create_product_excel_bytes(
    product_name: str,
    source_file: str,
    market_results_map: Dict[str, dict],
    selected_markets: Sequence[str],
    halal_result: Optional[dict] = None,
    language_code: str = "ko",
    translator: Optional[base.Translator] = None,
) -> bytes:
    workbook = Workbook()
    workbook.remove(workbook.active)
    sheet_number = 1

    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue
        profile = market_profile_for_result(target, result_data)
        label = clean_text(profile.get("display_name_ko"), MARKET_LABELS.get(target, target))

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_규제분석결과")
        _write_generic_sheet(
            ws,
            f"{label} 국가·시장별 규제 스크리닝 결과",
            [
                "번호",
                "입력 성분명",
                "INCI 명칭",
                "CAS 번호",
                "사업상 스크리닝 해석",
                "공식 규제 항목",
                "공식 문서",
                "참조 번호·항목",
                "관할",
                "공식 근거문구",
                "적용 조건·추가 확인",
                "권고 조치",
                "원문 파일",
                "원문 위치",
                "매칭 방식",
                "입력 파일",
            ],
            market_result_rows(result_data, source_file, target),
            [7, 18, 26, 14, 28, 30, 32, 20, 16, 46, 34, 38, 38, 24, 22, 24],
            4,
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_종합요약")
        base._write_summary_sheet(
            ws,
            f"{label} 종합요약",
            market_summary_rows(product_name, target, result_data, halal_result),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_기술추적")
        _write_generic_sheet(
            ws,
            f"{label} 시스템 기술추적 정보",
            [
                "번호",
                "입력 성분명",
                "INCI 명칭",
                "CAS 번호",
                "시스템 내부 상태 코드",
                "스크리닝 결정 코드",
                "고객 표시 해석",
                "매칭 소스 코드",
                "공식 참조문구",
                "원문 파일",
                "원문 위치",
                "보고서 번호",
                "API 버전",
                "보고서 스키마",
                "DB 버전",
                "DB 지문",
            ],
            technical_trace_rows(result_data),
            [7, 18, 26, 14, 20, 32, 28, 24, 46, 38, 24, 30, 22, 24, 24, 30],
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_수동확인")
        _write_manual_sheet(ws, f"{label} 추가확인 체크리스트", manual_rows(result_data, target))
        sheet_number += 1

    if halal_result is not None:
        ws = workbook.create_sheet(f"{sheet_number:02d}_HALAL_성분검토결과")
        base._write_result_sheet(
            ws,
            "HALAL 추가 성분 검토 결과",
            [
                "번호",
                "입력 성분명",
                "INCI 명칭",
                "CAS 번호",
                "할랄 검토 상태",
                "검출 방식",
                "우려 유형",
                "확인 사유",
                "필요한 증빙",
                "공급사 확인사항",
                "조치 완료 기준",
                "입력 파일",
            ],
            base.halal_result_rows(halal_result, source_file),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_HALAL_종합요약")
        base._write_summary_sheet(ws, "HALAL 종합요약", base.halal_summary_rows(product_name, halal_result))
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_HALAL_수동확인")
        base._write_manual_sheet(ws, "HALAL 수동확인 체크리스트", base.halal_manual_rows(halal_result))

    output = io.BytesIO()
    workbook.save(output)
    data = output.getvalue()
    protected_terms = base._result_protected_terms(
        product_name,
        source_file,
        market_results_map,
        halal_result,
    )
    return base._localize_excel_bytes(data, language_code, translator, protected_terms)


def _market_attention_rows(result_data: dict, target: str) -> List[List[str]]:
    details = base.result_details(result_data)
    chosen = details if len(details) <= 18 else _attention_details(result_data)
    rows: List[List[str]] = []
    for detail in chosen:
        reference = official_reference_text(detail)
        official_item = official_term_summary(detail)
        if entry_reference_summary(detail) != "-":
            official_item = f"{official_item} | {entry_reference_summary(detail)}"
        rows.append(
            [
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name"), "확인 필요"),
                detail_screening_label(detail),
                official_item,
                reference,
            ]
        )
    if len(details) > 18 and not rows:
        rows.append(
            [
                "전체 성분",
                "Excel 참조",
                SCREENING_DECISION_LABELS["NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"],
                "-",
                "상세 성분별 결과는 Excel 규제분석결과 시트를 참조하십시오.",
            ]
        )
    return rows


def _manual_report_rows(result_data: dict, target: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for detail in _attention_details(result_data):
        rows.append(
            [
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name"), "확인 필요"),
                detail_screening_label(detail),
                official_term_summary(detail),
                action_for_detail(detail, target),
            ]
        )
    return rows


def _screening_interpretation_rows(include_halal: bool = False) -> List[List[str]]:
    rows = [
        [
            SCREENING_DECISION_LABELS["IMMEDIATE_STOP_SIGNAL_IDENTIFIED"],
            "우선 중단하고 공식 적용범위를 확인",
            SCREENING_DECISION_DESCRIPTIONS["IMMEDIATE_STOP_SIGNAL_IDENTIFIED"],
        ],
        [
            SCREENING_DECISION_LABELS["CONDITIONS_REQUIRE_VERIFICATION"],
            "조건을 확인한 뒤 진행 여부 판단",
            SCREENING_DECISION_DESCRIPTIONS["CONDITIONS_REQUIRE_VERIFICATION"],
        ],
        [
            SCREENING_DECISION_LABELS["ADDITIONAL_INFORMATION_REQUIRED"],
            "추가 자료 없이는 결론 보류",
            SCREENING_DECISION_DESCRIPTIONS["ADDITIONAL_INFORMATION_REQUIRED"],
        ],
        [
            SCREENING_DECISION_LABELS["NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"],
            "현재 탑재 범위에서 즉시 중단 신호 없음",
            SCREENING_DECISION_DESCRIPTIONS["NO_IMMEDIATE_STOP_SIGNAL_IDENTIFIED"],
        ],
    ]
    if include_halal:
        rows.append(
            [
                "HALAL 위험 후보 미확인",
                "성분명 기준 후보가 확인되지 않음",
                "HALAL 인증 또는 최종 적합판정이 아니며 원료 기원·공정·교차오염·인증자료는 별도 확인합니다.",
            ]
        )
    return rows


def create_product_report_bytes(
    product_name: str,
    source_file: str,
    market_results_map: Dict[str, dict],
    selected_markets: Sequence[str],
    halal_result: Optional[dict] = None,
    language_code: str = "ko",
    translator: Optional[base.Translator] = None,
) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = base.DOC_TEXT
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(3)

    market_labels = []
    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue
        profile = market_profile_for_result(target, result_data)
        market_labels.append(clean_text(profile.get("display_name_ko"), MARKET_LABELS.get(target, target)))
    scope = " · ".join(market_labels)
    if halal_result is not None:
        scope += " + HALAL 추가검토"

    base._chapter_header(document, 1, "제품 및 분석 개요", f"{product_name} · {scope}")
    first_result = next(iter(market_results_map.values()), halal_result or {})
    ingredients = [
        [
            clean_text(item.get("original_ingredient")),
            clean_text(item.get("inci_name")),
            clean_text(item.get("cas_number"), "N/A"),
        ]
        for item in base.result_details(first_result)
    ]
    report_numbers = [
        clean_text(value.get("report_number"))
        for value in market_results_map.values()
        if clean_text(value.get("report_number"))
    ]
    overview_rows = [
        ["제품명", product_name],
        ["입력 파일", source_file],
        ["분석 시장", " · ".join(market_labels) or "없음"],
        ["HALAL 추가검토", "선택함" if halal_result is not None else "선택하지 않음"],
        ["분석 성분 수", len(ingredients)],
        ["보고서 식별번호", " / ".join(report_numbers) or "-"],
    ]
    base._add_simple_table(document, ["항목", "내용"], overview_rows, base.BLUE, 9.0)

    base._section_title(document, "1. 제품 성분")
    if len(ingredients) <= 24:
        base._add_simple_table(document, ["입력 성분명", "INCI 명칭", "CAS 번호"], ingredients, base.NAVY, 8.8)
    else:
        base._bullet(document, f"전체 {len(ingredients)}개 성분을 분석했습니다. 전체 성분 목록은 함께 제공되는 Excel 파일에서 확인할 수 있습니다.")

    base._section_title(document, "2. 시장별 스크리닝 요약")
    summary_rows: List[List[str]] = []
    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue
        profile = market_profile_for_result(target, result_data)
        decision = market_decision(result_data)
        summary_rows.append(
            [
                clean_text(profile.get("display_name_ko"), target),
                decision["overall"],
                clean_text(profile.get("screening_scope_ko"), "-"),
            ]
        )
    if halal_result is not None:
        hdecision = halal_decision(halal_result)
        summary_rows.append(["HALAL 추가검토", hdecision["overall"], FALLBACK_MARKET_PROFILES["HALAL"]["screening_scope_ko"]])
    base._add_simple_table(document, ["시장·검토", "사업상 스크리닝 해석", "확인 범위"], summary_rows, base.BLUE, 8.4, 1)

    base._section_title(document, "3. 보고서 사용 범위")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        "이 보고서는 각 시장의 탑재 규제자료를 제품 전성분과 빠르게 대조하여 시장 진입 가능성을 검토할 기초정보와 다음 행동을 제공합니다. "
        "법률 자문, 제품 등록 승인, 통관 허가, 판매 보증 또는 HALAL 인증을 대신하지 않습니다."
    )
    run.font.size = Pt(9.2)
    run.font.color.rgb = base.DOC_MUTED

    base._section_title(document, "4. 스크리닝 해석 기준")
    base._add_simple_table(
        document,
        ["스크리닝 해석", "짧은 설명", "정확한 의미"],
        _screening_interpretation_rows(include_halal=halal_result is not None),
        base.NAVY,
        7.7,
        0,
    )

    chapter_number = 2
    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue
        profile = market_profile_for_result(target, result_data)
        market_label = clean_text(profile.get("display_name_ko"), MARKET_LABELS.get(target, target))

        anchor = document.add_paragraph()
        anchor.paragraph_format.page_break_before = True
        anchor.paragraph_format.space_after = Pt(0)
        base._chapter_header(
            document,
            chapter_number,
            f"{market_label} 규제 스크리닝 결과",
            "공식 규제 항목과 사업상 해석을 분리하여 표시합니다.",
        )
        chapter_number += 1
        decision = market_decision(result_data)
        base._decision_box(document, "이번 스크리닝 해석", decision["overall"])

        base._section_title(document, "1. 확인 범위와 데이터 기준")
        metadata_rows = [
            ["확인 범위", clean_text(profile.get("screening_scope_ko"), "-")],
            ["보고서 번호", clean_text(result_data.get("report_number"), "-")],
            ["DB 버전", clean_text(result_data.get("database_version"), "-")],
            ["보고서 스키마", clean_text(result_data.get("report_schema_version"), "-")],
            ["보고서 생성일", clean_text(result_data.get("report_generated_at"), "-")],
        ]
        base._add_simple_table(document, ["항목", "내용"], metadata_rows, base.BLUE, 8.6)

        source_rows = official_sources_rows(result_data)
        if source_rows:
            base._add_simple_table(document, ["기관", "공식 자료", "URL"], source_rows, base.NAVY, 7.8)

        base._section_title(document, "2. 성분별 스크리닝 결과")
        rows = _market_attention_rows(result_data, target)
        base._add_simple_table(
            document,
            ["입력 성분명", "INCI Name", "사업상 해석", "공식 규제 항목", "공식 근거·참조"],
            rows,
            base.NAVY,
            7.6,
            2,
        )

        base._section_title(document, "3. 필요한 조치")
        action_rows = result_action_rows(result_data, target)
        if action_rows:
            base._add_simple_table(
                document,
                ["입력 성분명", "INCI Name", "스크리닝 해석", "해야 할 일", "필요한 자료", "완료 기준"],
                action_rows,
                base.BLUE,
                7.5,
                2,
            )
        else:
            base._bullet(document, "현재 탑재 범위에서는 성분 변경을 요구하는 즉시 중단 신호가 확인되지 않았습니다.")
            for item in _list(profile.get("next_checks_ko"))[:3]:
                base._bullet(document, clean_text(item))

        section_index = 4
        manual = _manual_report_rows(result_data, target)
        if manual:
            base._section_title(document, f"{section_index}. 추가 확인 항목")
            section_index += 1
            base._add_simple_table(
                document,
                ["입력 성분명", "INCI Name", "스크리닝 해석", "공식 규제 항목", "필요한 조치"],
                manual,
                base.BLUE,
                7.7,
                2,
            )

        if halal_result is not None:
            base._section_title(document, f"{section_index}. HALAL 검토와의 관계")
            section_index += 1
            relation = base.halal_market_relation(target)
            base._add_simple_table(
                document,
                ["질문", "답변"],
                [
                    ["이 시장에서 직접 필요한가?", relation["direct_required"]],
                    ["필요할 수 있는 경우", relation["required_when"]],
                    ["성분 규제 결과가 달라지는가?", relation["screening_effect"]],
                ],
                base.BLUE,
                8.4,
            )

        base._section_title(document, f"{section_index}. 다음 단계")
        next_checks = _list(profile.get("next_checks_ko"))
        for step in next_checks[:4]:
            if clean_text(step):
                base._bullet(document, clean_text(step))
        if not next_checks:
            for step in base.general_next_steps(target)[:3]:
                base._bullet(document, step)

    if halal_result is not None:
        anchor = document.add_paragraph()
        anchor.paragraph_format.page_break_before = True
        anchor.paragraph_format.space_after = Pt(0)
        base._chapter_header(
            document,
            chapter_number,
            "HALAL 추가 성분 검토",
            "국가별 화장품 성분 규제와 별도로 원료 기원·구성·공정·인증자료 후보를 선별합니다.",
        )
        hdecision = halal_decision(halal_result)
        base._decision_box(document, "이번 HALAL 검토 결과", hdecision["possibility"])
        rows = []
        for detail in base.result_details(halal_result):
            status = internal_status_for_detail(detail)
            rows.append(
                [
                    clean_text(detail.get("original_ingredient")),
                    clean_text(detail.get("inci_name"), "확인 필요"),
                    halal_status_label(detail),
                    "-" if status == "PASS" else clean_text(detail.get("verification_type"), base.manual_type(detail, "HALAL")),
                    required_evidence_for_detail(detail, status, "HALAL"),
                ]
            )
        base._section_title(document, "1. 성분별 검토 결과")
        base._add_simple_table(document, ["입력 성분명", "INCI Name", "검토 결과", "확인 이유", "필요한 자료"], rows, base.NAVY, 7.9, 2)
        base._section_title(document, "2. 중요 안내")
        base._bullet(document, "이 검토 결과는 HALAL 인증 완료 또는 특정 인증기관의 적합판정을 의미하지 않습니다.")
        base._bullet(document, "최종 판단은 공급사 증빙, 제조공정, 교차오염 관리 및 인증기관 심사를 통해 이루어집니다.")

    for document_section in document.sections:
        footer = document_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("K-Beauty Global Compliance | 국가·시장별 규제 사전 스크리닝")
        run.font.size = Pt(8)
        run.font.color.rgb = base.DOC_MUTED

    output = io.BytesIO()
    document.save(output)
    data = output.getvalue()
    protected_terms = base._result_protected_terms(
        product_name,
        source_file,
        market_results_map,
        halal_result,
    )
    return base._localize_docx_bytes(data, language_code, translator, protected_terms)


# Sheet-name localization: technical trace needs a distinct suffix so it cannot
# collide with the main result sheet in non-Korean output.
for language_code, suffixes in base.SHEET_SUFFIXES.items():
    suffixes.setdefault(
        "technical",
        {
            "ko": "기술추적",
            "en": "Technical_Trace",
            "zh-CN": "技术追踪",
            "ar": "تتبع_تقني",
            "fr": "Traçabilité",
            "id": "Jejak_Teknis",
            "ms": "Jejak_Teknikal",
            "sw": "Ufuatiliaji",
        }.get(language_code, "Technical_Trace"),
    )


def _localized_sheet_name(original_name: str, language_code: str) -> str:
    if language_code == "ko":
        return original_name[:31]
    prefix = original_name.split("_", 1)[0]
    if "HALAL" in original_name:
        if "수동확인" in original_name:
            key = "halal_manual"
        elif "종합요약" in original_name:
            key = "halal_summary"
        else:
            key = "halal_result"
        return f"{prefix}_{base.SHEET_SUFFIXES[language_code][key]}"[:31]
    parts = original_name.split("_")
    target = parts[1] if len(parts) > 1 else "MARKET"
    if "기술추적" in original_name:
        key = "technical"
    elif "수동확인" in original_name:
        key = "manual"
    elif "종합요약" in original_name:
        key = "summary"
    else:
        key = "result"
    return f"{prefix}_{target}_{base.SHEET_SUFFIXES[language_code][key]}"[:31]


base.ENGLISH_CANONICAL_EXACT.update(
    {
        "사업상 스크리닝 해석": "Business Screening Interpretation",
        "우선 중단 요인 확인": "Immediate Stop Signal Identified",
        "조건 확인 후 진행 판단": "Proceed After Verifying Conditions",
        "추가 정보·원문 확인 후 판단": "Decision Pending Additional Information and Source Verification",
        "현재 확인 범위에서 즉시 중단 신호 미확인": "No Immediate Stop Signal Identified Within the Current Screening Scope",
        "공식 규제 항목": "Official Regulatory Entry",
        "공식 문서": "Official Document",
        "참조 번호·항목": "Entry or Reference",
        "관할": "Jurisdiction",
        "공식 근거문구": "Official Reference Text",
        "적용 조건·추가 확인": "Conditions and Additional Verification",
        "원문 파일": "Source File",
        "원문 위치": "Source Location",
        "시스템 내부 상태 코드": "Internal System Status Code",
        "스크리닝 결정 코드": "Screening Decision Code",
        "고객 표시 해석": "Customer-Facing Interpretation",
        "매칭 소스 코드": "Match Source Code",
        "보고서 번호": "Report Number",
        "보고서 식별번호": "Report Identifier",
        "DB 버전": "Database Version",
        "DB 지문": "Database Fingerprint",
        "기술추적": "Technical Trace",
        "확인 범위와 데이터 기준": "Screening Scope and Data Basis",
        "국가·시장별 규제 스크리닝 결과": "Country- and Market-Specific Regulatory Screening Results",
        "국가·시장별 공식 규제자료와의 화장품 성분 사전 스크리닝": "Preliminary Cosmetic Ingredient Screening Against Country- and Market-Specific Official Regulatory Data",
    }
)

# Technical identifiers are not sentence-like English and must remain unchanged
# in localized reports. The inherited r3 heuristic counts alphabetic runs, so a
# SHA fingerprint or several slash-separated report identifiers can otherwise be
# misclassified as untranslated prose.
_original_needs_korean_translation = base._needs_korean_translation

_HEX_FINGERPRINT_RE = re.compile(r"[0-9A-Fa-f]{32,128}")
_UUID_RE = re.compile(
    r"[0-9A-Fa-f]{8}-[0-9A-Fa-f]{4}-[0-9A-Fa-f]{4}-"
    r"[0-9A-Fa-f]{4}-[0-9A-Fa-f]{12}"
)
_REPORT_IDENTIFIER_RE = re.compile(
    r"KBR-(?:[A-Z]{2,12}|[ \t]*)-"
    r"\d{8}T\d{6}Z-[0-9A-Fa-f]{8,64}"
)
_DATABASE_VERSION_RE = re.compile(
    r"KBRDB-(?:[A-Z0-9]{2,12}|[ \t]*)-[0-9A-Fa-f]{8,64}"
)


def _is_technical_identifier_component(value: str) -> bool:
    return any(
        pattern.fullmatch(value)
        for pattern in (
            _HEX_FINGERPRINT_RE,
            _UUID_RE,
            _REPORT_IDENTIFIER_RE,
            _DATABASE_VERSION_RE,
        )
    )


def _is_technical_identifier_sequence(value: str) -> bool:
    # Summary cells can contain several report IDs separated by slashes. During
    # translation validation protected market tokens may be hidden, producing
    # forms such as ``KBR- -20260627T172637Z-1DB54743``; the report-ID pattern
    # intentionally accepts that blank market-code slot as well as US/EU/EAC.
    parts = [
        clean_text(part)
        for part in re.split(r"\s*(?:/|\||;|,)\s*", value)
    ]
    return bool(parts) and all(
        part and _is_technical_identifier_component(part)
        for part in parts
    )


def _needs_korean_translation(text: str) -> bool:
    value = clean_text(text)
    if _is_technical_identifier_sequence(value):
        return False
    return _original_needs_korean_translation(value)


base._needs_korean_translation = _needs_korean_translation


# Make r3 bundle/localization helpers resolve the r4 implementations at runtime.
base.market_decision = market_decision
base.normalized_reason = normalized_reason
base.action_for_detail = action_for_detail
base.market_impact_for_status = market_impact_for_status
base.result_action_rows = result_action_rows
base.market_result_rows = market_result_rows
base.market_summary_rows = market_summary_rows
base.manual_rows = manual_rows
base.create_product_excel_bytes = create_product_excel_bytes
base.create_product_report_bytes = create_product_report_bytes
base._localized_sheet_name = _localized_sheet_name

create_bundle_zip_bytes = base.create_bundle_zip_bytes
