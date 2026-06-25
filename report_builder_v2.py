from __future__ import annotations

import csv
import io
import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation


MARKET_LABELS: Dict[str, str] = {
    "US": "미국(US)",
    "EU": "유럽연합(EU)",
    "UK": "영국(UK)",
    "CN": "중국(CN)",
    "ASEAN": "아세안(ASEAN)",
    "SFDA": "사우디아라비아(SFDA)",
    "EAC": "유라시아경제연합(EAC)",
    "HALAL": "HALAL 추가 성분 검토",
}

STATUS_LABELS: Dict[str, str] = {
    "PASS": "통과",
    "BANNED": "금지",
    "RESTRICTED": "제한",
    "WARNING_REQUIRED": "경고·표시 필요",
    "REGULATED": "규제 대상",
    "REVIEW_REQUIRED": "수동 검토",
    "VERIFICATION_REQUIRED": "명칭 검증 필요",
    "FAIL": "부적합",
}

STATUS_ORDER = [
    "BANNED",
    "REVIEW_REQUIRED",
    "VERIFICATION_REQUIRED",
    "REGULATED",
    "RESTRICTED",
    "WARNING_REQUIRED",
    "PASS",
]

MANUAL_STATUSES = {
    "REVIEW_REQUIRED",
    "VERIFICATION_REQUIRED",
    "REGULATED",
}

NAVY = "294A63"
BLUE = "5AA6C9"
LIGHT_BLUE = "DDEEF5"
PALE = "F5F7F8"
GRAY = "E9EDF0"
YELLOW = "FFF4CC"
RED = "F4CCCC"
GREEN = "D9EAD3"
WHITE = "FFFFFF"
BORDER_COLOR = "C8CDD1"

DOC_TEXT = RGBColor(45, 45, 45)
DOC_MUTED = RGBColor(105, 105, 105)
DOC_NAVY = RGBColor(41, 74, 99)
DOC_GOLD = RGBColor(92, 72, 0)


# ============================================================
# 공통 유틸리티
# ============================================================
def clean_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    try:
        if value != value:
            return default
    except Exception:
        pass
    text = " ".join(str(value).strip().split())
    if text.casefold() in {"", "nan", "none", "null"}:
        return default
    return text


def safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return default


def normalize_status(value: Any) -> str:
    status = clean_text(value).upper()
    aliases = {
        "FAIL": "BANNED",
        "MANUAL_REVIEW": "REVIEW_REQUIRED",
        "MANUAL REVIEW": "REVIEW_REQUIRED",
        "VERIFICATION": "VERIFICATION_REQUIRED",
        "WARNING": "WARNING_REQUIRED",
    }
    return aliases.get(status, status or "VERIFICATION_REQUIRED")


def safe_filename(value: str, fallback: str = "제품") -> str:
    value = clean_text(value, fallback)
    value = re.sub(r"[\\/:*?\"<>|]+", "_", value)
    value = re.sub(r"\s+", "_", value).strip("._")
    return value or fallback


def dedupe_preserve(values: Iterable[str]) -> List[str]:
    result: List[str] = []
    seen: set[str] = set()
    for value in values:
        text = clean_text(value)
        key = text.casefold()
        if text and key not in seen:
            seen.add(key)
            result.append(text)
    return result


def result_details(result_data: dict) -> List[dict]:
    details = result_data.get("report_details") or []
    return [item for item in details if isinstance(item, dict)]


def status_counts(result_data: dict) -> Dict[str, int]:
    raw = result_data.get("status_counts") or {}
    return {
        status: safe_int(raw.get(status))
        for status in (
            "PASS",
            "BANNED",
            "RESTRICTED",
            "WARNING_REQUIRED",
            "REGULATED",
            "REVIEW_REQUIRED",
            "VERIFICATION_REQUIRED",
        )
    }


def dominant_status(result_data: dict) -> str:
    counts = status_counts(result_data)
    for status in STATUS_ORDER:
        if counts.get(status, 0) > 0:
            return status
    return "PASS"


def market_decision(result_data: dict) -> Dict[str, str]:
    status = dominant_status(result_data)
    if status == "BANNED":
        return {
            "overall": "부적합",
            "possibility": "현재 처방으로 진행 불가",
            "conclusion": "현재 처방 기준 해당 시장 수입·유통·판매 진행 불가",
        }
    if status in {"REVIEW_REQUIRED", "VERIFICATION_REQUIRED"}:
        return {
            "overall": "확인 필요",
            "possibility": "가능성은 있으나 확인 완료 전 진행 보류",
            "conclusion": "금지 성분은 확인되지 않았으나 수동 검토 또는 명칭 검증이 완료되기 전까지 수입·유통·판매 진행 여부를 확정할 수 없습니다.",
        }
    if status == "REGULATED":
        return {
            "overall": "적용 범위 확인 필요",
            "possibility": "적용 범위 확인 후 결정",
            "conclusion": "관련 규제의 실제 제품 적용 범위를 확인한 뒤 시장 진입 진행 여부를 결정해야 합니다.",
        }
    if status == "RESTRICTED":
        return {
            "overall": "조건부 가능",
            "possibility": "제한조건 충족 후 가능",
            "conclusion": "금지 성분은 확인되지 않았으며, 농도·제품유형·사용조건 등 제한요건을 충족하면 시장 진입 준비가 가능합니다.",
        }
    if status == "WARNING_REQUIRED":
        return {
            "overall": "조건부 가능",
            "possibility": "표시·고지요건 충족 후 가능",
            "conclusion": "필요한 라벨·경고문구·고지·통지 요건을 충족하면 시장 진입 준비가 가능합니다.",
        }
    return {
        "overall": "성분 규제 단계 통과",
        "possibility": "시장 진입 준비 가능",
        "conclusion": "현재 로드된 규제 데이터 기준 금지·제한·추가 확인 대상이 확인되지 않아 성분 규제 단계상 시장 진입 준비가 가능합니다.",
    }


def halal_decision(result_data: dict) -> Dict[str, str]:
    status = dominant_status(result_data)
    if status == "BANNED":
        return {
            "overall": "할랄 금지 성분 확인",
            "possibility": "현재 처방 기준 할랄 인증 추진 곤란",
            "conclusion": "확정 할랄 금지 성분이 확인되어 현재 처방으로 할랄 인증을 추진하기 어렵습니다. 원료 대체 후 재검토가 필요합니다.",
        }
    if status in {"REVIEW_REQUIRED", "VERIFICATION_REQUIRED", "REGULATED"}:
        return {
            "overall": "증빙자료 확인 필요",
            "possibility": "원료 기원·구성·공정·인증자료 확인 전 판단 보류",
            "conclusion": "현재 확인된 확정 금지 성분은 없으나 원료 기원, 복합원료 구성, 제조공정 또는 인증자료 확인이 필요하여 할랄 인증 추진 가능성을 아직 확정할 수 없습니다.",
        }
    if status == "RESTRICTED":
        return {
            "overall": "조건 확인 필요",
            "possibility": "조건 확인 후 인증 준비 가능",
            "conclusion": "확정 금지 성분은 확인되지 않았으며 관련 사용조건과 인증기관 요구사항을 충족하면 할랄 인증 준비가 가능합니다.",
        }
    return {
        "overall": "할랄 위험 후보 미확인",
        "possibility": "인증 준비 가능",
        "conclusion": "현재 분석 데이터에서 할랄 위험 후보가 확인되지 않았으나, 완제품의 할랄 인증 적합성 또는 인증 완료를 의미하지는 않습니다.",
    }


def market_impact_for_status(status: str) -> str:
    status = normalize_status(status)
    return {
        "PASS": "시장 진입 가능성 있음",
        "BANNED": "현재 처방 진행 불가",
        "RESTRICTED": "제한조건 충족 후 가능",
        "WARNING_REQUIRED": "표시·고지요건 충족 후 가능",
        "REGULATED": "적용 범위 확인 후 결정",
        "REVIEW_REQUIRED": "확인 완료 전 보류",
        "VERIFICATION_REQUIRED": "확인 완료 전 보류",
    }.get(status, "확인 필요")


def applicable_condition(status: str) -> str:
    status = normalize_status(status)
    return {
        "PASS": "-",
        "BANNED": "금지 적용 범위, 동일물질 여부 및 예외조항 확인",
        "RESTRICTED": "실제 배합농도, 제품유형, 사용부위, 사용대상 및 용도 확인",
        "WARNING_REQUIRED": "라벨, 경고문구, 고지, 신고 또는 통지 요건 확인",
        "REGULATED": "해당 규정의 제품 적용 범위와 사용목적 확인",
        "REVIEW_REQUIRED": "공식 규정 원문, 사용조건 및 최신 시행상태 확인",
        "VERIFICATION_REQUIRED": "공식 INCI 명칭, CAS 번호 및 공급사 원료자료 확인",
    }.get(status, "추가 확인")


def recommended_action(status: str) -> str:
    status = normalize_status(status)
    return {
        "PASS": "별도 조치 없음",
        "BANNED": "처방 변경 또는 대체원료 검토 후 재분석",
        "RESTRICTED": "실제 처방조건과 규정 제한조건을 대조",
        "WARNING_REQUIRED": "필수 표시·경고·통지 문구를 라벨에 반영",
        "REGULATED": "규제 적용범위를 확인한 뒤 진행 여부 결정",
        "REVIEW_REQUIRED": "공식 원문 및 전문가 검토 완료 후 재판정",
        "VERIFICATION_REQUIRED": "공식 명칭과 공급사 자료를 확보한 뒤 재분석",
    }.get(status, "추가 확인")


def manual_type(detail: dict, target: str) -> str:
    status = normalize_status(
        detail.get("compliance_status")
        or detail.get("restriction_type")
    )
    if target == "HALAL":
        return clean_text(
            detail.get("verification_type"),
            {
                "REVIEW_REQUIRED": "공식 할랄 기준",
                "VERIFICATION_REQUIRED": "원료 기원·구성·인증자료",
                "REGULATED": "할랄 기준 적용범위",
                "BANNED": "금지 성분 확인",
                "RESTRICTED": "조건 확인",
            }.get(status, "추가 확인"),
        )
    return {
        "REVIEW_REQUIRED": "규제 원문·적용조건",
        "VERIFICATION_REQUIRED": "공식 명칭·CAS",
        "REGULATED": "규제 적용범위",
    }.get(status, "추가 확인")


def halal_status_label(detail: dict) -> str:
    status = normalize_status(
        detail.get("compliance_status")
        or detail.get("restriction_type")
    )
    if status == "PASS":
        return "할랄 위험 후보 미확인"
    if status == "BANNED":
        return "할랄 금지 성분 확인"
    if status == "RESTRICTED":
        return "조건 확인 필요"
    if status == "WARNING_REQUIRED":
        return "표시·인증 확인 필요"
    if status == "REGULATED":
        return "공식 기준 적용범위 확인"
    if status == "REVIEW_REQUIRED":
        return "공식 기준 수동 검토"
    if status == "VERIFICATION_REQUIRED":
        verification_type = clean_text(detail.get("verification_type"))
        if verification_type:
            return verification_type
        if clean_text(detail.get("match_source")) == "translation_verification":
            return "명칭·구성 확인 필요"
        return "원료 기원·구성·인증자료 확인 필요"
    return "추가 확인 필요"


def halal_market_relation(target: str) -> Dict[str, str]:
    target = clean_text(target).upper()
    if target == "SFDA":
        return {
            "level": "직접 관련 가능성이 높음",
            "impact": "현지 수입자·유통사·인증 요구에 따라 원료 기원 및 인증자료가 시장 진입 준비에 직접 영향을 줄 수 있습니다.",
        }
    if target == "ASEAN":
        return {
            "level": "국가별 직접 또는 조건부 관련",
            "impact": "ASEAN 권역 중 인도네시아·말레이시아 등에서는 할랄 인증 또는 관련 자료가 직접 영향을 줄 수 있으며, 기타 국가는 바이어·유통채널 요구에 따라 달라집니다.",
        }
    if target == "EAC":
        return {
            "level": "조건부 관련",
            "impact": "국가 화장품 규제판정에는 직접 영향을 주지 않지만 현지 바이어·유통채널 또는 할랄 표시 전략에 따라 추가 자료가 요구될 수 있습니다.",
        }
    return {
        "level": "별도 인증 목적 또는 조건부 관련",
        "impact": "해당 시장의 국가 화장품 규제판정을 변경하지는 않지만, 할랄 제품으로 판매하거나 바이어·유통채널이 인증자료를 요구하는 경우 시장 진입 준비에 영향을 줄 수 있습니다.",
    }


def attention_details(result_data: Optional[dict]) -> List[dict]:
    if not result_data:
        return []
    return [
        item
        for item in result_details(result_data)
        if normalize_status(
            item.get("compliance_status")
            or item.get("restriction_type")
        ) != "PASS"
    ]


def linked_halal_ingredients(halal_result: Optional[dict]) -> str:
    names = [
        clean_text(item.get("original_ingredient"))
        for item in attention_details(halal_result)
    ]
    return ", ".join(dedupe_preserve(names)) or "없음"


# ============================================================
# Excel 데이터 행 생성
# ============================================================
def market_result_rows(
    result_data: dict,
    source_file: str,
) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for index, detail in enumerate(result_details(result_data), start=1):
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                STATUS_LABELS.get(status, status),
                market_impact_for_status(status),
                clean_text(
                    detail.get("regulation_reason")
                    or detail.get("regulation_notice")
                ),
                applicable_condition(status),
                clean_text(
                    detail.get("recommended_action"),
                    recommended_action(status),
                ),
                clean_text(detail.get("match_source")),
                source_file,
            ]
        )
    return rows


def halal_result_rows(
    result_data: dict,
    source_file: str,
) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for index, detail in enumerate(result_details(result_data), start=1):
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                halal_status_label(detail),
                clean_text(detail.get("match_source")),
                clean_text(
                    detail.get("verification_type"),
                    manual_type(detail, "HALAL"),
                ),
                clean_text(
                    detail.get("regulation_reason")
                    or detail.get("regulation_notice")
                ),
                clean_text(
                    detail.get("required_evidence"),
                    "공급사 원료규격서·기원서·제조공정서·인증자료",
                ),
                clean_text(
                    detail.get("recommended_action"),
                    "공급사 자료 확인 후 재검토",
                ),
                source_file,
            ]
        )
    return rows


def market_summary_rows(
    product_name: str,
    target: str,
    result_data: dict,
    halal_result: Optional[dict],
) -> List[List[Any]]:
    counts = status_counts(result_data)
    decision = market_decision(result_data)
    rows: List[List[Any]] = [
        ["제품명", product_name],
        ["대상 시장", MARKET_LABELS.get(target, target)],
        ["검사 성분 수", safe_int(result_data.get("total_checked"))],
        ["통과", counts["PASS"]],
        ["금지", counts["BANNED"]],
        ["제한", counts["RESTRICTED"]],
        ["경고·표시 필요", counts["WARNING_REQUIRED"]],
        ["규제 대상", counts["REGULATED"]],
        ["수동 검토", counts["REVIEW_REQUIRED"]],
        ["명칭 검증 필요", counts["VERIFICATION_REQUIRED"]],
        ["전체 수동확인 필요", counts["REVIEW_REQUIRED"] + counts["VERIFICATION_REQUIRED"] + counts["REGULATED"]],
        ["종합 판정", decision["overall"]],
        ["수입·판매 가능성", decision["possibility"]],
        ["최종결론", decision["conclusion"]],
        ["우선 조치", priority_actions(result_data, target)],
        ["규제 DB", clean_text(result_data.get("database_file"))],
        ["DB 업데이트일", clean_text(result_data.get("database_last_updated"))],
        ["보고서 생성일", clean_text(result_data.get("report_generated_at"), datetime.now().strftime("%Y-%m-%d %H:%M:%S"))],
        ["사용 범위", "화장품 성분 규제 사전 스크리닝"],
    ]
    if halal_result is not None:
        relation = halal_market_relation(target)
        hdecision = halal_decision(halal_result)
        rows.extend(
            [
                ["HALAL 추가검토 여부", "적용"],
                ["해당 시장과의 관련성", relation["level"]],
                ["국가 규제판정 변경 여부", "변경 없음"],
                ["HALAL 시장 영향", relation["impact"]],
                ["HALAL 종합판정", hdecision["overall"]],
                ["연결된 HALAL 확인 성분", linked_halal_ingredients(halal_result)],
                ["HALAL 추가 조치", "공급사 기원서·전체 구성표·제조공정서·유효한 인증자료를 확보한 뒤 재검토"],
            ]
        )
    return rows


def halal_summary_rows(
    product_name: str,
    result_data: dict,
) -> List[List[Any]]:
    counts = status_counts(result_data)
    decision = halal_decision(result_data)
    return [
        ["제품명", product_name],
        ["검토 구분", "할랄 추가 성분 스크리닝"],
        ["검사 성분 수", safe_int(result_data.get("total_checked"))],
        ["위험 후보 미확인", counts["PASS"]],
        ["할랄 금지", counts["BANNED"]],
        ["조건 확인 필요", counts["RESTRICTED"]],
        ["공식 기준 적용범위 확인", counts["REGULATED"]],
        ["공식 기준 수동검토", counts["REVIEW_REQUIRED"]],
        ["원료 기원·구성·인증 확인 필요", counts["VERIFICATION_REQUIRED"]],
        ["전체 확인 필요", sum(counts[key] for key in counts if key != "PASS")],
        ["할랄 종합판정", decision["overall"]],
        ["할랄 인증 추진 가능성", decision["possibility"]],
        ["최종결론", decision["conclusion"]],
        ["우선 확보 자료", "원료규격서·원료 기원서·전체 INCI·제조공정서·유효한 할랄 인증서"],
        ["확정 DB", clean_text(result_data.get("database_file"))],
        ["검토대기 DB", clean_text(result_data.get("review_database_file"))],
        ["DB 업데이트일", clean_text(result_data.get("database_last_updated"))],
        ["보고서 생성일", clean_text(result_data.get("report_generated_at"), datetime.now().strftime("%Y-%m-%d %H:%M:%S"))],
        ["사용 범위", "할랄 인증서 또는 최종 인증판정이 아닌 사전 위험 선별"],
    ]


def priority_actions(result_data: dict, target: str) -> str:
    actions: List[str] = []
    for detail in attention_details(result_data):
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        action = clean_text(
            detail.get("recommended_action"),
            recommended_action(status),
        )
        if action and action not in actions:
            actions.append(action)
    if not actions:
        return "시장별 등록·표시·통관 등 비성분 요건 확인"
    return " → ".join(actions[:4])


def manual_rows(
    result_data: dict,
    target: str,
    halal_result: Optional[dict],
) -> List[List[Any]]:
    rows: List[List[Any]] = []
    index = 1
    for detail in result_details(result_data):
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        if status not in MANUAL_STATUSES:
            continue
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                manual_type(detail, target),
                STATUS_LABELS.get(status, status),
                clean_text(detail.get("regulation_reason") or detail.get("regulation_notice")),
                applicable_condition(status),
                clean_text(detail.get("required_evidence"), recommended_action(status)),
                "",
                "",
                "",
                "",
            ]
        )
        index += 1

    if halal_result is not None:
        for detail in attention_details(halal_result):
            rows.append(
                [
                    index,
                    clean_text(detail.get("original_ingredient")),
                    clean_text(detail.get("inci_name")),
                    clean_text(detail.get("cas_number"), "N/A"),
                    f"HALAL - {manual_type(detail, 'HALAL')}",
                    halal_status_label(detail),
                    clean_text(detail.get("regulation_reason") or detail.get("regulation_notice")),
                    halal_market_relation(target)["impact"],
                    clean_text(detail.get("required_evidence"), "공급사 원료기원·구성·공정·인증자료"),
                    "",
                    "",
                    "",
                    "",
                ]
            )
            index += 1
    return rows


def halal_manual_rows(result_data: dict) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for index, detail in enumerate(attention_details(result_data), start=1):
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                manual_type(detail, "HALAL"),
                halal_status_label(detail),
                clean_text(detail.get("regulation_reason") or detail.get("regulation_notice")),
                clean_text(detail.get("recommended_action"), "원료 기원·구성·공정·인증 여부 확인"),
                clean_text(detail.get("required_evidence"), "원료규격서·기원서·공정서·인증서"),
                "",
                "",
                "",
                "",
            ]
        )
    return rows


# ============================================================
# Excel 생성
# ============================================================
def _xlsx_title(ws, title: str, last_col: int) -> None:
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=last_col)
    cell = ws.cell(1, 1, title)
    cell.fill = PatternFill("solid", fgColor=LIGHT_BLUE)
    cell.font = Font(bold=True, size=15, color=NAVY)
    cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 30


def _xlsx_header(ws, row: int, headers: Sequence[str]) -> None:
    thin = Side(style="thin", color=BORDER_COLOR)
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row, col, header)
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(bold=True, color=WHITE, size=10)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    ws.row_dimensions[row].height = 28


def _xlsx_body(ws, start_row: int, end_row: int, max_col: int) -> None:
    thin = Side(style="thin", color=BORDER_COLOR)
    for row in ws.iter_rows(min_row=start_row, max_row=end_row, min_col=1, max_col=max_col):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.font = Font(size=9, color="333333")
    for row_index in range(start_row, end_row + 1):
        ws.row_dimensions[row_index].height = 42


def _set_widths(ws, widths: Sequence[float]) -> None:
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width


def _write_result_sheet(ws, title: str, headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> None:
    _xlsx_title(ws, title, len(headers))
    _xlsx_header(ws, 3, headers)
    for r_idx, row in enumerate(rows, start=4):
        for c_idx, value in enumerate(row, start=1):
            ws.cell(r_idx, c_idx, value)
    end_row = max(4, 3 + len(rows))
    _xlsx_body(ws, 4, end_row, len(headers))
    ws.freeze_panes = "A4"
    ws.auto_filter.ref = f"A3:{get_column_letter(len(headers))}{end_row}"
    _set_widths(ws, [7, 18, 26, 14, 20, 22, 32, 26, 30, 24, 24])

    for row_index in range(4, end_row + 1):
        status = clean_text(ws.cell(row_index, 5).value)
        fill = None
        if "금지" in status:
            fill = RED
        elif any(term in status for term in ("제한", "확인", "검토", "규제")):
            fill = YELLOW
        elif "통과" in status or "미확인" in status:
            fill = GREEN
        if fill:
            ws.cell(row_index, 5).fill = PatternFill("solid", fgColor=fill)
            ws.cell(row_index, 5).font = Font(bold=True, size=9)


def _write_summary_sheet(ws, title: str, rows: Sequence[Sequence[Any]]) -> None:
    _xlsx_title(ws, title, 2)
    _xlsx_header(ws, 3, ["항목", "내용"])
    for r_idx, row in enumerate(rows, start=4):
        ws.cell(r_idx, 1, row[0])
        ws.cell(r_idx, 2, row[1])
        ws.cell(r_idx, 1).fill = PatternFill("solid", fgColor=GRAY)
        ws.cell(r_idx, 1).font = Font(bold=True, size=9)
    end_row = max(4, 3 + len(rows))
    _xlsx_body(ws, 4, end_row, 2)
    _set_widths(ws, [28, 80])
    for r_idx in range(4, end_row + 1):
        ws.row_dimensions[r_idx].height = 30


def _write_manual_sheet(ws, title: str, rows: Sequence[Sequence[Any]]) -> None:
    headers = [
        "번호",
        "성분명",
        "INCI 명칭",
        "CAS 번호",
        "수동확인 유형",
        "현재 판정",
        "확인 사유",
        "확인할 사항",
        "필요한 자료",
        "확인 결과",
        "최종 판정",
        "검토자",
        "검토일",
    ]
    _xlsx_title(ws, title, len(headers))
    _xlsx_header(ws, 3, headers)
    if rows:
        for r_idx, row in enumerate(rows, start=4):
            for c_idx, value in enumerate(row, start=1):
                ws.cell(r_idx, c_idx, value)
    else:
        rows = [[1, "해당 항목 없음", "", "", "", "", "", "", "", "", "", "", ""]]
        for c_idx, value in enumerate(rows[0], start=1):
            ws.cell(4, c_idx, value)
    end_row = 3 + len(rows)
    _xlsx_body(ws, 4, end_row, len(headers))
    ws.freeze_panes = "A4"
    ws.auto_filter.ref = f"A3:M{end_row}"
    _set_widths(ws, [7, 18, 25, 14, 24, 20, 32, 32, 28, 22, 20, 14, 14])

    status_validation = DataValidation(
        type="list",
        formula1='"확인 중,적합,조건부 적합,부적합,재검토 필요"',
        allow_blank=True,
    )
    ws.add_data_validation(status_validation)
    status_validation.add(f"K4:K{max(end_row, 100)}")


def create_product_excel_bytes(
    product_name: str,
    source_file: str,
    market_results_map: Dict[str, dict],
    selected_markets: Sequence[str],
    halal_result: Optional[dict] = None,
) -> bytes:
    workbook = Workbook()
    workbook.remove(workbook.active)
    sheet_number = 1

    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue
        label = MARKET_LABELS.get(target, target)

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_규제분석결과")
        _write_result_sheet(
            ws,
            f"{label} 규제 분석 결과",
            [
                "번호",
                "입력 성분명",
                "INCI 명칭",
                "CAS 번호",
                "규제 판정",
                "시장 진입 영향",
                "규제 근거",
                "적용 조건",
                "필요한 조치",
                "판정 출처",
                "입력 파일",
            ],
            market_result_rows(result_data, source_file),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_종합요약")
        _write_summary_sheet(
            ws,
            f"{label} 종합요약",
            market_summary_rows(
                product_name,
                target,
                result_data,
                halal_result,
            ),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_수동확인")
        _write_manual_sheet(
            ws,
            f"{label} 수동확인 체크리스트",
            manual_rows(result_data, target, halal_result),
        )
        sheet_number += 1

    if halal_result is not None:
        ws = workbook.create_sheet(f"{sheet_number:02d}_HALAL_성분검토결과")
        _write_result_sheet(
            ws,
            "HALAL 추가 성분 검토 결과",
            [
                "번호",
                "입력 성분명",
                "INCI 명칭",
                "CAS 번호",
                "할랄 검토 상태",
                "검출 방식",
                "우려 유형",
                "확인 사유",
                "필요한 증빙",
                "공급사 확인사항",
                "입력 파일",
            ],
            halal_result_rows(halal_result, source_file),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_HALAL_종합요약")
        _write_summary_sheet(
            ws,
            "HALAL 종합요약",
            halal_summary_rows(product_name, halal_result),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_HALAL_수동확인")
        _write_manual_sheet(
            ws,
            "HALAL 수동확인 체크리스트",
            halal_manual_rows(halal_result),
        )

    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


# ============================================================
# DOCX 생성 도우미
# ============================================================
def _doc_shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _doc_cell_margins(
    cell: Any,
    top: int = 70,
    start: int = 90,
    bottom: int = 70,
    end: int = 90,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _doc_set_cell(
    cell: Any,
    text: Any,
    bold: bool = False,
    size: float = 9,
    color: RGBColor = DOC_TEXT,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(clean_text(text))
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _doc_cell_margins(cell)


def _doc_borders(table: Any, color: str = BORDER_COLOR, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _doc_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B9C2C8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _chapter_header(
    document: Document,
    number: int,
    title: str,
    subtitle: str,
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(0.85)
    table.columns[1].width = Inches(5.8)
    _doc_borders(table, WHITE, 0)
    _doc_shade(table.cell(0, 0), "3C3C3C")
    _doc_set_cell(
        table.cell(0, 0),
        f"제{number}장",
        True,
        9.2,
        RGBColor(255, 255, 255),
        WD_ALIGN_PARAGRAPH.CENTER,
    )
    _doc_set_cell(table.cell(0, 1), title, True, 17.2)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(subtitle)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DOC_MUTED
    _doc_rule(document)


def _section_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11.2)
    run.font.color.rgb = DOC_NAVY


def _bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.08)
    paragraph.paragraph_format.space_after = Pt(1.5)
    run = paragraph.add_run(text)
    run.font.size = Pt(9.6)


def _decision_box(document: Document, label: str, text: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.15)
    _doc_borders(table, "D8BE63", 6)
    _doc_shade(table.cell(0, 0), YELLOW)
    _doc_shade(table.cell(0, 1), "FFF9E8")
    _doc_set_cell(
        table.cell(0, 0),
        label,
        True,
        9.4,
        DOC_TEXT,
        WD_ALIGN_PARAGRAPH.CENTER,
    )
    _doc_set_cell(
        table.cell(0, 1),
        text,
        True,
        10.2,
        DOC_GOLD,
        WD_ALIGN_PARAGRAPH.CENTER,
    )


def _add_simple_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    header_fill: str = NAVY,
    font_size: float = 8.7,
    highlight_col: Optional[int] = None,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _doc_borders(table)
    for index, header in enumerate(headers):
        _doc_shade(table.cell(0, index), header_fill)
        _doc_set_cell(
            table.cell(0, index),
            header,
            True,
            font_size,
            RGBColor(255, 255, 255),
            WD_ALIGN_PARAGRAPH.CENTER,
        )
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            if row_index % 2:
                _doc_shade(cells[col_index], PALE)
            if highlight_col is not None and col_index == highlight_col:
                text = clean_text(value)
                if any(term in text for term in ("금지", "제한", "확인", "검토", "규제")):
                    _doc_shade(cells[col_index], YELLOW)
            _doc_set_cell(
                cells[col_index],
                value,
                col_index == highlight_col,
                font_size,
                DOC_TEXT,
                WD_ALIGN_PARAGRAPH.CENTER if col_index == highlight_col else WD_ALIGN_PARAGRAPH.LEFT,
            )


def _market_attention_rows(result_data: dict) -> List[List[str]]:
    details = result_details(result_data)
    chosen = details if len(details) <= 18 else attention_details(result_data)
    rows: List[List[str]] = []
    for detail in chosen:
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        rows.append(
            [
                clean_text(detail.get("original_ingredient")),
                STATUS_LABELS.get(status, status),
                clean_text(detail.get("regulation_reason") or detail.get("regulation_notice")),
            ]
        )
    if len(details) > 18 and not rows:
        rows.append(["전체 성분", "통과", "상세 성분별 결과는 Excel 규제분석결과 시트를 참조하십시오."])
    return rows


def _manual_report_rows(result_data: dict, target: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for detail in result_details(result_data):
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        if status not in MANUAL_STATUSES:
            continue
        rows.append(
            [
                clean_text(detail.get("original_ingredient")),
                STATUS_LABELS.get(status, status),
                clean_text(detail.get("regulation_reason") or detail.get("regulation_notice")),
                clean_text(detail.get("recommended_action"), recommended_action(status)),
            ]
        )
    return rows


def _halal_link_rows(halal_result: dict, target: str) -> List[List[str]]:
    relation = halal_market_relation(target)
    rows: List[List[str]] = []
    for detail in attention_details(halal_result)[:12]:
        rows.append(
            [
                clean_text(detail.get("original_ingredient")),
                halal_status_label(detail),
                relation["impact"],
                clean_text(
                    detail.get("required_evidence"),
                    "공급사 원료 기원·구성·공정·인증자료",
                ),
            ]
        )
    return rows


def _common_glossary_rows() -> List[Tuple[str, str]]:
    return [
        ("통과", "현재 규제 DB에서 일치 항목이 확인되지 않은 상태"),
        ("금지", "현재 처방 기준 해당 시장 수입·유통·판매 진행이 불가능한 상태"),
        ("제한", "농도·제품유형·사용대상 등 조건 충족이 필요한 상태"),
        ("경고·표시 필요", "라벨·고지·경고문구·통지 요건 검토가 필요한 상태"),
        ("규제 대상", "관련 규제는 있으나 실제 적용범위를 추가 확인해야 하는 상태"),
        ("수동 검토", "공식 원문이나 공급사 자료를 사람이 확인해야 하는 상태"),
        ("명칭 검증 필요", "공식 INCI가 확정되지 않아 판정을 완료할 수 없는 상태"),
        ("확인 필요", "추가 자료 없이는 최종 결론을 내릴 수 없는 상태"),
        ("부적합", "현재 처방 또는 표시 상태로 규제요건을 충족하지 못하는 상태"),
        ("조건 충족 후 진행 가능", "필수 제한·표시·자료 요건을 충족하면 진행 가능한 상태"),
        ("CAS 번호", "Chemical Abstracts Service가 화학물질에 부여한 고유 식별번호"),
    ]


def _halal_glossary_rows() -> List[Tuple[str, str]]:
    return [
        ("할랄", "이슬람 기준에 따라 허용되는 원료·공정·제품"),
        ("하람", "이슬람 기준상 허용되지 않는 원료·공정·제품"),
        ("원료 기원", "식물성·합성·동물성 등 원료가 유래한 근원"),
        ("동물 유래", "동물 조직·지방·분비물 등에서 유래한 원료"),
        ("가공보조제", "최종 성분표에 남지 않더라도 제조 중 사용되는 물질"),
        ("교차오염", "비할랄 원료·설비·보관·운송 과정에서 혼입될 가능성"),
        ("할랄 인증서", "인증기관이 원료 또는 제품의 적합성을 확인해 발급한 문서"),
        ("위험 후보 미확인", "현재 분석에서 우려 후보가 발견되지 않았으나 인증 완료는 아닌 상태"),
        ("원료 기원 확인 필요", "공급사 자료로 원료 유래와 제조공정을 확인해야 하는 상태"),
        ("CAS 번호", "Chemical Abstracts Service가 화학물질에 부여한 고유 식별번호"),
    ]


def _glossary_table(document: Document, rows: Sequence[Tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _doc_borders(table)
    for term, definition in rows:
        cells = table.add_row().cells
        _doc_shade(cells[0], GRAY)
        _doc_set_cell(cells[0], term, True, 8.3)
        _doc_set_cell(cells[1], definition, False, 8.3)


def create_product_report_bytes(
    product_name: str,
    source_file: str,
    market_results_map: Dict[str, dict],
    selected_markets: Sequence[str],
    halal_result: Optional[dict] = None,
) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = DOC_TEXT
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(3)

    market_labels = [
        MARKET_LABELS.get(target, target)
        for target in selected_markets
        if target in market_results_map
    ]
    scope = " · ".join(market_labels)
    if halal_result is not None:
        scope += " + HALAL 추가검토"

    _chapter_header(
        document,
        1,
        "제품 및 분석 개요",
        f"{product_name} · {scope}",
    )

    info = document.add_table(rows=5, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.LEFT
    _doc_borders(info)
    overview_rows = [
        ("제품명", product_name),
        ("입력 파일", source_file),
        ("국가 규제시장", " · ".join(market_labels) or "없음"),
        ("추가 검토", "HALAL 성분 기원·구성·인증자료 검토" if halal_result is not None else "미선택"),
        ("검사 성분 수", max([safe_int(data.get("total_checked")) for data in market_results_map.values()] + ([safe_int(halal_result.get("total_checked"))] if halal_result else [0]))),
    ]
    for row_index, (label, value) in enumerate(overview_rows):
        _doc_shade(info.cell(row_index, 0), GRAY)
        _doc_set_cell(info.cell(row_index, 0), label, True, 9.5)
        _doc_set_cell(info.cell(row_index, 1), value, False, 9.5)

    first_result = next(iter(market_results_map.values()), halal_result or {})
    ingredients = [
        [
            clean_text(item.get("original_ingredient")),
            clean_text(item.get("inci_name")),
            clean_text(item.get("cas_number"), "N/A"),
        ]
        for item in result_details(first_result)
    ]
    _section_title(document, "1. 제품 성분 구성")
    if len(ingredients) <= 24:
        _add_simple_table(
            document,
            ["입력 성분명", "INCI 명칭", "CAS 번호"],
            ingredients,
            NAVY,
            8.8,
        )
    else:
        paragraph = document.add_paragraph()
        paragraph.add_run(
            f"전체 {len(ingredients)}개 성분을 분석했습니다. 전체 성분 목록과 성분별 판정은 함께 제공되는 Excel 파일을 참조하십시오."
        )

    _section_title(document, "2. 시장별 종합판정")
    summary_rows: List[List[str]] = []
    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue
        decision = market_decision(result_data)
        summary_rows.append(
            [
                "국가 규제",
                target,
                decision["overall"],
                decision["possibility"],
            ]
        )
    if halal_result is not None:
        decision = halal_decision(halal_result)
        summary_rows.append(
            [
                "추가 검토",
                "HALAL",
                decision["overall"],
                decision["possibility"],
            ]
        )
    _add_simple_table(
        document,
        ["구분", "대상", "종합판정", "가능성 판단"],
        summary_rows,
        BLUE,
        8.8,
        2,
    )

    _section_title(document, "3. 통합 최종결론")
    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue
        decision = market_decision(result_data)
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{MARKET_LABELS.get(target, target)}: ").bold = True
        paragraph.add_run(decision["conclusion"])
    if halal_result is not None:
        decision = halal_decision(halal_result)
        paragraph = document.add_paragraph()
        paragraph.add_run("HALAL: ").bold = True
        paragraph.add_run(decision["conclusion"])
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            "HALAL 결과는 독립 장에서 상세히 설명하며, 각 시장 장에서는 해당 결과가 현지 인증·바이어·유통채널 요구에 미치는 영향만 연결하여 표시합니다."
        )
        run.font.size = Pt(9.2)
        run.font.color.rgb = DOC_MUTED

    _section_title(document, "4. 면책 및 사용 범위")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        "본 보고서는 국가별 화장품 성분 규제 및 할랄 위험 후보를 사전에 선별하는 문서입니다. 제품 등록, 통관 승인, 판매 허가, 할랄 인증서 또는 최종 법률 판단을 대신하지 않습니다. 실제 진행 전 최신 공식 규정, 처방 농도, 제품 유형, 라벨, 공급사 자료 및 인증기관 요구사항을 별도로 확인해야 합니다."
    )
    run.font.size = Pt(9.2)
    run.font.color.rgb = DOC_MUTED

    chapter_number = 2
    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue

        document.add_page_break()
        market_label = MARKET_LABELS.get(target, target)
        _chapter_header(
            document,
            chapter_number,
            f"{market_label} 규제 분석 결과",
            "시장별 분석은 이전 시장과 분리하여 새 장에서 시작합니다.",
        )
        chapter_number += 1
        decision = market_decision(result_data)
        _decision_box(document, "시장별 최종결론", decision["possibility"])

        _section_title(document, "1. 분석 세부내용")
        rows = _market_attention_rows(result_data)
        if rows:
            _add_simple_table(
                document,
                ["성분", "판정", "결과 해석"],
                rows,
                NAVY,
                8.6,
                1,
            )
        else:
            _bullet(document, "추가 확인이 필요한 성분이 없습니다. 전체 결과는 Excel을 참조하십시오.")

        _section_title(document, "2. 세부 확인사항")
        attention = attention_details(result_data)
        if attention:
            seen_actions: List[str] = []
            for detail in attention:
                status = normalize_status(
                    detail.get("compliance_status")
                    or detail.get("restriction_type")
                )
                action = clean_text(
                    detail.get("recommended_action"),
                    recommended_action(status),
                )
                if action and action not in seen_actions:
                    seen_actions.append(action)
            for action in seen_actions[:8]:
                _bullet(document, action)
        else:
            _bullet(document, "성분 규제 단계 외의 제품 등록·표시·통관 요건을 별도로 확인하십시오.")

        section_index = 3
        if halal_result is not None:
            _section_title(
                document,
                f"{section_index}. HALAL 추가검토의 해당 시장 영향",
            )
            section_index += 1
            relation = halal_market_relation(target)
            relation_table = [
                ["해당 시장과의 관련성", relation["level"]],
                ["국가 규제판정 변경 여부", "변경 없음"],
                ["시장 영향", relation["impact"]],
                ["HALAL 종합판정", halal_decision(halal_result)["overall"]],
            ]
            _add_simple_table(
                document,
                ["구분", "내용"],
                relation_table,
                BLUE,
                8.7,
            )
            link_rows = _halal_link_rows(halal_result, target)
            if link_rows:
                _add_simple_table(
                    document,
                    ["HALAL 확인 대상", "검토 상태", "해당 시장 영향", "필요한 자료"],
                    link_rows,
                    NAVY,
                    8.2,
                    1,
                )
            else:
                _bullet(document, "HALAL 추가검토에서 별도의 확인 대상이 발견되지 않았습니다.")

        _section_title(document, f"{section_index}. 수동확인 항목")
        section_index += 1
        manual = _manual_report_rows(result_data, target)
        if manual:
            _add_simple_table(
                document,
                ["성분", "현재 판정", "확인 사유", "필요한 조치"],
                manual,
                BLUE,
                8.4,
                1,
            )
        else:
            _bullet(document, "국가 규제 수동확인 항목이 없습니다.")

        _section_title(document, f"{section_index}. 판정 해석 시 주의사항")
        section_index += 1
        _bullet(document, "통과는 현재 로드된 데이터에서 일치 항목이 없다는 의미이며 최종 수입·판매 허가를 보증하지 않습니다.")
        _bullet(document, "제한·규제 대상은 실제 농도, 제품유형, 용도와 표시정보를 함께 확인해야 합니다.")
        _bullet(document, "수동검토와 명칭검증 항목은 확인 전 최종 적합 판정으로 취급하면 안 됩니다.")
        if halal_result is not None:
            _bullet(document, "HALAL 결과는 국가 화장품 규제판정을 변경하지 않으며 인증·바이어·유통요건에 대한 별도 영향으로 해석해야 합니다.")

        _section_title(document, f"{section_index}. 후속 절차")
        section_index += 1
        _bullet(document, "공급사 자료 및 공식 INCI 확보")
        _bullet(document, "배합농도·제품유형·표시요건 확인")
        if halal_result is not None:
            _bullet(document, "HALAL 관련 원료 기원·구성·공정·인증자료 확보")
        _bullet(document, "확인자료 반영 후 동일 시장 재분석")

        _section_title(document, f"{section_index}. 용어의 정의")
        _glossary_table(document, _common_glossary_rows())

    if halal_result is not None:
        document.add_page_break()
        _chapter_header(
            document,
            chapter_number,
            "HALAL 추가 성분 검토",
            "국가 규제판정과 별도로 원료 기원·구성·제조공정·인증자료를 검토합니다.",
        )
        decision = halal_decision(halal_result)
        _decision_box(document, "할랄 최종결론", decision["possibility"])

        _section_title(document, "1. 분석 세부내용")
        rows: List[List[str]] = []
        for detail in result_details(halal_result):
            rows.append(
                [
                    clean_text(detail.get("original_ingredient")),
                    halal_status_label(detail),
                    clean_text(detail.get("verification_type"), manual_type(detail, "HALAL")),
                    clean_text(detail.get("required_evidence"), "공급사 원료규격서·기원서·공정서·인증자료"),
                ]
            )
        if len(rows) > 22:
            attention_names = {
                clean_text(item.get("original_ingredient"))
                for item in attention_details(halal_result)
            }
            rows = [row for row in rows if row[0] in attention_names]
        _add_simple_table(
            document,
            ["성분", "할랄 검토 상태", "우려 유형", "필요한 증빙"],
            rows,
            NAVY,
            8.4,
            1,
        )

        _section_title(document, "2. 세부 확인사항")
        attention = attention_details(halal_result)
        if attention:
            for detail in attention[:10]:
                _bullet(
                    document,
                    f"{clean_text(detail.get('original_ingredient'))}: {clean_text(detail.get('recommended_action'), '원료 기원·구성·제조공정 및 인증자료 확인')}",
                )
        else:
            _bullet(document, "현재 분석 데이터에서 별도의 할랄 위험 후보가 확인되지 않았습니다.")

        _section_title(document, "3. 수동확인 항목")
        manual = [
            [row[1], row[4], row[7], row[8]]
            for row in halal_manual_rows(halal_result)
        ]
        if manual:
            _add_simple_table(
                document,
                ["성분", "수동확인 유형", "확인할 사항", "필요한 자료"],
                manual,
                BLUE,
                8.4,
                1,
            )
        else:
            _bullet(document, "HALAL 수동확인 항목이 없습니다.")

        _section_title(document, "4. 판정 해석 시 주의사항")
        _bullet(document, "할랄 위험 후보 미확인은 할랄 인증 완료를 의미하지 않습니다.")
        _bullet(document, "성분명만으로 동물 종, 제조공정, 발효 배지, 가공보조제와 교차오염을 확정할 수 없습니다.")
        _bullet(document, "최종 인증 가능 여부는 공급사 증빙과 적용 인증기관의 심사를 통해 확인해야 합니다.")

        _section_title(document, "5. 후속 절차")
        _bullet(document, "공급사 원료규격서·기원서·전체 INCI 확보")
        _bullet(document, "제조공정·가공보조제·교차오염 여부 확인")
        _bullet(document, "할랄 인증서 발급기관·번호·유효기간 확인")
        _bullet(document, "자료 반영 후 HALAL 재분석 및 인증기관 검토")

        _section_title(document, "6. 용어의 정의")
        _glossary_table(document, _halal_glossary_rows())

    for document_section in document.sections:
        footer = document_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(
            "K-Beauty Global Compliance | 다중시장 규제 및 HALAL 추가검토"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = DOC_MUTED

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


# ============================================================
# ZIP 일괄 다운로드
# ============================================================
def create_bundle_zip_bytes(
    product_outputs: Sequence[dict],
    failures: Sequence[dict],
) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for product in product_outputs:
            product_name = safe_filename(clean_text(product.get("product_name"), "제품"))
            excel_data = product.get("excel_data")
            report_data = product.get("report_data")
            if excel_data:
                archive.writestr(
                    f"{product_name}_다중시장_규제분석결과.xlsx",
                    excel_data,
                )
            if report_data:
                archive.writestr(
                    f"{product_name}_다중시장_규제스크리닝_보고서.docx",
                    report_data,
                )

        if failures:
            text_stream = io.StringIO()
            writer = csv.DictWriter(
                text_stream,
                fieldnames=["제품명", "시장", "오류"],
            )
            writer.writeheader()
            for failure in failures:
                writer.writerow(
                    {
                        "제품명": clean_text(failure.get("product_name")),
                        "시장": clean_text(failure.get("target")),
                        "오류": clean_text(failure.get("error")),
                    }
                )
            archive.writestr(
                "분석실패목록.csv",
                "\ufeff" + text_stream.getvalue(),
            )

    return output.getvalue()
