from __future__ import annotations

import csv
import io
import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Sequence, Set, Tuple

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation


REPORT_BUILD_ID = "v9.1.1-report-r2"


MARKET_LABELS: Dict[str, str] = {
    "US": "미국(US)",
    "EU": "유럽연합(EU)",
    "UK": "영국(UK)",
    "CN": "중국(CN)",
    "ASEAN": "아세안(ASEAN)",
    "SFDA": "사우디아라비아(SFDA)",
    "EAC": "유라시아경제연합(EAEU)",
    "HALAL": "HALAL 추가 성분 검토",
}


OUTPUT_LANGUAGES: Dict[str, str] = {
    "ko": "한국어",
    "en": "English",
    "zh-CN": "简体中文",
    "ar": "العربية",
    "fr": "Français",
    "id": "Bahasa Indonesia",
    "ms": "Bahasa Melayu",
    "sw": "Kiswahili",
}

LANGUAGE_FILE_SUFFIXES: Dict[str, str] = {
    "ko": "KO",
    "en": "EN",
    "zh-CN": "ZH-CN",
    "ar": "AR",
    "fr": "FR",
    "id": "ID",
    "ms": "MS",
    "sw": "SW",
}

CHAPTER_LABELS: Dict[str, str] = {
    "ko": "제{number}장",
    "en": "Chapter {number}",
    "zh-CN": "第{number}章",
    "ar": "الفصل {number}",
    "fr": "Chapitre {number}",
    "id": "Bab {number}",
    "ms": "Bab {number}",
    "sw": "Sura ya {number}",
}

CHAPTER_BOX_WIDTHS: Dict[str, float] = {
    "ko": 0.85,
    "en": 1.18,
    "zh-CN": 0.85,
    "ar": 1.12,
    "fr": 1.18,
    "id": 0.88,
    "ms": 0.88,
    "sw": 1.05,
}

SHEET_SUFFIXES: Dict[str, Dict[str, str]] = {
    "ko": {"result": "규제분석결과", "summary": "종합요약", "manual": "수동확인", "halal_result": "HALAL_성분검토", "halal_summary": "HALAL_요약", "halal_manual": "HALAL_수동확인"},
    "en": {"result": "Screening_Result", "summary": "Summary", "manual": "Manual_Review", "halal_result": "HALAL_Result", "halal_summary": "HALAL_Summary", "halal_manual": "HALAL_Review"},
    "zh-CN": {"result": "结果", "summary": "汇总", "manual": "人工复核", "halal_result": "HALAL_结果", "halal_summary": "HALAL_汇总", "halal_manual": "HALAL_复核"},
    "ar": {"result": "نتيجة", "summary": "ملخص", "manual": "مراجعة", "halal_result": "HALAL_نتيجة", "halal_summary": "HALAL_ملخص", "halal_manual": "HALAL_مراجعة"},
    "fr": {"result": "Résultat", "summary": "Résumé", "manual": "Vérification", "halal_result": "HALAL_Résultat", "halal_summary": "HALAL_Résumé", "halal_manual": "HALAL_Vérif"},
    "id": {"result": "Hasil", "summary": "Ringkasan", "manual": "Verifikasi", "halal_result": "HALAL_Hasil", "halal_summary": "HALAL_Ringkas", "halal_manual": "HALAL_Verif"},
    "ms": {"result": "Keputusan", "summary": "Ringkasan", "manual": "Semakan", "halal_result": "HALAL_Keputusan", "halal_summary": "HALAL_Ringkas", "halal_manual": "HALAL_Semakan"},
    "sw": {"result": "Matokeo", "summary": "Muhtasari", "manual": "Ukaguzi", "halal_result": "HALAL_Matokeo", "halal_summary": "HALAL_Muhtasari", "halal_manual": "HALAL_Ukaguzi"},
}

Translator = Callable[[Sequence[str], str, Sequence[str]], Dict[str, str]]

STATUS_LABELS: Dict[str, str] = {
    "PASS": "통과",
    "BANNED": "금지",
    "RESTRICTED": "제한",
    "WARNING_REQUIRED": "경고·표시 필요",
    "REGULATED": "규제 대상",
    "REVIEW_REQUIRED": "수동 검토",
    "VERIFICATION_REQUIRED": "명칭 검증 필요",
    "FAIL": "부적합",
}


# Controlled English terminology for the default report.
# These terms are deterministic and are not delegated to a general translator.
ENGLISH_CANONICAL_EXACT: Dict[str, str] = {
    # Core status labels
    "통과": "Passed",
    "금지": "Prohibited",
    "제한": "Restricted",
    "경고·표시 필요": "Warning/Labeling Required",
    "규제 대상": "Regulatory Scope Verification Required",
    "수동 검토": "Manual Regulatory Review Required",
    "명칭 검증 필요": "Ingredient Identity Verification Required",
    "부적합": "Cannot Proceed with Current Formula",
    "확인 필요": "Verification Required",
    "조건부 가능": "Proceed Subject to Conditions",
    "적용 범위 확인 필요": "Regulatory Scope Verification Required",
    "증빙자료 확인 필요": "Supporting Documentation Required",
    "위험 후보 미확인": "No HALAL Risk Indicators Identified",
    "할랄 위험 후보 미확인": "No HALAL Risk Indicators Identified",
    "할랄 금지": "HALAL-Prohibited",
    "할랄 금지 성분 확인": "HALAL-Prohibited Ingredient Identified",
    "조건 확인 필요": "Condition Verification Required",
    "원료 기원 확인 필요": "Ingredient Origin Verification Required",
    "명칭·구성 확인 필요": "Ingredient Identity and Composition Verification Required",
    "표시·인증 확인 필요": "Labeling and Certification Verification Required",
    "공식 기준 수동검토": "Manual Review of Official Standards Required",
    "공식 기준 적용범위 확인": "Official-Standard Scope Verification Required",
    "원료 기원·구성·인증 확인 필요": "Origin, Composition, and Certification Verification Required",

    # Report structure
    "제품 및 분석 개요": "Product and Analysis Overview",
    "제품명": "Product Name",
    "입력 파일": "Input File",
    "국가 규제시장": "Regulatory Markets",
    "추가 검토": "Additional Review",
    "검사 성분 수": "Number of Ingredients Screened",
    "제품 성분 구성": "Product Ingredient Composition",
    "시장별 종합판정": "Overall Screening Decision by Market",
    "통합 최종결론": "Integrated Screening Conclusion",
    "면책 및 사용 범위": "Disclaimer and Scope of Use",
    "규제 분석 결과": "Regulatory Screening Results",
    "시장별 최종결론": "Market-Specific Screening Conclusion",
    "분석 세부내용": "Detailed Screening Results",
    "분석결과에 따른 조치": "Actions Based on Screening Results",
    "분석결과에 따른 필수 조치": "Required Actions Based on Screening Results",
    "조치 완료 기준": "Completion Criteria",
    "조치 완료 후 다음 단계": "Next Steps After Completing Required Actions",
    "다음 진행 단계": "Next Steps",
    "일반적인 시장진입 다음 단계": "General Market-Entry Preparation Steps",
    "수동확인 항목": "Manual Regulatory Review Items",
    "수동확인 유형": "Manual Review Type",
    "판정 해석 시 주의사항": "Important Interpretation Notes",
    "용어의 정의": "Definitions",
    "결과 해석": "Result Interpretation",
    "현재 판정": "Current Screening Decision",
    "최종 판정": "Final Review Decision",
    "확인할 사항": "Items to Verify",
    "확인 결과": "Verification Result",
    "검토자": "Reviewer",
    "검토일": "Review Date",
    "필요한 자료": "Required Documentation",
    "필요한 증빙": "Required Evidence",
    "적용 조건": "Applicable Conditions",
    "규제 근거": "Regulatory Basis",
    "판정 출처": "Decision Source",
    "대상 시장": "Target Market",
    "시장 영향": "Market Impact",
    "시장 진입 영향": "Market-Entry Impact",
    "국가 규제판정 변경 여부": "Change to National Regulatory Screening Decision",
    "해당 시장과의 관련성": "Relevance to the Market",
    "HALAL 추가 성분 검토": "Additional HALAL Ingredient Review",
    "HALAL 추가 성분 검토 결과": "Additional HALAL Ingredient Review Results",
    "HALAL 수동확인 체크리스트": "HALAL Manual Verification Checklist",
    "할랄 검토 상태": "HALAL Review Status",
    "할랄 종합판정": "Overall HALAL Review Decision",
    "할랄 최종결론": "HALAL Review Conclusion",
    "우려 유형": "Concern Type",
    "원료 기원": "Ingredient Origin",
    "가공보조제": "Processing Aid",
    "교차오염": "Cross-Contamination",
    "할랄 인증서": "HALAL Certificate",

    # Decisions and action labels
    "성분 규제 스크리닝 통과": "Ingredient Regulatory Screening Passed",
    "현재 처방 유지 가능": "Current Formula May Be Maintained",
    "현재 처방 진행 불가": "Current Formula Cannot Proceed",
    "현재 처방으로 진행 불가": "Current Formula Cannot Proceed",
    "확인 완료 전 판단 보류": "Decision Pending Verification",
    "명칭·자료 확인 전 판단 보류": "Decision Pending Ingredient Identity and Documentation Verification",
    "적용 범위 확인 후 결정": "Proceed Only After Regulatory Scope Verification",
    "제한조건 충족 후 진행 가능": "Proceed After All Restriction Conditions Are Met",
    "표시·고지요건 반영 후 진행 가능": "Proceed After Labeling and Notification Requirements Are Implemented",
    "인증 준비 가능": "May Proceed with Certification Preparation",
    "조건 확인 후 인증 준비 가능": "May Proceed with Certification Preparation After Condition Verification",
    "현재 처방 기준 할랄 인증 추진 곤란": "HALAL Certification Cannot Proceed with the Current Formula",
    "원료 기원·구성·공정·인증자료 확인 전 판단 보류": "Decision Pending Verification of Ingredient Origin, Composition, Manufacturing Process, and Certification Documentation",
    "성분 규제상 필수 보완조치 없음": "No Ingredient-Regulatory Corrective Action Required",
    "해당 없음(현재 처방 유지)": "Not Applicable (Maintain Current Formula)",
    "현재 성분명 기준 별도 HALAL 보완조치 없음": "No Additional HALAL Corrective Action Required Based on the Current Ingredient Identity",
    "해당 없음(현재 성분명 기준 별도 HALAL 보완조치 없음)": "Not Applicable (No Additional HALAL Corrective Action Required Based on the Current Ingredient Identity)",

    # Common data-source and evidence labels
    "현재 DB 일치 없음": "No Match in the Currently Loaded Database",
    "확정 규제 DB": "Confirmed Regulatory Database",
    "검토대기 DB": "Pending-Review Database",
    "명칭 검증": "Ingredient Identity Verification",
    "HALAL 사전검토": "HALAL Pre-Screening Review",
    "보조 키워드 검토": "Supplementary Keyword Review",
    "해당 항목 없음": "No Applicable Items",
    "없음": "None",
    "미선택": "Not Selected",
    "미확인": "Not Verified",
    "확인 중": "Under Review",
}

ENGLISH_CANONICAL_GLOSSARY: Tuple[Tuple[str, str], ...] = (
    ("통과", "Passed"),
    ("금지", "Prohibited"),
    ("제한", "Restricted"),
    ("경고·표시 필요", "Warning/Labeling Required"),
    ("규제 대상", "Regulatory Scope Verification Required"),
    ("수동 검토", "Manual Regulatory Review Required"),
    ("명칭 검증 필요", "Ingredient Identity Verification Required"),
    ("성분 규제 스크리닝 통과", "Ingredient Regulatory Screening Passed"),
    ("분석결과에 따른 필수 조치", "Required Actions Based on Screening Results"),
    ("필요한 자료", "Required Documentation"),
    ("조치 완료 기준", "Completion Criteria"),
    ("HALAL 위험 후보 미확인", "No HALAL Risk Indicators Identified"),
)

ENGLISH_DISALLOWED_PATTERNS: Tuple[Tuple[str, str], ...] = (
    (r"\bBanned\b", "Use 'Prohibited'."),
    (r"\bName Verification Required\b", "Use 'Ingredient Identity Verification Required'."),
    (r"\bManual Review Required\b", "Use 'Manual Regulatory Review Required'."),
    (r"\bWarning or Labeling Required\b", "Use 'Warning/Labeling Required'."),
    (r"\bFully Compliant\b", "A screening result must not claim full compliance."),
    (r"\bCompliance Confirmed\b", "A screening result must not claim confirmed compliance."),
    (r"\bMarket Entry Approved\b", "A screening result must not claim market authorization."),
    (r"\bCurrent Formula May Be Maintained\b", "Use a direct statement about whether an ingredient change is required."),
    (r"\bChange to National Regulatory Screening Decision\b", "Use a plain question-and-answer statement instead."),
    (r"\bRelevant for Separate Certification Purposes or Under Specific Conditions\b", "Use a direct statement of when HALAL documentation is required."),
)

ENGLISH_TRANSLATION_GLOSSARY_PROMPT = """
Use the following controlled terminology exactly whenever the corresponding Korean
concept appears:
- 통과 = Passed
- 금지 = Prohibited
- 제한 = Restricted
- 경고·표시 필요 = Warning/Labeling Required
- 규제 대상 = Regulatory Scope Verification Required
- 수동 검토 = Manual Regulatory Review Required
- 명칭 검증 필요 = Ingredient Identity Verification Required
- 성분 규제 스크리닝 통과 = Ingredient Regulatory Screening Passed
- 분석결과에 따른 필수 조치 = Required Actions Based on Screening Results
- 필요한 자료 = Required Documentation
- 조치 완료 기준 = Completion Criteria
- 할랄 위험 후보 미확인 = No HALAL Risk Indicators Identified

A 'Passed' screening result means only that no matching item was found in the
currently loaded screening data. Never translate it as 'fully compliant',
'compliance confirmed', 'approved', or 'market entry approved'.
Use 'regulatory screening' rather than 'regulatory approval' or 'legal clearance'.
Write HALAL in uppercase except when preserving an official quoted title.

""".strip()


# Controlled terminology for every optional output language. These labels are
# deterministic and bypass general-purpose translation so the regulatory
# meaning stays aligned with the approved English terminology.
ADDITIONAL_LANGUAGE_CANONICAL_EXACT: Dict[str, Dict[str, str]] = {
    "zh-CN": {
        "통과": "筛查通过",
        "금지": "禁止",
        "제한": "受限",
        "경고·표시 필요": "需要警示/标签标示",
        "규제 대상": "需要核实法规适用范围",
        "수동 검토": "需要人工法规审查",
        "명칭 검증 필요": "需要核实成分身份",
        "부적합": "当前配方无法继续推进",
        "확인 필요": "需要核验",
        "조건부 가능": "满足条件后可推进",
        "성분 규제 스크리닝 통과": "成分法规筛查通过",
        "분석결과에 따른 필수 조치": "根据筛查结果需采取的措施",
        "필요한 자료": "所需文件",
        "조치 완료 기준": "完成标准",
        "현재 처방 유지 가능": "可维持当前配方",
        "현재 처방 진행 불가": "当前配方无法继续推进",
        "현재 처방으로 진행 불가": "当前配方无法继续推进",
        "확인 완료 전 판단 보류": "核验完成前暂缓判断",
        "제한조건 충족 후 진행 가능": "满足全部限制条件后可推进",
        "표시·고지요건 반영 후 진행 가능": "落实标签和告知要求后可推进",
        "증빙자료 확인 필요": "需要核验支持性文件",
        "할랄 위험 후보 미확인": "未发现HALAL风险指标",
        "위험 후보 미확인": "未发现HALAL风险指标",
        "원료 기원 확인 필요": "需要核实原料来源",
    },
    "ar": {
        "통과": "اجتاز الفحص",
        "금지": "محظور",
        "제한": "مقيّد",
        "경고·표시 필요": "يلزم تحذير/بيانات على الملصق",
        "규제 대상": "يلزم التحقق من نطاق انطباق اللوائح",
        "수동 검토": "تلزم مراجعة تنظيمية يدوية",
        "명칭 검증 필요": "يلزم التحقق من هوية المكوّن",
        "부적합": "لا يمكن المتابعة بالتركيبة الحالية",
        "확인 필요": "يلزم التحقق",
        "조건부 가능": "يمكن المتابعة وفقًا للشروط",
        "성분 규제 스크리닝 통과": "اجتاز فحص لوائح المكونات",
        "분석결과에 따른 필수 조치": "الإجراءات المطلوبة بناءً على نتائج الفحص",
        "필요한 자료": "المستندات المطلوبة",
        "조치 완료 기준": "معايير الإكمال",
        "현재 처방 유지 가능": "يمكن الإبقاء على التركيبة الحالية",
        "현재 처방 진행 불가": "لا يمكن المتابعة بالتركيبة الحالية",
        "현재 처방으로 진행 불가": "لا يمكن المتابعة بالتركيبة الحالية",
        "확인 완료 전 판단 보류": "القرار معلّق إلى حين اكتمال التحقق",
        "제한조건 충족 후 진행 가능": "يمكن المتابعة بعد استيفاء جميع شروط التقييد",
        "표시·고지요건 반영 후 진행 가능": "يمكن المتابعة بعد تنفيذ متطلبات الملصق والإخطار",
        "증빙자료 확인 필요": "المستندات الداعمة مطلوبة",
        "할랄 위험 후보 미확인": "لم تُحدَّد مؤشرات مخاطر HALAL",
        "위험 후보 미확인": "لم تُحدَّد مؤشرات مخاطر HALAL",
        "원료 기원 확인 필요": "يلزم التحقق من منشأ المكوّن",
    },
    "fr": {
        "통과": "Filtrage réussi",
        "금지": "Interdit",
        "제한": "Restreint",
        "경고·표시 필요": "Avertissement/étiquetage requis",
        "규제 대상": "Vérification du champ d’application réglementaire requise",
        "수동 검토": "Examen réglementaire manuel requis",
        "명칭 검증 필요": "Vérification de l’identité de l’ingrédient requise",
        "부적합": "Impossible de poursuivre avec la formule actuelle",
        "확인 필요": "Vérification requise",
        "조건부 가능": "Poursuite possible sous réserve de conditions",
        "성분 규제 스크리닝 통과": "Filtrage réglementaire des ingrédients réussi",
        "분석결과에 따른 필수 조치": "Mesures requises d’après les résultats du filtrage",
        "필요한 자료": "Documents requis",
        "조치 완료 기준": "Critères d’achèvement",
        "현재 처방 유지 가능": "La formule actuelle peut être maintenue",
        "현재 처방 진행 불가": "Impossible de poursuivre avec la formule actuelle",
        "현재 처방으로 진행 불가": "Impossible de poursuivre avec la formule actuelle",
        "확인 완료 전 판단 보류": "Décision suspendue jusqu’à la fin de la vérification",
        "제한조건 충족 후 진행 가능": "Poursuite possible après satisfaction de toutes les conditions de restriction",
        "표시·고지요건 반영 후 진행 가능": "Poursuite possible après mise en œuvre des exigences d’étiquetage et de notification",
        "증빙자료 확인 필요": "Documents justificatifs requis",
        "할랄 위험 후보 미확인": "Aucun indicateur de risque HALAL identifié",
        "위험 후보 미확인": "Aucun indicateur de risque HALAL identifié",
        "원료 기원 확인 필요": "Vérification de l’origine de l’ingrédient requise",
    },
    "id": {
        "통과": "Lulus penyaringan",
        "금지": "Dilarang",
        "제한": "Dibatasi",
        "경고·표시 필요": "Peringatan/Pelabelan Wajib",
        "규제 대상": "Verifikasi Cakupan Regulasi Diperlukan",
        "수동 검토": "Tinjauan Regulasi Manual Diperlukan",
        "명칭 검증 필요": "Verifikasi Identitas Bahan Diperlukan",
        "부적합": "Tidak Dapat Dilanjutkan dengan Formula Saat Ini",
        "확인 필요": "Verifikasi Diperlukan",
        "조건부 가능": "Dapat Dilanjutkan dengan Syarat",
        "성분 규제 스크리닝 통과": "Lulus Penyaringan Regulasi Bahan",
        "분석결과에 따른 필수 조치": "Tindakan Wajib Berdasarkan Hasil Penyaringan",
        "필요한 자료": "Dokumen yang Diperlukan",
        "조치 완료 기준": "Kriteria Penyelesaian",
        "현재 처방 유지 가능": "Formula Saat Ini Dapat Dipertahankan",
        "현재 처방 진행 불가": "Tidak Dapat Dilanjutkan dengan Formula Saat Ini",
        "현재 처방으로 진행 불가": "Tidak Dapat Dilanjutkan dengan Formula Saat Ini",
        "확인 완료 전 판단 보류": "Keputusan Ditunda hingga Verifikasi Selesai",
        "제한조건 충족 후 진행 가능": "Dapat Dilanjutkan setelah Semua Ketentuan Pembatasan Dipenuhi",
        "표시·고지요건 반영 후 진행 가능": "Dapat Dilanjutkan setelah Persyaratan Pelabelan dan Pemberitahuan Diterapkan",
        "증빙자료 확인 필요": "Dokumen Pendukung Diperlukan",
        "할랄 위험 후보 미확인": "Tidak Ada Indikator Risiko HALAL yang Teridentifikasi",
        "위험 후보 미확인": "Tidak Ada Indikator Risiko HALAL yang Teridentifikasi",
        "원료 기원 확인 필요": "Verifikasi Asal Bahan Diperlukan",
    },
    "ms": {
        "통과": "Lulus saringan",
        "금지": "Dilarang",
        "제한": "Terhad",
        "경고·표시 필요": "Amaran/Pelabelan Diperlukan",
        "규제 대상": "Pengesahan Skop Peraturan Diperlukan",
        "수동 검토": "Semakan Peraturan Manual Diperlukan",
        "명칭 검증 필요": "Pengesahan Identiti Ramuan Diperlukan",
        "부적합": "Tidak Boleh Diteruskan dengan Formula Semasa",
        "확인 필요": "Pengesahan Diperlukan",
        "조건부 가능": "Boleh Diteruskan Tertakluk pada Syarat",
        "성분 규제 스크리닝 통과": "Lulus Saringan Peraturan Ramuan",
        "분석결과에 따른 필수 조치": "Tindakan Diperlukan Berdasarkan Hasil Saringan",
        "필요한 자료": "Dokumen Diperlukan",
        "조치 완료 기준": "Kriteria Penyelesaian",
        "현재 처방 유지 가능": "Formula Semasa Boleh Dikekalkan",
        "현재 처방 진행 불가": "Tidak Boleh Diteruskan dengan Formula Semasa",
        "현재 처방으로 진행 불가": "Tidak Boleh Diteruskan dengan Formula Semasa",
        "확인 완료 전 판단 보류": "Keputusan Ditangguhkan sehingga Pengesahan Selesai",
        "제한조건 충족 후 진행 가능": "Boleh Diteruskan selepas Semua Syarat Sekatan Dipenuhi",
        "표시·고지요건 반영 후 진행 가능": "Boleh Diteruskan selepas Keperluan Pelabelan dan Pemberitahuan Dilaksanakan",
        "증빙자료 확인 필요": "Dokumen Sokongan Diperlukan",
        "할랄 위험 후보 미확인": "Tiada Petunjuk Risiko HALAL Dikenal Pasti",
        "위험 후보 미확인": "Tiada Petunjuk Risiko HALAL Dikenal Pasti",
        "원료 기원 확인 필요": "Pengesahan Asal Ramuan Diperlukan",
    },
    "sw": {
        "통과": "Imepita uchunguzi",
        "금지": "Imepigwa marufuku",
        "제한": "Imewekewa masharti",
        "경고·표시 필요": "Onyo/Uwekaji lebo unahitajika",
        "규제 대상": "Uthibitishaji wa upeo wa kanuni unahitajika",
        "수동 검토": "Mapitio ya kanuni kwa mkono yanahitajika",
        "명칭 검증 필요": "Uthibitishaji wa utambulisho wa kiambato unahitajika",
        "부적합": "Haiwezi kuendelea kwa fomula ya sasa",
        "확인 필요": "Uthibitishaji unahitajika",
        "조건부 가능": "Inaweza kuendelea kwa masharti",
        "성분 규제 스크리닝 통과": "Uchunguzi wa kanuni za viambato umepita",
        "분석결과에 따른 필수 조치": "Hatua zinazohitajika kulingana na matokeo ya uchunguzi",
        "필요한 자료": "Nyaraka zinazohitajika",
        "조치 완료 기준": "Vigezo vya kukamilisha",
        "현재 처방 유지 가능": "Fomula ya sasa inaweza kudumishwa",
        "현재 처방 진행 불가": "Haiwezi kuendelea kwa fomula ya sasa",
        "현재 처방으로 진행 불가": "Haiwezi kuendelea kwa fomula ya sasa",
        "확인 완료 전 판단 보류": "Uamuzi umesitishwa hadi uthibitishaji ukamilike",
        "제한조건 충족 후 진행 가능": "Inaweza kuendelea baada ya masharti yote ya vizuizi kutimizwa",
        "표시·고지요건 반영 후 진행 가능": "Inaweza kuendelea baada ya mahitaji ya uwekaji lebo na taarifa kutekelezwa",
        "증빙자료 확인 필요": "Nyaraka za kuthibitisha zinahitajika",
        "할랄 위험 후보 미확인": "Hakuna viashiria vya hatari ya HALAL vilivyotambuliwa",
        "위험 후보 미확인": "Hakuna viashiria vya hatari ya HALAL vilivyotambuliwa",
        "원료 기원 확인 필요": "Uthibitishaji wa asili ya kiambato unahitajika",
    },
}

ADDITIONAL_LANGUAGE_DISALLOWED_PATTERNS: Dict[str, Tuple[Tuple[str, str], ...]] = {
    "zh-CN": (
        (r"完全合规", "筛查结果不得表述为完全合规。"),
        (r"合规已确认", "筛查结果不得表述为已确认合规。"),
        (r"获准进入市场|已批准上市", "筛查结果不得表述为已获市场准入批准。"),
    ),
    "ar": (
        (r"متوافق بالكامل", "لا يجوز وصف نتيجة الفحص بأنها امتثال كامل."),
        (r"تم تأكيد الامتثال", "لا يجوز وصف نتيجة الفحص بأنها امتثال مؤكد."),
        (r"تمت الموافقة على دخول السوق", "لا يجوز وصف الفحص بأنه موافقة على دخول السوق."),
    ),
    "fr": (
        (r"Entièrement conforme", "Le filtrage ne doit pas être présenté comme une conformité totale."),
        (r"Conformité confirmée", "Le filtrage ne doit pas être présenté comme une conformité confirmée."),
        (r"Entrée sur le marché approuvée", "Le filtrage ne constitue pas une autorisation de mise sur le marché."),
    ),
    "id": (
        (r"Sepenuhnya patuh", "Hasil penyaringan tidak boleh dinyatakan sebagai kepatuhan penuh."),
        (r"Kepatuhan dikonfirmasi", "Hasil penyaringan tidak boleh dinyatakan sebagai kepatuhan yang telah dikonfirmasi."),
        (r"Masuk pasar disetujui|Akses pasar disetujui", "Penyaringan bukan persetujuan masuk pasar."),
    ),
    "ms": (
        (r"Patuh sepenuhnya", "Hasil saringan tidak boleh dinyatakan sebagai pematuhan penuh."),
        (r"Pematuhan disahkan", "Hasil saringan tidak boleh dinyatakan sebagai pematuhan yang disahkan."),
        (r"Kemasukan pasaran diluluskan", "Saringan bukan kelulusan kemasukan pasaran."),
    ),
    "sw": (
        (r"Inatii kikamilifu", "Matokeo ya uchunguzi hayapaswi kudai uzingatiaji kamili."),
        (r"Uzingatiaji umethibitishwa", "Matokeo ya uchunguzi hayapaswi kudai uzingatiaji uliothibitishwa."),
        (r"Kuingia sokoni kumeidhinishwa", "Uchunguzi si idhini ya kuingia sokoni."),
    ),
}

ADDITIONAL_LANGUAGE_NORMALIZATION_PATTERNS: Dict[str, Tuple[Tuple[str, str], ...]] = {
    "zh-CN": ((r"禁用", "禁止"), (r"需人工审核", "需要人工法规审查")),
    "ar": ((r"ممنوع", "محظور"), (r"مراجعة يدوية مطلوبة", "تلزم مراجعة تنظيمية يدوية")),
    "fr": ((r"Prohibé", "Interdit"), (r"Vérification manuelle requise", "Examen réglementaire manuel requis")),
    "id": ((r"Terlarang", "Dilarang"), (r"Tinjauan manual diperlukan", "Tinjauan Regulasi Manual Diperlukan")),
    "ms": ((r"Terlarang", "Dilarang"), (r"Semakan manual diperlukan", "Semakan Peraturan Manual Diperlukan")),
    "sw": ((r"^Marufuku$", "Imepigwa marufuku"), (r"Ukaguzi wa mikono unahitajika", "Mapitio ya kanuni kwa mkono yanahitajika")),
}

ADDITIONAL_LANGUAGE_FORBIDDEN_ENGLISH_TERMS: Tuple[str, ...] = (
    "Passed",
    "Prohibited",
    "Restricted",
    "Warning/Labeling Required",
    "Regulatory Scope Verification Required",
    "Manual Regulatory Review Required",
    "Ingredient Identity Verification Required",
    "Required Documentation",
    "Completion Criteria",
    "Ingredient Regulatory Screening Passed",
)


# Plain-language English templates for customer-facing reports.
# These override older internal or literal translations.
ENGLISH_CANONICAL_EXACT.update({
    "분석 시장": "Markets Screened",
    "HALAL 추가검토": "Additional HALAL Review",
    "선택함": "Selected",
    "선택하지 않음": "Not Selected",
    "분석 성분 수": "Number of Ingredients Screened",
    "제품 성분": "Product Ingredients",
    "시장별 결과 요약": "Summary by Market",
    "검토 구분": "Review Type",
    "시장": "Market",
    "분석 결과": "Screening Result",
    "의미": "What This Means",
    "화장품 성분 규제": "Cosmetics Ingredient Regulation",
    "시장별 핵심 결론": "Key Findings by Market",
    "보고서의 범위": "Scope and Limitations",
    "이번 분석 결과": "Screening Result",
    "성분별 분석 결과": "Ingredient-Level Results",
    "설명": "Explanation",
    "해야 할 일": "Required Action",
    "완료 기준": "Completion Criteria",
    "HALAL 검토가 이 시장에서 필요한 경우": "When HALAL Review May Be Relevant to This Market",
    "질문": "Question",
    "답변": "Answer",
    "일반 화장품 판매에 직접 필요한가?": "Is HALAL review directly required for ordinary cosmetics sales?",
    "HALAL 자료가 필요할 수 있는 경우": "When HALAL documentation may be required",
    "이번 성분 규제 결과가 달라지는가?": "Does the HALAL review change this ingredient-regulatory result?",
    "이번 HALAL 검토 결과": "HALAL Finding in This Review",
    "이번 분석 범위에서는 성분 변경 필요 없음": "No Ingredient Change Required Based on This Screening",
    "현재 성분명 기준으로 추가 확인할 항목 없음": "No Additional Ingredient-Level Verification Identified",
    "HALAL 확인 성분": "Ingredient Requiring HALAL Verification",
    "검토 상태": "Review Status",
    "시장에 미치는 영향": "Potential Market Impact",
    "추가로 확인할 항목": "Items Requiring Additional Verification",
    "확인 이유": "Reason for Verification",
    "다음 단계": "Next Steps",
    "성분별 검토 결과": "Ingredient-Level HALAL Review",
    "확인 이유": "Reason for Verification",
    "확인 유형": "Verification Type",
    "중요 안내": "Important Note",
    "규제 스크리닝 및 HALAL 추가검토": "Regulatory Screening and Additional HALAL Review",

    "금지 성분 확인": "Prohibited Ingredient Identified",
    "성분 변경 후 재검토 필요": "Revise the Formula and Re-Screen",
    "금지 성분이 확인되었습니다. 해당 성분을 제거하거나 대체한 뒤 전체 성분 구성을 다시 분석해야 합니다.": "A prohibited ingredient was identified. Remove or replace the ingredient, then re-screen the complete ingredient list.",
    "추가 확인 필요": "Additional Verification Required",
    "필요한 자료 확인 전 결론 보류": "Decision Pending Required Documentation",
    "수동 검토 또는 성분 명칭 확인이 필요한 항목이 남아 있습니다. 지정된 자료를 확인하기 전에는 이 시장의 성분 규제 결과를 확정할 수 없습니다.": "One or more items require manual regulatory review or ingredient-identity verification. The screening result cannot be finalized until the specified documentation has been reviewed.",
    "규정 적용 여부 확인 필요": "Regulatory Scope Verification Required",
    "제품 유형과 용도 확인 후 결정": "Determine Applicability After Verifying Product Type and Intended Use",
    "규정 적용 가능성이 있는 성분이 확인되었습니다. 제품 유형, 사용 목적, 표시·광고 문구를 해당 규정의 적용 범위와 비교한 뒤 진행 여부를 결정해야 합니다.": "An ingredient that may fall within a regulatory scope was identified. Compare the product type, intended use, and claims with the applicable rule before deciding how to proceed.",
    "제한조건 확인 필요": "Restriction Conditions Require Verification",
    "제한조건 충족 여부 확인 후 진행 판단": "Proceed Only After Confirming All Restriction Conditions",
    "제한 성분이 확인되었습니다. 실제 배합농도와 제품 유형, 사용 부위 및 사용 대상을 해당 제한조건과 비교해야 합니다. 조건을 충족한다는 자료를 확보한 뒤 진행 여부를 결정하십시오.": "A restricted ingredient was identified. Compare the actual concentration, product type, application area, and target user with the applicable restriction conditions. Proceed only after documented confirmation that all conditions are met.",
    "표시 또는 고지사항 반영 필요": "Labeling or Notification Requirements Identified",
    "필수 표시사항 반영 후 진행 판단": "Proceed After Implementing the Required Labeling or Notification Measures",
    "표시, 경고, 고지 또는 통지 요건이 확인되었습니다. 필요한 문구와 절차를 라벨 및 제출자료에 반영한 뒤 진행 여부를 결정하십시오.": "A labeling, warning, disclosure, or notification requirement was identified. Implement the required wording and procedures in the label and submission materials before proceeding.",
    "성분 규제상 즉시 수정할 사항 없음": "No Immediate Ingredient-Regulatory Action Identified",
    "이번 분석 범위에서는 성분 변경 불필요": "No Ingredient Change Required Based on This Screening",
    "현재 분석에 사용된 규제 데이터에서는 금지 성분, 제한 성분, 표시 의무 또는 추가 확인 대상이 발견되지 않았습니다. 따라서 성분 규제와 관련하여 즉시 수정할 항목은 없습니다. 다만 제품 등록, 라벨, 광고 및 통관 요건은 별도로 확인해야 합니다.": "No prohibited, restricted, labeling, or verification item was identified in the currently loaded screening data. No immediate ingredient-related change is required. Product registration, labeling, claims, and customs requirements must still be reviewed separately.",

    "HALAL 금지 성분 확인": "HALAL-Prohibited Ingredient Identified",
    "원료 변경 후 재검토 필요": "Replace the Ingredient and Re-Screen",
    "HALAL 기준상 금지 성분이 확인되었습니다. 해당 원료를 제거하거나 대체한 뒤 다시 검토해야 합니다.": "An ingredient prohibited under the applicable HALAL criteria was identified. Remove or replace the ingredient, then repeat the HALAL review.",
    "원료 또는 증빙자료 확인 필요": "Ingredient or Supporting Documentation Requires Verification",
    "자료 확인 전 인증 가능 여부 판단 불가": "Certification Readiness Cannot Be Determined Until Documentation Is Reviewed",
    "확정적인 금지 성분은 확인되지 않았습니다. 다만 원료 기원, 복합원료 구성, 제조공정 또는 인증자료를 확인해야 하므로 현재 단계에서는 HALAL 인증 가능 여부를 판단할 수 없습니다.": "No confirmed prohibited ingredient was identified. However, ingredient origin, compound-ingredient composition, manufacturing process, or certification documentation requires verification. HALAL certification readiness cannot be determined at this stage.",
    "인증기관 조건 확인 후 준비 가능": "Certification Preparation May Begin After Verifying Certification-Body Requirements",
    "금지 성분은 확인되지 않았습니다. 관련 사용조건과 인증기관 요구사항을 확인한 뒤 인증 준비 여부를 결정하십시오.": "No prohibited ingredient was identified. Verify the applicable use conditions and certification-body requirements before deciding whether to begin certification preparation.",
    "성분명 기준 추가 확인 대상 없음": "No Additional Ingredient-Level Verification Identified",
    "현재 성분명과 분석 데이터에서는 HALAL 위험 후보가 확인되지 않았습니다. 이 결과는 완제품의 HALAL 인증 적합성 또는 인증 완료를 의미하지 않습니다.": "No HALAL risk indicator was identified from the ingredient names and screening data reviewed. This does not mean that the finished product is HALAL-compliant or certified.",

    "일반 화장품 판매에는 직접 필요하지 않습니다.": "No, not for ordinary cosmetics sales.",
    "수입자, 유통사 또는 인증기관의 요구에 따라 필요할 수 있습니다.": "It may be required by the importer, distributor, or certification body.",
    "판매 국가와 유통채널에 따라 다릅니다.": "It depends on the destination country and sales channel.",
    "회원국과 거래 상대방의 요구에 따라 다릅니다.": "It depends on the member country and the requirements of the trading partner.",
    "제품을 할랄 제품으로 판매하거나 바이어·유통채널이 HALAL 인증자료를 요구하는 경우": "When the product is marketed as HALAL or a buyer or distribution channel requires HALAL certification documentation.",
    "할랄 표시를 사용하거나 거래 상대방이 원료 기원 또는 인증자료를 요구하는 경우": "When HALAL labeling is used or the trading partner requires ingredient-origin or certification documentation.",
    "할랄 인증이 요구되는 국가에 판매하거나 바이어·유통채널이 관련 자료를 요구하는 경우": "When selling in a country that requires HALAL certification, or when a buyer or distribution channel requests supporting documentation.",
    "수입자, 바이어 또는 유통채널이 할랄 표시나 인증자료를 요구하는 경우": "When an importer, buyer, or distribution channel requires HALAL labeling or certification documentation.",
    "달라지지 않습니다. 이 시장의 성분 규제 분석과 HALAL 검토는 별도로 판단합니다.": "No. The ingredient-regulatory screening and the HALAL review are separate assessments.",
    "달라지지 않습니다. 화장품 성분 규제 분석과 HALAL 검토는 별도로 판단합니다.": "No. The cosmetics ingredient-regulatory screening and the HALAL review are separate assessments.",
    "달라지지 않습니다. 국가별 화장품 성분 규제와 HALAL 요건은 별도로 확인합니다.": "No. Country-specific cosmetics ingredient regulations and HALAL requirements are assessed separately.",
    "달라지지 않습니다. EAEU 성분 규제 분석과 HALAL 검토는 별도로 판단합니다.": "No. The EAC ingredient-regulatory screening and the HALAL review are separate assessments.",

    "현재 분석 결과, 성분 규제와 관련하여 즉시 수정해야 할 항목은 없습니다.": "No immediate ingredient-related change is required based on this screening.",
    "제품 등록, 라벨, 광고 및 통관 요건은 별도로 확인해야 합니다.": "Product registration, labeling, claims, and customs requirements must still be reviewed separately.",
    "현재 성분명 기준으로 별도 확인이 필요한 HALAL 위험 후보는 발견되지 않았습니다.": "No HALAL risk indicator requiring additional ingredient-level verification was identified from the current ingredient names.",
    "위 조치를 완료하고, 완료를 입증할 자료를 보관하십시오.": "Complete the required actions and retain evidence showing that each action has been completed.",
    "HALAL 인증 또는 관련 유통채널을 추진하는 경우 HALAL 장의 확인사항도 완료하십시오.": "When pursuing HALAL certification or a HALAL-related sales channel, also complete the verification items in the HALAL chapter.",
    "국가별 화장품 성분 규제와 별도로 원료 기원, 제조공정 및 인증자료를 확인합니다.": "This chapter reviews ingredient origin, manufacturing processes, and certification documentation separately from country-specific cosmetics ingredient regulations.",
    "최종 인증 여부는 공급사 자료, 제조공정 및 인증기관 심사를 통해 결정됩니다.": "Final certification status is determined through supplier documentation, manufacturing-process review, and assessment by the certification body.",
    "필요한 자료를 확보하고 확인 결과를 기록하십시오.": "Obtain the required documentation and record the verification outcome.",
    "적용할 HALAL 인증기관과 인증 범위를 정하십시오.": "Select the HALAL certification body and define the certification scope.",
    "원료 기원, 완제품 제조공정, 가공보조제 및 교차오염 관리자료를 준비하십시오.": "Prepare documentation covering ingredient origin, finished-product manufacturing, processing aids, and cross-contamination controls.",
    "인증기관이 요구하는 자료를 제출하고 최종 심사를 진행하십시오.": "Submit the required documentation to the certification body and proceed with the final assessment.",
    "이 검토 결과는 HALAL 인증 완료를 의미하지 않습니다.": "This review does not mean that HALAL certification has been completed.",
    "최종 판단은 공급사 증빙, 제조공정, 교차오염 관리 및 인증기관 심사를 통해 이루어집니다.": "The final determination depends on supplier evidence, manufacturing processes, cross-contamination controls, and review by the certification body.",

    "이 장에서는 해당 시장의 성분 분석 결과와 필요한 조치를 설명합니다.": "This chapter explains the ingredient-level findings and required actions for this market.",
    "HALAL 검토는 국가별 화장품 성분 규제 분석과 별개입니다. 각 시장 장에서는 HALAL 자료가 실제로 필요한 경우만 간단히 설명합니다.": "The HALAL review is separate from country-specific cosmetics ingredient regulation. Each market chapter explains only when HALAL documentation may be relevant.",
    "이 보고서는 화장품 성분 규제와 HALAL 관련 위험 후보를 사전에 확인하기 위한 자료입니다. 제품 등록, 통관 승인, 판매 허가 또는 HALAL 인증을 보장하지 않습니다. 실제 진행 전 최신 공식 규정, 배합농도, 제품 유형, 라벨, 공급사 자료 및 인증기관 요구사항을 별도로 확인해야 합니다.": "This report is a preliminary screening of cosmetics ingredient regulations and potential HALAL risk indicators. It does not guarantee product registration, customs clearance, sales authorization, or HALAL certification. Before proceeding, verify the latest official rules, actual concentrations, product type, labeling, supplier documentation, and certification-body requirements.",
    "전체 {len(ingredients)}개 성분을 분석했습니다. 전체 성분 목록과 성분별 결과는 함께 제공되는 Excel 파일에서 확인할 수 있습니다.": "The full ingredient list and ingredient-level results are provided in the accompanying Excel file.",
    "추가 확인이 필요한 성분은 없습니다. 전체 결과는 Excel 파일에서 확인할 수 있습니다.": "No ingredient requires additional verification. The complete results are provided in the accompanying Excel file.",
})


def _build_translation_glossary_prompt(language_code: str) -> str:
    glossary = ADDITIONAL_LANGUAGE_CANONICAL_EXACT.get(language_code)
    if not glossary:
        return ""
    lines = [
        "The source text is polished English. Use the following controlled regulatory terminology exactly:"
    ]
    preferred_keys = (
        "통과", "금지", "제한", "경고·표시 필요", "규제 대상", "수동 검토",
        "명칭 검증 필요", "성분 규제 스크리닝 통과", "분석결과에 따른 필수 조치",
        "필요한 자료", "조치 완료 기준", "할랄 위험 후보 미확인",
    )
    for key in preferred_keys:
        english_term = ENGLISH_CANONICAL_EXACT.get(key, key)
        lines.append(f"- {english_term} = {glossary[key]}")
    lines.extend([
        "Do not translate a screening result as full compliance, confirmed compliance, approval, or market-entry authorization.",
        "Use the local-language equivalent of regulatory screening, not regulatory approval or legal clearance.",
        "Write HALAL in uppercase except when preserving an official quoted title.",
        "Prefer a short natural sentence over a literal translation of English word order.",
    ])
    return "\n".join(lines)


TRANSLATION_GLOSSARY_PROMPTS: Dict[str, str] = {
    "en": ENGLISH_TRANSLATION_GLOSSARY_PROMPT,
    **{
        code: _build_translation_glossary_prompt(code)
        for code in ADDITIONAL_LANGUAGE_CANONICAL_EXACT
    },
}


# Full-sentence canonical English used by the default report. Keeping these
# statements deterministic prevents semantic drift between products and runs.
ENGLISH_CANONICAL_EXACT.update({
    "항목": "Item",
    "내용": "Details",
    "추가검토": "Additional Review",
    "성분 기원·구성·인증자료 검토": "Ingredient Origin, Composition, and Certification Documentation Review",
    "제1장": "Chapter 1",
    "제2장": "Chapter 2",
    "제3장": "Chapter 3",
    "제4장": "Chapter 4",
    "번호": "No.",
    "입력 성분명": "Input Ingredient Name",
    "INCI 명칭": "INCI Name",
    "CAS 번호": "CAS Number",
    "규제 판정": "Regulatory Screening Status",
    "필요한 조치": "Required Action",
    "필수 조치": "Required Action",
    "현재 로드된 해당 시장의 확정 규제 및 검토대기 데이터에서 일치하는 규제 항목이 확인되지 않았습니다.": "No matching regulatory item was identified in the confirmed or pending-review data currently loaded for this market.",
    "현재 로드된 규제 데이터에서 금지·제한·표시·추가 확인 대상이 확인되지 않았습니다. 성분 규제상 필수 보완조치는 없으며, 현재 처방을 유지한 상태로 제품 등록·표시·통관 등 다음 준비단계를 진행할 수 있습니다.": "No prohibited, restricted, labeling, or additional-verification item was identified in the currently loaded regulatory data. No ingredient-regulatory corrective action is required, and the current formula may proceed to product registration, labeling, customs, and other market-entry preparation steps.",
    "회원국·수입자·유통채널별 조건부 관련": "Conditionally Relevant by Member State, Importer, or Distribution Channel",
    "성분 규제상 필수 보완조치가 없습니다. 현재 처방을 유지할 수 있습니다.": "No ingredient-regulatory corrective action is required. The current formula may be maintained.",
    "이 결론은 현재 로드된 성분 규제 데이터에 대한 스크리닝 결과이며 제품 등록·표시·통관 완료를 의미하지 않습니다.": "This conclusion is based on screening against the currently loaded ingredient-regulatory data and does not mean that product registration, labeling review, or customs clearance has been completed.",
    "국가 규제 수동확인 항목이 없습니다.": "No national-regulatory manual review item was identified.",
    "확정 규제 데이터에서 금지 항목과 일치했습니다.": "The ingredient matched a prohibited-item entry in the confirmed regulatory data.",
    "금지 적용 범위, 동일물질 여부 및 예외조항 확인": "Verify the scope of the prohibition, substance identity, and any applicable exemptions.",
    "해당 성분을 제거하거나 허용 가능한 대체원료로 변경한 뒤 전체 처방을 재분석": "Remove the ingredient or replace it with a permitted alternative, then re-screen the complete formula.",
    "금지 성분 제거·대체 및 변경 처방 재분석 완료": "Complete formula revision and re-screening after removing or replacing the prohibited ingredient.",
    "제한 규제 항목과 일치했습니다. 실제 배합농도, 제품 유형, 사용 부위·대상과 허용조건을 확인해야 합니다.": "The ingredient matched a restricted-item entry. Verify the actual concentration, product type, application area, target users, and all applicable conditions of use.",
    "실제 배합농도, 제품유형, 사용부위, 사용대상 및 용도 확인": "Verify the actual concentration, product type, application area, target users, and intended use.",
    "실제 배합농도·제품유형·사용부위·사용대상을 제한조건과 대조하고 필요 시 처방 조정": "Compare the actual concentration, product type, application area, and target users against the restriction conditions, and revise the formula if necessary.",
    "제한조건 충족 증빙 확보 또는 기준에 맞춘 처방 변경 완료": "Obtain evidence that all restriction conditions are met, or complete a formula revision that meets the applicable requirements.",
    "표시·경고·고지 또는 통지 요건이 적용될 수 있어 최종 라벨과 제출자료 확인이 필요합니다.": "Warning, labeling, disclosure, or notification requirements may apply. Review the final label and submission materials.",
    "라벨, 경고문구, 고지, 신고 또는 통지 요건 확인": "Verify all applicable label, warning statement, disclosure, filing, or notification requirements.",
    "필수 표시·경고·고지·통지 요건을 라벨과 제출자료에 반영": "Implement all required labeling, warning, disclosure, and notification requirements in the label and submission materials.",
    "필수 표시·경고·고지·통지 반영 및 최종 라벨 검토 완료": "Complete implementation of all required labeling, warning, disclosure, and notification requirements and finalize the label review.",
    "관련 규제 항목과 일치했으나 실제 제품 적용 범위와 사용 목적을 추가 확인해야 합니다.": "The ingredient matched a regulatory entry, but the scope of application to the actual product and intended use requires further verification.",
    "해당 규정의 제품 적용 범위와 사용목적 확인": "Verify the scope of the regulation as applied to the product and its intended use.",
    "제품 유형·사용목적·표시광고와 규제 적용범위를 확인한 뒤 진행 여부 결정": "Verify the product type, intended use, claims, and regulatory scope before deciding whether to proceed.",
    "적용 범위와 제품 분류 확인 및 진행 근거 문서화": "Verify the regulatory scope and product classification, and document the basis for proceeding.",
    "검토대기 규제 자료와 일치하여 공식 원문, 시행상태와 사용조건을 수동으로 확인해야 합니다.": "The ingredient matched a pending-review regulatory entry. Manually verify the official source text, current effective status, and conditions of use.",
    "공식 규정 원문, 사용조건 및 최신 시행상태 확인": "Verify the official regulatory text, conditions of use, and current effective status.",
    "공식 규정 원문·시행상태·사용조건을 확인하고 검토 완료 후 재판정": "Verify the official regulatory text, effective status, and conditions of use, then reassess the screening decision.",
    "공식 원문·시행상태·사용조건 검토 및 최종 판정 기록": "Complete the review of the official source text, effective status, and conditions of use, and record the final review decision.",
    "공식 INCI 명칭 또는 CAS 번호가 확정되지 않아 규제 적용 여부를 완료할 수 없습니다.": "The regulatory applicability assessment cannot be completed because the official INCI name or CAS number has not been confirmed.",
    "공식 INCI 명칭, CAS 번호 및 공급사 원료자료 확인": "Verify the official INCI name, CAS number, and supplier ingredient documentation.",
    "공식 INCI·CAS·공급사 원료자료를 확보한 뒤 재분석": "Obtain the official INCI name, CAS number, and supplier ingredient documentation, then re-screen the ingredient.",
    "공식 명칭·CAS·원료규격서 확인 및 재분석 완료": "Confirm the official ingredient identity, CAS number, and raw-material specification, and complete re-screening.",
    "전체 수동확인 필요": "Total Items Requiring Manual Review",
    "종합 판정": "Overall Screening Decision",
    "성분 규제 진행 상태": "Ingredient-Regulatory Screening Status",
    "최종결론": "Screening Conclusion",
    "금지 성분이 확인되어 현재 처방으로는 해당 시장의 성분 규제 단계를 통과할 수 없습니다. 해당 성분을 제거하거나 대체한 뒤 전체 처방을 다시 분석해야 합니다.": "A prohibited ingredient match was identified. The current formula cannot pass ingredient regulatory screening for this market. Remove or replace the ingredient and re-screen the complete formula.",
    "현지 제품 등록·신고 또는 사전통지 절차 확인": "Verify the applicable local product registration, filing, or pre-notification procedure.",
    "제품 분류·효능표현·표시광고 및 라벨 언어 요건 검토": "Review product classification, efficacy claims, advertising statements, and label-language requirements.",
    "책임자·수입자·현지 대리인과 통관 제출자료 준비": "Prepare the responsible-person, importer, local-representative, and customs-submission documentation.",
    "처방·제품유형·표시내용 또는 적용 규정이 변경된 경우 재분석": "Re-screen the product if the formula, product type, labeling content, or applicable regulation changes.",
    "검토 기준 DB": "Screening Database",
    "DB 업데이트일": "Database Update Date",
    "보고서 생성일": "Report Generation Date",
    "사용 범위": "Scope of Use",
    "화장품 성분 규제 사전 스크리닝": "Preliminary Cosmetics Ingredient Regulatory Screening",
    "적용": "Selected",
    "별도 인증 목적 또는 조건부 관련": "Relevant for Separate Certification Purposes or Under Specific Conditions",
    "변경 없음": "No Change",
    "해당 시장의 국가 화장품 규제판정을 변경하지는 않지만, 할랄 제품으로 판매하거나 바이어·유통채널이 인증자료를 요구하는 경우 시장 진입 준비에 영향을 줄 수 있습니다.": "This does not change the national cosmetics regulatory screening decision for the market. However, it may affect market-entry preparation when the product is marketed as HALAL or when buyers or distribution channels require certification documentation.",
    "연결된 HALAL 확인 성분": "Linked HALAL Verification Ingredients",
    "HALAL 필수 조치": "Required HALAL Actions",
    "수동확인 체크리스트": "Manual Regulatory Review Checklist",
    "성분명": "Ingredient",
    "확인 사유": "Reason for Verification",
    "규제 적용범위": "Regulatory Scope",
    "제품 분류·용도·표시광고·규제 적용범위 근거": "Product classification, intended use, claims, and evidence supporting the regulatory scope determination.",
    "규제 원문·적용조건": "Official Regulatory Text and Applicable Conditions",
    "공식 규정 원문·시행일·사용조건 검토기록": "Review record covering the official regulatory text, effective date, and conditions of use.",
    "공식 명칭·CAS": "Official Ingredient Identity and CAS Number",
    "공식 INCI·CAS·공급사 원료규격서": "Official INCI name, CAS number, and supplier raw-material specification.",
    "검출 방식": "Identification Method",
    "공급사 확인사항": "Supplier Verification Item",
    "현재 성분명 기준 HALAL 위험 후보가 확인되지 않았습니다.": "No HALAL risk indicator was identified based on the current ingredient identity.",
    "원료 기원·구성·제조공정 또는 인증자료를 공급사 자료로 확인해야 합니다.": "Verify the ingredient origin, composition, manufacturing process, and certification documentation using supplier records.",
    "원료규격서·기원서·전체 구성표·제조공정서·유효한 인증자료": "Raw-material specification, origin statement, full composition, manufacturing-process documentation, and valid certification documents.",
    "식물성·합성·동물성 기원을 공급사에 확인하십시오.": "Confirm with the supplier whether the ingredient is plant-derived, synthetic, or animal-derived.",
    "원료 기원·전체 구성·제조공정·가공보조제·인증자료 확인 완료": "Complete verification of ingredient origin, full composition, manufacturing process, processing aids, and certification documentation.",
    "검토 구분": "Review Category",
    "할랄 추가 성분 스크리닝": "Additional HALAL Ingredient Screening",
    "전체 확인 필요": "Total Items Requiring Verification",
    "할랄 인증 추진 가능성": "HALAL Certification Readiness",
    "현재 확인된 확정 금지 성분은 없으나 원료 기원, 복합원료 구성, 제조공정 또는 인증자료 확인이 필요하여 할랄 인증 추진 가능성을 아직 확정할 수 없습니다.": "No confirmed HALAL-prohibited ingredient was identified, but ingredient origin, compound-ingredient composition, manufacturing process, or certification documentation still requires verification. HALAL certification readiness cannot yet be determined.",
    "할랄 인증서 또는 최종 인증판정이 아닌 사전 위험 선별": "Preliminary risk screening only; not a HALAL certificate or final certification decision.",
    "1. 제품 성분 구성": "1. Product Ingredient Composition",
    "2. 시장별 종합판정": "2. Overall Screening Decision by Market",
    "3. 통합 최종결론": "3. Integrated Screening Conclusion",
    "HALAL 결과는 독립 장에서 상세히 설명하며, 각 시장 장에서는 해당 결과가 현지 인증·바이어·유통채널 요구에 미치는 영향만 연결하여 표시합니다.": "The HALAL review is presented in a separate chapter. Each market chapter only explains how the HALAL findings may affect local certification, buyer, or distribution-channel requirements.",
    "4. 면책 및 사용 범위": "4. Disclaimer and Scope of Use",
    "본 보고서는 국가별 화장품 성분 규제 및 할랄 위험 후보를 사전에 선별하는 문서입니다. 제품 등록, 통관 승인, 판매 허가, 할랄 인증서 또는 최종 법률 판단을 대신하지 않습니다. 실제 진행 전 최신 공식 규정, 처방 농도, 제품 유형, 라벨, 공급사 자료 및 인증기관 요구사항을 별도로 확인해야 합니다.": "This report is a preliminary screening document for country-specific cosmetics ingredient regulations and HALAL risk indicators. It does not replace product registration, customs clearance, sales authorization, a HALAL certificate, or a final legal determination. Before proceeding, independently verify the latest official regulations, formula concentrations, product type, labeling, supplier documentation, and certification-body requirements.",
    "시장별 분석은 이전 시장과 분리하여 새 장에서 시작합니다.": "Each market analysis begins in a separate chapter.",
    "1. 분석 세부내용": "1. Detailed Screening Results",
    "2. 분석결과에 따른 조치": "2. Actions Based on Screening Results",
    "4. 수동확인 항목": "4. Manual Regulatory Review Items",
    "5. 판정 해석 시 주의사항": "5. Important Interpretation Notes",
    "통과는 현재 로드된 데이터에서 일치 항목이 없다는 의미이며 최종 수입·판매 허가를 보증하지 않습니다.": "Passed means that no matching item was identified in the currently loaded screening data. It does not guarantee final import or sales authorization.",
    "제한·규제 대상은 실제 농도, 제품유형, 용도와 표시정보를 함께 확인해야 합니다.": "For Restricted or Regulatory Scope Verification Required items, review the actual concentration, product type, intended use, and labeling information together.",
    "수동검토와 명칭검증 항목은 확인 전 최종 적합 판정으로 취급하면 안 됩니다.": "Do not treat Manual Regulatory Review Required or Ingredient Identity Verification Required items as a final acceptable decision before verification is complete.",
    "HALAL 결과는 국가 화장품 규제판정을 변경하지 않으며 인증·바이어·유통요건에 대한 별도 영향으로 해석해야 합니다.": "HALAL findings do not change the national cosmetics regulatory screening decision and must be interpreted separately in relation to certification, buyer, and distribution requirements.",
    "6. 조치 완료 후 다음 단계": "6. Next Steps After Completing Required Actions",
    "위 분석결과별 필수 조치를 완료하고 완료 근거를 문서화하십시오.": "Complete each required action identified by the screening results and document the evidence of completion.",
    "HALAL 인증 또는 관련 유통채널을 추진하는 경우 HALAL 독립 장의 확인사항을 별도로 완료하십시오.": "When pursuing HALAL certification or related distribution channels, separately complete the verification items in the HALAL chapter.",
    "7. 용어의 정의": "7. Definitions",
    "국가 규제판정과 별도로 원료 기원·구성·제조공정·인증자료를 검토합니다.": "Ingredient origin, composition, manufacturing process, and certification documentation are reviewed separately from the national regulatory screening decision.",
    "3. 수동확인 항목": "3. Manual Verification Items",
    "4. 판정 해석 시 주의사항": "4. Important Interpretation Notes",
    "할랄 위험 후보 미확인은 할랄 인증 완료를 의미하지 않습니다.": "No HALAL Risk Indicators Identified does not mean that HALAL certification has been completed.",
    "성분명만으로 동물 종, 제조공정, 발효 배지, 가공보조제와 교차오염을 확정할 수 없습니다.": "The ingredient name alone cannot confirm animal species, manufacturing process, fermentation media, processing aids, or cross-contamination.",
    "최종 인증 가능 여부는 공급사 증빙과 적용 인증기관의 심사를 통해 확인해야 합니다.": "Final certification eligibility must be determined through supplier evidence and review by the applicable certification body.",
    "5. 조치 완료 후 다음 단계": "5. Next Steps After Completing Required Actions",
    "위 확인 대상별 자료를 확보하고 결과를 기록한 뒤 HALAL 재분석을 수행하십시오.": "Obtain the required documentation for each verification item, record the findings, and then re-screen the HALAL review.",
    "적용할 할랄 인증기관과 인증 범위를 선정하십시오.": "Select the applicable HALAL certification body and certification scope.",
    "완제품 제조공정·가공보조제·교차오염 관리자료를 준비하십시오.": "Prepare finished-product manufacturing-process, processing-aid, and cross-contamination control documentation.",
    "인증기관이 요구하는 원료·공정·시설 자료를 제출해 최종 심사를 진행하십시오.": "Submit the ingredient, process, and facility documentation required by the certification body for final review.",
    "6. 용어의 정의": "6. Definitions",
    "구분": "Category",
    "대상": "Target",
    "종합판정": "Overall Decision",
    "가능성 판단": "Readiness Assessment",
    "국가 규제": "National Regulation",
    "성분": "Ingredient",
    "판정": "Screening Status",
    "변경 처방·대체원료 규격서": "Revised formula and replacement-ingredient specification.",
    "배합농도·제품유형·사용조건·원료규격서": "Formula concentration, product type, conditions of use, and raw-material specification.",
    "최종 라벨·경고문구·고지 또는 통지 자료": "Final label, warning statements, disclosure, or notification documentation.",
    "HALAL 확인 대상": "HALAL Verification Ingredient",
    "검토 상태": "Review Status",
    "해당 시장 영향": "Impact on This Market",
    "현재 규제 DB에서 일치 항목이 확인되지 않은 상태": "No matching item was identified in the currently loaded regulatory database.",
    "현재 처방 기준 해당 시장 수입·유통·판매 진행이 불가능한 상태": "Based on the screening result, the current formula should not proceed to import, distribution, or sale in this market.",
    "농도·제품유형·사용대상 등 조건 충족이 필요한 상태": "Applicable concentration, product-type, target-user, or other conditions must be met.",
    "라벨·고지·경고문구·통지 요건 검토가 필요한 상태": "Label, disclosure, warning-statement, or notification requirements require review.",
    "관련 규제는 있으나 실제 적용범위를 추가 확인해야 하는 상태": "A relevant regulatory entry exists, but its actual scope of application requires further verification.",
    "공식 원문이나 공급사 자료를 사람이 확인해야 하는 상태": "The official source text or supplier documentation requires manual regulatory review.",
    "공식 INCI가 확정되지 않아 판정을 완료할 수 없는 상태": "The screening decision cannot be completed because the official INCI name has not been confirmed.",
    "추가 자료 없이는 최종 결론을 내릴 수 없는 상태": "A final screening conclusion cannot be reached without additional documentation.",
    "현재 처방 또는 표시 상태로 규제요건을 충족하지 못하는 상태": "The current formula or labeling does not meet the identified regulatory requirement.",
    "조건 충족 후 진행 가능": "Proceed After Requirements Are Met",
    "필수 제한·표시·자료 요건을 충족하면 진행 가능한 상태": "The product may proceed after all applicable restriction, labeling, and documentation requirements are met.",
    "Chemical Abstracts Service가 화학물질에 부여한 고유 식별번호": "A unique substance identifier assigned by the Chemical Abstracts Service.",
    "할랄": "HALAL",
    "이슬람 기준에 따라 허용되는 원료·공정·제품": "An ingredient, process, or product permitted under Islamic requirements.",
    "하람": "Haram",
    "이슬람 기준상 허용되지 않는 원료·공정·제품": "An ingredient, process, or product not permitted under Islamic requirements.",
    "식물성·합성·동물성 등 원료가 유래한 근원": "The source from which an ingredient is derived, such as plant, synthetic, or animal origin.",
    "동물 유래": "Animal-Derived",
    "동물 조직·지방·분비물 등에서 유래한 원료": "An ingredient derived from animal tissue, fat, secretions, or other animal sources.",
    "최종 성분표에 남지 않더라도 제조 중 사용되는 물질": "A substance used during manufacturing even if it is not present in the final ingredient list.",
    "비할랄 원료·설비·보관·운송 과정에서 혼입될 가능성": "The possibility of contamination from non-HALAL ingredients, equipment, storage, or transportation.",
    "인증기관이 원료 또는 제품의 적합성을 확인해 발급한 문서": "A document issued by a certification body confirming that an ingredient or product meets its HALAL requirements.",
    "현재 분석에서 우려 후보가 발견되지 않았으나 인증 완료는 아닌 상태": "No risk indicator was identified in the current review, but certification has not been completed.",
    "공급사 자료로 원료 유래와 제조공정을 확인해야 하는 상태": "Ingredient origin and manufacturing process must be verified using supplier documentation.",
    "K-Beauty Global Compliance | 다중시장 규제 및 HALAL 추가검토": "K-Beauty Global Compliance | Multi-Market Regulatory Screening and Additional HALAL Review",
})

# Final customer-facing English overrides. These must remain after all legacy
# dictionaries so older internal labels cannot overwrite the plain-language text.
ENGLISH_CANONICAL_EXACT.update({
    "검토 구분": "Review Type",
    "시장별 결과 요약": "Summary by Market",
    "HALAL 위험 후보 미확인": "No HALAL Risk Indicators Identified",
    "검토 결과": "Review Result",
    "보고서 사용 범위": "Scope of Use",
    "이 보고서는 성분 규제와 HALAL 위험 후보를 사전에 확인하기 위한 자료입니다. 제품 등록, 통관 승인, 판매 허가 또는 HALAL 인증을 보장하지 않으므로 실제 진행 전 최신 공식 기준과 제품 자료를 별도로 확인해야 합니다.": "This report is a preliminary screening of ingredient regulations and potential HALAL risk indicators. It does not guarantee product registration, customs clearance, sales authorization, or HALAL certification. Verify the latest official requirements and product documentation before proceeding.",
    "HALAL 검토는 국가별 화장품 성분 규제 분석과 별도로 판단합니다.": "The HALAL review is assessed separately from country-specific cosmetics ingredient regulation.",
    "HALAL 추가검토": "Additional HALAL Review",
    "이번 HALAL 검토 결과": "HALAL Finding in This Review",
    "이번 분석 결과": "Screening Result",
    "필요한 조치": "Required Actions",
    "즉시 필요한 성분 수정 없음": "No Immediate Ingredient Change Required",
    "금지 성분 제거 또는 대체 후 전체 성분 재분석": "Remove or replace the prohibited ingredient, then re-screen the complete ingredient list.",
    "제한 성분의 배합농도, 제품 유형 및 사용조건 확인": "Verify the concentration, product type, and conditions of use for each restricted ingredient.",
    "필수 표시·경고·고지 또는 통지 요건 반영": "Implement all required labeling, warning, disclosure, or notification measures.",
    "제품 유형과 사용 목적에 따른 규정 적용 여부 확인": "Verify whether the regulation applies to the product type and intended use.",
    "공식 규정 원문과 적용조건 수동 확인": "Manually verify the official regulatory text and applicable conditions.",
    "공식 INCI, CAS 번호 또는 공급사 자료 확인": "Verify the official INCI name, CAS number, or supplier documentation.",
    "HALAL 추가검토 여부": "Additional HALAL Review Selected",
    "이 시장과 HALAL 검토의 관련성": "Is HALAL review directly required for ordinary cosmetics sales?",
    "HALAL 자료가 필요한 경우": "When HALAL documentation may be required",
    "이 시장의 성분 규제 결과에 미치는 영향": "Does the HALAL review change this ingredient-regulatory result?",
    "HALAL 추가검토상 필수 보완조치 없음": "No Additional HALAL Corrective Action Required",
    "성분 규제 외 별도 절차": "Separate Non-Ingredient Market Procedures",
    "성분 규제상 추가 조치나 재분석은 필요하지 않습니다.": "No additional ingredient-regulatory action or re-screening is required.",
    "제품 등록·라벨·표시광고·통관은 성분 규제와 별도 절차이므로 필요한 경우 확인하십시오.": "Product registration, labeling, claims, and customs are separate from ingredient screening and should be reviewed only as applicable.",
    "성분 규제상 추가 조치 없음. 제품 등록·라벨·표시광고·통관은 별도 절차": "No additional ingredient-regulatory action. Product registration, labeling, claims, and customs are separate procedures.",
})

STATUS_ORDER = [
    "BANNED",
    "REVIEW_REQUIRED",
    "VERIFICATION_REQUIRED",
    "REGULATED",
    "RESTRICTED",
    "WARNING_REQUIRED",
    "PASS",
]

MANUAL_STATUSES = {
    "REVIEW_REQUIRED",
    "VERIFICATION_REQUIRED",
    "REGULATED",
}

NAVY = "294A63"
BLUE = "5AA6C9"
LIGHT_BLUE = "DDEEF5"
PALE = "F5F7F8"
GRAY = "E9EDF0"
YELLOW = "FFF4CC"
RED = "F4CCCC"
GREEN = "D9EAD3"
WHITE = "FFFFFF"
BORDER_COLOR = "C8CDD1"

DOC_TEXT = RGBColor(45, 45, 45)
DOC_MUTED = RGBColor(105, 105, 105)
DOC_NAVY = RGBColor(41, 74, 99)
DOC_GOLD = RGBColor(92, 72, 0)


# ============================================================
# 공통 유틸리티
# ============================================================
def clean_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    try:
        if value != value:
            return default
    except Exception:
        pass
    text = " ".join(str(value).strip().split())
    if text.casefold() in {"", "nan", "none", "null"}:
        return default
    return text


def safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return default


def normalize_status(value: Any) -> str:
    status = clean_text(value).upper()
    aliases = {
        "FAIL": "BANNED",
        "MANUAL_REVIEW": "REVIEW_REQUIRED",
        "MANUAL REVIEW": "REVIEW_REQUIRED",
        "VERIFICATION": "VERIFICATION_REQUIRED",
        "WARNING": "WARNING_REQUIRED",
    }
    return aliases.get(status, status or "VERIFICATION_REQUIRED")


def safe_filename(value: str, fallback: str = "제품") -> str:
    value = clean_text(value, fallback)
    value = re.sub(r"[\\/:*?\"<>|]+", "_", value)
    value = re.sub(r"\s+", "_", value).strip("._")
    return value or fallback


def dedupe_preserve(values: Iterable[str]) -> List[str]:
    result: List[str] = []
    seen: set[str] = set()
    for value in values:
        text = clean_text(value)
        key = text.casefold()
        if text and key not in seen:
            seen.add(key)
            result.append(text)
    return result


def result_details(result_data: dict) -> List[dict]:
    details = result_data.get("report_details") or []
    return [item for item in details if isinstance(item, dict)]


def status_counts(result_data: dict) -> Dict[str, int]:
    raw = result_data.get("status_counts") or {}
    return {
        status: safe_int(raw.get(status))
        for status in (
            "PASS",
            "BANNED",
            "RESTRICTED",
            "WARNING_REQUIRED",
            "REGULATED",
            "REVIEW_REQUIRED",
            "VERIFICATION_REQUIRED",
        )
    }


def dominant_status(result_data: dict) -> str:
    counts = status_counts(result_data)
    for status in STATUS_ORDER:
        if counts.get(status, 0) > 0:
            return status
    return "PASS"


def market_decision(result_data: dict) -> Dict[str, str]:
    status = dominant_status(result_data)
    if status == "BANNED":
        return {
            "overall": "금지 성분 확인",
            "possibility": "성분 변경 후 재검토 필요",
            "conclusion": "금지 성분이 확인되었습니다. 해당 성분을 제거하거나 대체한 뒤 전체 성분 구성을 다시 분석해야 합니다.",
        }
    if status in {"REVIEW_REQUIRED", "VERIFICATION_REQUIRED"}:
        return {
            "overall": "추가 확인 필요",
            "possibility": "필요한 자료 확인 전 결론 보류",
            "conclusion": "수동 검토 또는 성분 명칭 확인이 필요한 항목이 남아 있습니다. 지정된 자료를 확인하기 전에는 이 시장의 성분 규제 결과를 확정할 수 없습니다.",
        }
    if status == "REGULATED":
        return {
            "overall": "규정 적용 여부 확인 필요",
            "possibility": "제품 유형과 용도 확인 후 결정",
            "conclusion": "규정 적용 가능성이 있는 성분이 확인되었습니다. 제품 유형, 사용 목적, 표시·광고 문구를 해당 규정의 적용 범위와 비교한 뒤 진행 여부를 결정해야 합니다.",
        }
    if status == "RESTRICTED":
        return {
            "overall": "제한조건 확인 필요",
            "possibility": "제한조건 충족 여부 확인 후 진행 판단",
            "conclusion": "제한 성분이 확인되었습니다. 실제 배합농도와 제품 유형, 사용 부위 및 사용 대상을 해당 제한조건과 비교해야 합니다. 조건을 충족한다는 자료를 확보한 뒤 진행 여부를 결정하십시오.",
        }
    if status == "WARNING_REQUIRED":
        return {
            "overall": "표시 또는 고지사항 반영 필요",
            "possibility": "필수 표시사항 반영 후 진행 판단",
            "conclusion": "표시, 경고, 고지 또는 통지 요건이 확인되었습니다. 필요한 문구와 절차를 라벨 및 제출자료에 반영한 뒤 진행 여부를 결정하십시오.",
        }
    return {
        "overall": "성분 규제상 즉시 수정할 사항 없음",
        "possibility": "이번 분석 범위에서는 성분 변경 필요 없음",
        "conclusion": "현재 분석에 사용된 규제 데이터에서는 금지 성분, 제한 성분, 표시 의무 또는 추가 확인 대상이 발견되지 않았습니다. 따라서 성분 규제와 관련하여 즉시 수정할 항목은 없습니다. 다만 제품 등록, 라벨, 광고 및 통관 요건은 별도로 확인해야 합니다.",
    }


def halal_decision(result_data: dict) -> Dict[str, str]:
    status = dominant_status(result_data)
    if status == "BANNED":
        return {
            "overall": "HALAL 금지 성분 확인",
            "possibility": "원료 변경 후 재검토 필요",
            "conclusion": "HALAL 기준상 금지 성분이 확인되었습니다. 해당 원료를 제거하거나 대체한 뒤 다시 검토해야 합니다.",
        }
    if status in {"REVIEW_REQUIRED", "VERIFICATION_REQUIRED", "REGULATED"}:
        return {
            "overall": "원료 또는 증빙자료 확인 필요",
            "possibility": "자료 확인 전 인증 가능 여부 판단 불가",
            "conclusion": "확정적인 금지 성분은 확인되지 않았습니다. 다만 원료 기원, 복합원료 구성, 제조공정 또는 인증자료를 확인해야 하므로 현재 단계에서는 HALAL 인증 가능 여부를 판단할 수 없습니다.",
        }
    if status == "RESTRICTED":
        return {
            "overall": "조건 확인 필요",
            "possibility": "인증기관 조건 확인 후 준비 가능",
            "conclusion": "금지 성분은 확인되지 않았습니다. 관련 사용조건과 인증기관 요구사항을 확인한 뒤 인증 준비 여부를 결정하십시오.",
        }
    return {
        "overall": "HALAL 위험 후보 미확인",
        "possibility": "현재 성분명 기준으로 추가 확인할 항목 없음",
        "conclusion": "현재 성분명과 분석 데이터에서는 HALAL 위험 후보가 확인되지 않았습니다. 이 결과는 완제품의 HALAL 인증 적합성 또는 인증 완료를 의미하지 않습니다.",
    }


def market_impact_for_status(status: str) -> str:
    status = normalize_status(status)
    return {
        "PASS": "성분 규제상 필수 보완조치 없음",
        "BANNED": "현재 처방 진행 불가",
        "RESTRICTED": "제한조건 충족 후 진행 가능",
        "WARNING_REQUIRED": "표시·고지요건 반영 후 진행 가능",
        "REGULATED": "적용 범위 확인 후 결정",
        "REVIEW_REQUIRED": "확인 완료 전 판단 보류",
        "VERIFICATION_REQUIRED": "명칭·자료 확인 전 판단 보류",
    }.get(status, "확인 필요")

def applicable_condition(status: str) -> str:
    status = normalize_status(status)
    return {
        "PASS": "-",
        "BANNED": "금지 적용 범위, 동일물질 여부 및 예외조항 확인",
        "RESTRICTED": "실제 배합농도, 제품유형, 사용부위, 사용대상 및 용도 확인",
        "WARNING_REQUIRED": "라벨, 경고문구, 고지, 신고 또는 통지 요건 확인",
        "REGULATED": "해당 규정의 제품 적용 범위와 사용목적 확인",
        "REVIEW_REQUIRED": "공식 규정 원문, 사용조건 및 최신 시행상태 확인",
        "VERIFICATION_REQUIRED": "공식 INCI 명칭, CAS 번호 및 공급사 원료자료 확인",
    }.get(status, "추가 확인")


def recommended_action(status: str) -> str:
    status = normalize_status(status)
    return {
        "PASS": "성분 규제상 필수 보완조치 없음",
        "BANNED": "해당 성분을 제거하거나 허용 가능한 대체원료로 변경한 뒤 전체 처방을 재분석",
        "RESTRICTED": "실제 배합농도·제품유형·사용부위·사용대상을 제한조건과 대조하고 필요 시 처방 조정",
        "WARNING_REQUIRED": "필수 표시·경고·고지·통지 요건을 라벨과 제출자료에 반영",
        "REGULATED": "제품 유형·사용목적·표시광고와 규제 적용범위를 확인한 뒤 진행 여부 결정",
        "REVIEW_REQUIRED": "공식 규정 원문·시행상태·사용조건을 확인하고 검토 완료 후 재판정",
        "VERIFICATION_REQUIRED": "공식 INCI·CAS·공급사 원료자료를 확보한 뒤 재분석",
    }.get(status, "추가 확인")

def completion_criterion(status: str, target: str = "") -> str:
    status = normalize_status(status)
    if target == "HALAL":
        return {
            "PASS": "해당 없음(현재 성분명 기준 별도 HALAL 보완조치 없음)",
            "BANNED": "금지 원료 제거·대체 및 변경 처방 HALAL 재검토 완료",
            "RESTRICTED": "사용조건과 인증기관 요구사항 충족 증빙 확보",
            "WARNING_REQUIRED": "표시·인증 관련 요구사항 반영 완료",
            "REGULATED": "적용 할랄 기준과 인증기관 판단 범위 확인 완료",
            "REVIEW_REQUIRED": "공식 기준·인증기관 요구사항 검토 및 결과 기록",
            "VERIFICATION_REQUIRED": "원료 기원·전체 구성·제조공정·가공보조제·인증자료 확인 완료",
        }.get(status, "HALAL 확인사항 완료")
    return {
        "PASS": "해당 없음(현재 처방 유지)",
        "BANNED": "금지 성분 제거·대체 및 변경 처방 재분석 완료",
        "RESTRICTED": "제한조건 충족 증빙 확보 또는 기준에 맞춘 처방 변경 완료",
        "WARNING_REQUIRED": "필수 표시·경고·고지·통지 반영 및 최종 라벨 검토 완료",
        "REGULATED": "적용 범위와 제품 분류 확인 및 진행 근거 문서화",
        "REVIEW_REQUIRED": "공식 원문·시행상태·사용조건 검토 및 최종 판정 기록",
        "VERIFICATION_REQUIRED": "공식 명칭·CAS·원료규격서 확인 및 재분석 완료",
    }.get(status, "필요한 확인 완료")

def match_source_label(value: Any) -> str:
    source = clean_text(value)
    labels = {
        "no_database_match": "현재 DB 일치 없음",
        "confirmed_database": "확정 규제 DB",
        "review_database": "검토대기 DB",
        "translation_verification": "명칭 검증",
        "halal_prescreen_verification": "HALAL 사전검토",
        "fallback_keyword_check": "보조 키워드 검토",
    }
    return labels.get(source, source.replace("_", " ") if source else "-")


def normalized_reason(detail: dict, target: str = "") -> str:
    status = normalize_status(
        detail.get("compliance_status") or detail.get("restriction_type")
    )
    source = clean_text(detail.get("match_source"))
    if status == "PASS" and source == "no_database_match":
        if target == "HALAL":
            return "현재 성분명 기준 HALAL 위험 후보가 확인되지 않았습니다."
        return "현재 로드된 해당 시장의 확정 규제 및 검토대기 데이터에서 일치하는 규제 항목이 확인되지 않았습니다."
    if status == "VERIFICATION_REQUIRED" and source == "translation_verification":
        return "공식 INCI 명칭 또는 CAS 번호가 확정되지 않아 규제 적용 여부를 완료할 수 없습니다."
    if status == "VERIFICATION_REQUIRED" and target == "HALAL":
        return "원료 기원·구성·제조공정 또는 인증자료를 공급사 자료로 확인해야 합니다."

    raw = clean_text(
        detail.get("regulation_reason") or detail.get("regulation_notice")
    )
    # API 또는 규제 DB의 원문이 영문인 경우 한국어 보고서에 그대로 섞지
    # 않고 상태별 표준 해석문을 사용한다. 원문 식별자는 기술 데이터에 유지한다.
    if raw and not re.search(r"[가-힣]", raw) and len(re.findall(r"[A-Za-z]{2,}", raw)) >= 4:
        return {
            "BANNED": "확정 규제 데이터에서 금지 항목과 일치했습니다. 동일물질 여부, 적용 범위와 예외조항을 확인해야 합니다.",
            "RESTRICTED": "제한 규제 항목과 일치했습니다. 실제 배합농도, 제품 유형, 사용 부위·대상과 허용조건을 확인해야 합니다.",
            "WARNING_REQUIRED": "표시·경고·고지 또는 통지 요건이 적용될 수 있어 최종 라벨과 제출자료 확인이 필요합니다.",
            "REGULATED": "관련 규제 항목과 일치했으나 실제 제품 적용 범위와 사용 목적을 추가 확인해야 합니다.",
            "REVIEW_REQUIRED": "검토대기 규제 자료와 일치하여 공식 원문, 시행상태와 사용조건을 수동으로 확인해야 합니다.",
            "VERIFICATION_REQUIRED": "명칭 또는 원료자료가 충분하지 않아 공식 INCI, CAS 번호와 공급사 자료 확인이 필요합니다.",
        }.get(status, "규제 적용 근거와 조건을 추가 확인해야 합니다.")
    return raw or "규제 적용 근거와 조건을 추가 확인해야 합니다."

def action_for_detail(detail: dict, target: str = "") -> str:
    status = normalize_status(
        detail.get("compliance_status") or detail.get("restriction_type")
    )
    if target == "HALAL":
        if status == "PASS":
            return "현재 성분명 기준 별도 HALAL 보완조치 없음"
        return clean_text(
            detail.get("recommended_action"),
            {
                "BANNED": "금지 원료를 제거하거나 인증 가능한 대체원료로 변경한 뒤 재검토",
                "RESTRICTED": "사용조건과 인증기관 요구사항을 확인하고 관련 증빙 확보",
                "WARNING_REQUIRED": "표시·인증 관련 요구사항을 확인해 제출자료에 반영",
                "REGULATED": "적용 할랄 기준과 인증기관의 판단 범위를 확인",
                "REVIEW_REQUIRED": "공식 할랄 기준과 인증기관 요구사항을 수동 검토",
                "VERIFICATION_REQUIRED": "원료 기원·전체 구성·제조공정·가공보조제·인증자료를 공급사에 확인",
            }.get(status, "원료 기원·구성·공정·인증자료 확인"),
        )

    # 규제 판정 단계에서 성분별 구체 조치가 제공된 경우 우선 사용한다.
    # 한국어 기반 원본 보고서에서 영문 내부 문구가 그대로 섞이지 않도록
    # 한국어가 포함된 조치만 직접 채택하고, 그 외에는 상태별 표준 조치를 사용한다.
    specific_action = clean_text(detail.get("recommended_action"))
    if specific_action and re.search(r"[가-힣]", specific_action):
        return specific_action
    return recommended_action(status)


def required_evidence_for_detail(detail: dict, status: str, target: str = "") -> str:
    value = clean_text(detail.get("required_evidence"))
    if value:
        return value
    status = normalize_status(status)
    if target == "HALAL":
        return {
            "PASS": "-",
            "BANNED": "대체원료 규격서·기원서·할랄 인증자료",
            "RESTRICTED": "원료규격서·사용조건 자료·인증기관 요구자료",
            "WARNING_REQUIRED": "표시안·인증자료·고지자료",
            "REGULATED": "적용 기준·원료규격서·인증기관 회신",
            "REVIEW_REQUIRED": "공식 기준 원문·인증기관 검토자료",
            "VERIFICATION_REQUIRED": "원료규격서·기원서·전체 구성표·제조공정서·유효한 인증자료",
        }.get(status, "공급사 확인자료")
    return {
        "PASS": "-",
        "BANNED": "변경 처방·대체원료 규격서",
        "RESTRICTED": "배합농도·제품유형·사용조건·원료규격서",
        "WARNING_REQUIRED": "최종 라벨·경고문구·고지 또는 통지 자료",
        "REGULATED": "제품 분류·용도·표시광고·규제 적용범위 근거",
        "REVIEW_REQUIRED": "공식 규정 원문·시행일·사용조건 검토기록",
        "VERIFICATION_REQUIRED": "공식 INCI·CAS·공급사 원료규격서",
    }.get(status, "확인자료")


def result_action_rows(result_data: dict, target: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for detail in attention_details(result_data):
        status = normalize_status(
            detail.get("compliance_status") or detail.get("restriction_type")
        )
        status_label = (
            halal_status_label(detail)
            if target == "HALAL"
            else STATUS_LABELS.get(status, status)
        )
        rows.append([
            clean_text(detail.get("original_ingredient")),
            status_label,
            action_for_detail(detail, target),
            required_evidence_for_detail(detail, status, target),
            completion_criterion(status, target),
        ])
    return rows


def general_next_steps(target: str) -> List[str]:
    return [
        "현지 제품 등록·신고 또는 사전통지 절차 확인",
        "제품 분류·효능표현·표시광고 및 라벨 언어 요건 검토",
        "책임자·수입자·현지 대리인과 통관 제출자료 준비",
        "처방·제품유형·표시내용 또는 적용 규정이 변경된 경우 재분석",
    ]


def manual_type(detail: dict, target: str) -> str:
    status = normalize_status(
        detail.get("compliance_status")
        or detail.get("restriction_type")
    )
    if target == "HALAL":
        return clean_text(
            detail.get("verification_type"),
            {
                "REVIEW_REQUIRED": "공식 할랄 기준",
                "VERIFICATION_REQUIRED": "원료 기원·구성·인증자료",
                "REGULATED": "할랄 기준 적용범위",
                "BANNED": "금지 성분 확인",
                "RESTRICTED": "조건 확인",
            }.get(status, "추가 확인"),
        )
    return {
        "REVIEW_REQUIRED": "규제 원문·적용조건",
        "VERIFICATION_REQUIRED": "공식 명칭·CAS",
        "REGULATED": "규제 적용범위",
    }.get(status, "추가 확인")


def halal_status_label(detail: dict) -> str:
    status = normalize_status(
        detail.get("compliance_status")
        or detail.get("restriction_type")
    )
    if status == "PASS":
        return "할랄 위험 후보 미확인"
    if status == "BANNED":
        return "할랄 금지 성분 확인"
    if status == "RESTRICTED":
        return "조건 확인 필요"
    if status == "WARNING_REQUIRED":
        return "표시·인증 확인 필요"
    if status == "REGULATED":
        return "공식 기준 적용범위 확인"
    if status == "REVIEW_REQUIRED":
        return "공식 기준 수동 검토"
    if status == "VERIFICATION_REQUIRED":
        verification_type = clean_text(detail.get("verification_type"))
        if verification_type:
            return verification_type
        if clean_text(detail.get("match_source")) == "translation_verification":
            return "명칭·구성 확인 필요"
        return "원료 기원·구성·인증자료 확인 필요"
    return "추가 확인 필요"


def halal_market_relation(target: str) -> Dict[str, str]:
    target = clean_text(target).upper()
    if target == "SFDA":
        return {
            "direct_required": "수입자, 유통사 또는 인증기관의 요구에 따라 필요할 수 있습니다.",
            "required_when": "할랄 표시를 사용하거나 거래 상대방이 원료 기원 또는 인증자료를 요구하는 경우",
            "screening_effect": "달라지지 않습니다. 화장품 성분 규제 분석과 HALAL 검토는 별도로 판단합니다.",
            "relevance": "수입자, 유통사 또는 인증기관의 요구에 따라 필요할 수 있습니다.",
            "regulatory_effect": "달라지지 않습니다. 화장품 성분 규제 분석과 HALAL 검토는 별도로 판단합니다.",
        }
    if target == "ASEAN":
        return {
            "direct_required": "판매 국가와 유통채널에 따라 다릅니다.",
            "required_when": "할랄 인증이 요구되는 국가에 판매하거나 바이어·유통채널이 관련 자료를 요구하는 경우",
            "screening_effect": "달라지지 않습니다. 국가별 화장품 성분 규제와 HALAL 요건은 별도로 확인합니다.",
            "relevance": "판매 국가와 유통채널에 따라 다릅니다.",
            "regulatory_effect": "달라지지 않습니다. 국가별 화장품 성분 규제와 HALAL 요건은 별도로 확인합니다.",
        }
    if target == "EAC":
        return {
            "direct_required": "회원국과 거래 상대방의 요구에 따라 다릅니다.",
            "required_when": "수입자, 바이어 또는 유통채널이 할랄 표시나 인증자료를 요구하는 경우",
            "screening_effect": "달라지지 않습니다. EAEU 성분 규제 분석과 HALAL 검토는 별도로 판단합니다.",
            "relevance": "회원국과 거래 상대방의 요구에 따라 다릅니다.",
            "regulatory_effect": "달라지지 않습니다. EAEU 성분 규제 분석과 HALAL 검토는 별도로 판단합니다.",
        }
    return {
        "direct_required": "일반 화장품 판매에는 직접 필요하지 않습니다.",
        "required_when": "제품을 할랄 제품으로 판매하거나 바이어·유통채널이 HALAL 인증자료를 요구하는 경우",
        "screening_effect": "달라지지 않습니다. 이 시장의 성분 규제 분석과 HALAL 검토는 별도로 판단합니다.",
        "relevance": "일반 화장품 판매에는 직접 필요하지 않습니다.",
        "regulatory_effect": "달라지지 않습니다. 이 시장의 성분 규제 분석과 HALAL 검토는 별도로 판단합니다.",
    }


def attention_details(result_data: Optional[dict]) -> List[dict]:
    if not result_data:
        return []
    return [
        item
        for item in result_details(result_data)
        if normalize_status(
            item.get("compliance_status")
            or item.get("restriction_type")
        ) != "PASS"
    ]


def linked_halal_ingredients(halal_result: Optional[dict]) -> str:
    names = [
        clean_text(item.get("original_ingredient"))
        for item in attention_details(halal_result)
    ]
    return ", ".join(dedupe_preserve(names)) or "없음"


# ============================================================
# Excel 데이터 행 생성
# ============================================================
def market_result_rows(
    result_data: dict,
    source_file: str,
    target: str = "",
) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for index, detail in enumerate(result_details(result_data), start=1):
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                STATUS_LABELS.get(status, status),
                market_impact_for_status(status),
                normalized_reason(detail, target),
                applicable_condition(status),
                action_for_detail(detail, target),
                completion_criterion(status, target),
                match_source_label(detail.get("match_source")),
                source_file,
            ]
        )
    return rows

def halal_result_rows(
    result_data: dict,
    source_file: str,
) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for index, detail in enumerate(result_details(result_data), start=1):
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        is_pass = status == "PASS"
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                halal_status_label(detail),
                match_source_label(detail.get("match_source")),
                "-" if is_pass else clean_text(
                    detail.get("verification_type"),
                    manual_type(detail, "HALAL"),
                ),
                normalized_reason(detail, "HALAL"),
                required_evidence_for_detail(detail, status, "HALAL"),
                action_for_detail(detail, "HALAL"),
                completion_criterion(status, "HALAL"),
                source_file,
            ]
        )
    return rows

def summary_required_action(result_data: dict) -> str:
    status = dominant_status(result_data)
    return {
        "PASS": "즉시 필요한 성분 수정 없음",
        "BANNED": "금지 성분 제거 또는 대체 후 전체 성분 재분석",
        "RESTRICTED": "제한 성분의 배합농도, 제품 유형 및 사용조건 확인",
        "WARNING_REQUIRED": "필수 표시·경고·고지 또는 통지 요건 반영",
        "REGULATED": "제품 유형과 사용 목적에 따른 규정 적용 여부 확인",
        "REVIEW_REQUIRED": "공식 규정 원문과 적용조건 수동 확인",
        "VERIFICATION_REQUIRED": "공식 INCI, CAS 번호 또는 공급사 자료 확인",
    }.get(status, "추가 확인 필요")


def market_summary_rows(
    product_name: str,
    target: str,
    result_data: dict,
    halal_result: Optional[dict],
) -> List[List[Any]]:
    counts = status_counts(result_data)
    decision = market_decision(result_data)
    rows: List[List[Any]] = [
        ["제품명", product_name],
        ["대상 시장", MARKET_LABELS.get(target, target)],
        ["검사 성분 수", safe_int(result_data.get("total_checked"))],
        ["통과", counts["PASS"]],
        ["금지", counts["BANNED"]],
        ["제한", counts["RESTRICTED"]],
        ["경고·표시 필요", counts["WARNING_REQUIRED"]],
        ["규제 대상", counts["REGULATED"]],
        ["수동 검토", counts["REVIEW_REQUIRED"]],
        ["명칭 검증 필요", counts["VERIFICATION_REQUIRED"]],
        ["전체 수동확인 필요", counts["REVIEW_REQUIRED"] + counts["VERIFICATION_REQUIRED"] + counts["REGULATED"]],
        ["종합 판정", decision["overall"]],
        ["성분 규제 진행 상태", decision["possibility"]],
        ["최종결론", decision["conclusion"]],
        ["필요한 조치", summary_required_action(result_data)],
        [
            ("조치 완료 후 다음 단계" if attention_details(result_data) else "성분 규제 외 별도 절차"),
            (
                " / ".join(general_next_steps(target)[:2])
                if attention_details(result_data)
                else "성분 규제상 추가 조치 없음. 제품 등록·라벨·표시광고·통관은 별도 절차"
            ),
        ],
        ["검토 기준 DB", clean_text(result_data.get("database_file"))],
        ["DB 업데이트일", clean_text(result_data.get("database_last_updated"))],
        ["보고서 생성일", clean_text(result_data.get("report_generated_at"), datetime.now().strftime("%Y-%m-%d %H:%M:%S"))],
        ["사용 범위", "화장품 성분 규제 사전 스크리닝"],
    ]
    if halal_result is not None:
        relation = halal_market_relation(target)
        hdecision = halal_decision(halal_result)
        rows.extend(
            [
                ["HALAL 추가검토 여부", "적용"],
                ["이 시장과 HALAL 검토의 관련성", relation["relevance"]],
                ["HALAL 자료가 필요한 경우", relation["required_when"]],
                ["이 시장의 성분 규제 결과에 미치는 영향", relation["regulatory_effect"]],
                ["이번 HALAL 검토 결과", hdecision["overall"]],
                ["연결된 HALAL 확인 성분", linked_halal_ingredients(halal_result)],
                ["HALAL 필수 조치", " / ".join(
                    prioritized_action_summaries(halal_result, "HALAL")
                ) or "HALAL 추가검토상 필수 보완조치 없음"],
            ]
        )
    return rows

def halal_summary_rows(
    product_name: str,
    result_data: dict,
) -> List[List[Any]]:
    counts = status_counts(result_data)
    decision = halal_decision(result_data)
    actions = prioritized_action_summaries(result_data, "HALAL")
    return [
        ["제품명", product_name],
        ["검토 구분", "할랄 추가 성분 스크리닝"],
        ["검사 성분 수", safe_int(result_data.get("total_checked"))],
        ["위험 후보 미확인", counts["PASS"]],
        ["할랄 금지", counts["BANNED"]],
        ["조건 확인 필요", counts["RESTRICTED"]],
        ["공식 기준 적용범위 확인", counts["REGULATED"]],
        ["공식 기준 수동검토", counts["REVIEW_REQUIRED"]],
        ["원료 기원·구성·인증 확인 필요", counts["VERIFICATION_REQUIRED"]],
        ["전체 확인 필요", sum(counts[key] for key in counts if key != "PASS")],
        ["할랄 종합판정", decision["overall"]],
        ["할랄 인증 추진 가능성", decision["possibility"]],
        ["최종결론", decision["conclusion"]],
        ["분석결과에 따른 필수 조치", " / ".join(actions) if actions else "현재 성분명 기준 별도 HALAL 보완조치 없음"],
        ["검토 기준 DB", clean_text(result_data.get("database_file"))],
        ["검토대기 DB", clean_text(result_data.get("review_database_file"))],
        ["DB 업데이트일", clean_text(result_data.get("database_last_updated"))],
        ["보고서 생성일", clean_text(result_data.get("report_generated_at"), datetime.now().strftime("%Y-%m-%d %H:%M:%S"))],
        ["사용 범위", "할랄 인증서 또는 최종 인증판정이 아닌 사전 위험 선별"],
    ]

def prioritized_action_summaries(
    result_data: dict,
    target: str,
) -> List[str]:
    details = attention_details(result_data)
    order = {status: index for index, status in enumerate(STATUS_ORDER)}
    details = sorted(
        details,
        key=lambda item: order.get(
            normalize_status(
                item.get("compliance_status")
                or item.get("restriction_type")
            ),
            len(order),
        ),
    )

    actions: List[str] = []
    for detail in details:
        ingredient = clean_text(
            detail.get("original_ingredient"),
            clean_text(detail.get("inci_name"), "해당 성분"),
        )
        action = action_for_detail(detail, target)
        summarized = f"{ingredient}: {action}" if ingredient else action
        if summarized and summarized not in actions:
            actions.append(summarized)
    return actions


def priority_actions(result_data: dict, target: str) -> str:
    actions = prioritized_action_summaries(result_data, target)
    if not actions:
        return "성분 규제상 필수 보완조치 없음"
    return " → ".join(actions[:6])

def manual_rows(
    result_data: dict,
    target: str,
    halal_result: Optional[dict] = None,
) -> List[List[Any]]:
    """국가 규제 수동확인 시트에는 국가 규제 항목만 포함한다.

    HALAL 결과는 별도의 HALAL 수동확인 시트에서 관리하여 국가 규제
    통과 결과와 혼동되지 않도록 한다. halal_result 인자는 하위 호환용이다.
    """
    rows: List[List[Any]] = []
    index = 1
    for detail in result_details(result_data):
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        if status not in MANUAL_STATUSES:
            continue
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                manual_type(detail, target),
                STATUS_LABELS.get(status, status),
                normalized_reason(detail, target),
                applicable_condition(status),
                required_evidence_for_detail(detail, status, target),
                "",
                "",
                "",
                "",
            ]
        )
        index += 1
    return rows

def halal_manual_rows(result_data: dict) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for index, detail in enumerate(attention_details(result_data), start=1):
        status = normalize_status(
            detail.get("compliance_status") or detail.get("restriction_type")
        )
        rows.append(
            [
                index,
                clean_text(detail.get("original_ingredient")),
                clean_text(detail.get("inci_name")),
                clean_text(detail.get("cas_number"), "N/A"),
                manual_type(detail, "HALAL"),
                halal_status_label(detail),
                normalized_reason(detail, "HALAL"),
                action_for_detail(detail, "HALAL"),
                required_evidence_for_detail(detail, status, "HALAL"),
                "",
                "",
                "",
                "",
            ]
        )
    return rows

# ============================================================
# Excel 생성
# ============================================================
def _xlsx_title(ws, title: str, last_col: int) -> None:
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=last_col)
    cell = ws.cell(1, 1, title)
    cell.fill = PatternFill("solid", fgColor=LIGHT_BLUE)
    cell.font = Font(bold=True, size=15, color=NAVY)
    cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 30


def _xlsx_header(ws, row: int, headers: Sequence[str]) -> None:
    thin = Side(style="thin", color=BORDER_COLOR)
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row, col, header)
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(bold=True, color=WHITE, size=10)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    ws.row_dimensions[row].height = 28


def _xlsx_body(ws, start_row: int, end_row: int, max_col: int) -> None:
    thin = Side(style="thin", color=BORDER_COLOR)
    for row in ws.iter_rows(min_row=start_row, max_row=end_row, min_col=1, max_col=max_col):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.font = Font(size=9, color="333333")
    for row_index in range(start_row, end_row + 1):
        ws.row_dimensions[row_index].height = 42


def _set_widths(ws, widths: Sequence[float]) -> None:
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width


def _write_result_sheet(ws, title: str, headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> None:
    _xlsx_title(ws, title, len(headers))
    _xlsx_header(ws, 3, headers)
    for r_idx, row in enumerate(rows, start=4):
        for c_idx, value in enumerate(row, start=1):
            ws.cell(r_idx, c_idx, value)
    end_row = max(4, 3 + len(rows))
    _xlsx_body(ws, 4, end_row, len(headers))
    ws.freeze_panes = "A4"
    ws.auto_filter.ref = f"A3:{get_column_letter(len(headers))}{end_row}"
    _set_widths(ws, [7, 18, 26, 14, 20, 22, 34, 28, 34, 30, 22, 24])

    for row_index in range(4, end_row + 1):
        status = clean_text(ws.cell(row_index, 5).value)
        fill = None
        if "금지" in status:
            fill = RED
        elif any(term in status for term in ("제한", "확인", "검토", "규제")):
            fill = YELLOW
        elif "통과" in status or "미확인" in status:
            fill = GREEN
        if fill:
            ws.cell(row_index, 5).fill = PatternFill("solid", fgColor=fill)
            ws.cell(row_index, 5).font = Font(bold=True, size=9)


def _write_summary_sheet(ws, title: str, rows: Sequence[Sequence[Any]]) -> None:
    _xlsx_title(ws, title, 2)
    _xlsx_header(ws, 3, ["항목", "내용"])
    for r_idx, row in enumerate(rows, start=4):
        ws.cell(r_idx, 1, row[0])
        ws.cell(r_idx, 2, row[1])
        ws.cell(r_idx, 1).fill = PatternFill("solid", fgColor=GRAY)
        ws.cell(r_idx, 1).font = Font(bold=True, size=9)
    end_row = max(4, 3 + len(rows))
    _xlsx_body(ws, 4, end_row, 2)
    _set_widths(ws, [28, 80])
    for r_idx in range(4, end_row + 1):
        ws.row_dimensions[r_idx].height = 30


def _write_manual_sheet(ws, title: str, rows: Sequence[Sequence[Any]]) -> None:
    headers = [
        "번호",
        "성분명",
        "INCI 명칭",
        "CAS 번호",
        "수동확인 유형",
        "현재 판정",
        "확인 사유",
        "확인할 사항",
        "필요한 자료",
        "확인 결과",
        "최종 판정",
        "검토자",
        "검토일",
    ]
    _xlsx_title(ws, title, len(headers))
    _xlsx_header(ws, 3, headers)
    if rows:
        for r_idx, row in enumerate(rows, start=4):
            for c_idx, value in enumerate(row, start=1):
                ws.cell(r_idx, c_idx, value)
    else:
        rows = [[1, "해당 항목 없음", "", "", "", "", "", "", "", "", "", "", ""]]
        for c_idx, value in enumerate(rows[0], start=1):
            ws.cell(4, c_idx, value)
    end_row = 3 + len(rows)
    _xlsx_body(ws, 4, end_row, len(headers))
    ws.freeze_panes = "A4"
    ws.auto_filter.ref = f"A3:M{end_row}"
    _set_widths(ws, [7, 18, 25, 14, 24, 20, 32, 32, 28, 22, 20, 14, 14])

    status_validation = DataValidation(
        type="list",
        formula1='"확인 중,적합,조건부 적합,부적합,재검토 필요"',
        allow_blank=True,
    )
    ws.add_data_validation(status_validation)
    status_validation.add(f"K4:K{max(end_row, 100)}")


def create_product_excel_bytes(
    product_name: str,
    source_file: str,
    market_results_map: Dict[str, dict],
    selected_markets: Sequence[str],
    halal_result: Optional[dict] = None,
    language_code: str = "ko",
    translator: Optional[Translator] = None,
) -> bytes:
    workbook = Workbook()
    workbook.remove(workbook.active)
    sheet_number = 1

    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue
        label = MARKET_LABELS.get(target, target)

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_규제분석결과")
        _write_result_sheet(
            ws,
            f"{label} 규제 분석 결과",
            [
                "번호",
                "입력 성분명",
                "INCI 명칭",
                "CAS 번호",
                "규제 판정",
                "시장 진입 영향",
                "규제 근거",
                "적용 조건",
                "필요한 조치",
                "조치 완료 기준",
                "판정 출처",
                "입력 파일",
            ],
            market_result_rows(result_data, source_file, target),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_종합요약")
        _write_summary_sheet(
            ws,
            f"{label} 종합요약",
            market_summary_rows(
                product_name,
                target,
                result_data,
                halal_result,
            ),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_{target}_수동확인")
        _write_manual_sheet(
            ws,
            f"{label} 수동확인 체크리스트",
            manual_rows(result_data, target),
        )
        sheet_number += 1

    if halal_result is not None:
        ws = workbook.create_sheet(f"{sheet_number:02d}_HALAL_성분검토결과")
        _write_result_sheet(
            ws,
            "HALAL 추가 성분 검토 결과",
            [
                "번호",
                "입력 성분명",
                "INCI 명칭",
                "CAS 번호",
                "할랄 검토 상태",
                "검출 방식",
                "우려 유형",
                "확인 사유",
                "필요한 증빙",
                "공급사 확인사항",
                "조치 완료 기준",
                "입력 파일",
            ],
            halal_result_rows(halal_result, source_file),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_HALAL_종합요약")
        _write_summary_sheet(
            ws,
            "HALAL 종합요약",
            halal_summary_rows(product_name, halal_result),
        )
        sheet_number += 1

        ws = workbook.create_sheet(f"{sheet_number:02d}_HALAL_수동확인")
        _write_manual_sheet(
            ws,
            "HALAL 수동확인 체크리스트",
            halal_manual_rows(halal_result),
        )

    output = io.BytesIO()
    workbook.save(output)
    data = output.getvalue()
    protected_terms = _result_protected_terms(
        product_name,
        source_file,
        market_results_map,
        halal_result,
    )
    return _localize_excel_bytes(
        data,
        language_code,
        translator,
        protected_terms,
    )


# ============================================================
# DOCX 생성 도우미
# ============================================================
def _doc_shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _doc_cell_margins(
    cell: Any,
    top: int = 70,
    start: int = 90,
    bottom: int = 70,
    end: int = 90,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _doc_set_cell(
    cell: Any,
    text: Any,
    bold: bool = False,
    size: float = 9,
    color: RGBColor = DOC_TEXT,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(clean_text(text))
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _doc_cell_margins(cell)


def _doc_borders(table: Any, color: str = BORDER_COLOR, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _doc_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B9C2C8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _chapter_header(
    document: Document,
    number: int,
    title: str,
    subtitle: str,
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(0.85)
    table.columns[1].width = Inches(5.8)
    _doc_borders(table, WHITE, 0)
    _doc_shade(table.cell(0, 0), "3C3C3C")
    _doc_set_cell(
        table.cell(0, 0),
        f"제{number}장",
        True,
        9.2,
        RGBColor(255, 255, 255),
        WD_ALIGN_PARAGRAPH.CENTER,
    )
    _doc_set_cell(table.cell(0, 1), title, True, 17.2)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(subtitle)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DOC_MUTED
    _doc_rule(document)


def _section_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11.2)
    run.font.color.rgb = DOC_NAVY


def _bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.08)
    paragraph.paragraph_format.space_after = Pt(1.5)
    run = paragraph.add_run(text)
    run.font.size = Pt(9.6)


def _decision_box(document: Document, label: str, text: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.15)
    _doc_borders(table, "D8BE63", 6)
    _doc_shade(table.cell(0, 0), YELLOW)
    _doc_shade(table.cell(0, 1), "FFF9E8")
    _doc_set_cell(
        table.cell(0, 0),
        label,
        True,
        9.4,
        DOC_TEXT,
        WD_ALIGN_PARAGRAPH.CENTER,
    )
    _doc_set_cell(
        table.cell(0, 1),
        text,
        True,
        10.2,
        DOC_GOLD,
        WD_ALIGN_PARAGRAPH.CENTER,
    )


def _add_simple_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    header_fill: str = NAVY,
    font_size: float = 8.7,
    highlight_col: Optional[int] = None,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _doc_borders(table)
    for index, header in enumerate(headers):
        _doc_shade(table.cell(0, index), header_fill)
        _doc_set_cell(
            table.cell(0, index),
            header,
            True,
            font_size,
            RGBColor(255, 255, 255),
            WD_ALIGN_PARAGRAPH.CENTER,
        )
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            if row_index % 2:
                _doc_shade(cells[col_index], PALE)
            if highlight_col is not None and col_index == highlight_col:
                text = clean_text(value)
                if any(term in text for term in ("금지", "제한", "확인", "검토", "규제")):
                    _doc_shade(cells[col_index], YELLOW)
            _doc_set_cell(
                cells[col_index],
                value,
                col_index == highlight_col,
                font_size,
                DOC_TEXT,
                WD_ALIGN_PARAGRAPH.CENTER if col_index == highlight_col else WD_ALIGN_PARAGRAPH.LEFT,
            )


def _market_attention_rows(result_data: dict) -> List[List[str]]:
    details = result_details(result_data)
    chosen = details if len(details) <= 18 else attention_details(result_data)
    rows: List[List[str]] = []
    for detail in chosen:
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        rows.append(
            [
                clean_text(detail.get("original_ingredient")),
                STATUS_LABELS.get(status, status),
                normalized_reason(detail),
            ]
        )
    if len(details) > 18 and not rows:
        rows.append(["전체 성분", "통과", "상세 성분별 결과는 Excel 규제분석결과 시트를 참조하십시오."])
    return rows


def _manual_report_rows(result_data: dict, target: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for detail in result_details(result_data):
        status = normalize_status(
            detail.get("compliance_status")
            or detail.get("restriction_type")
        )
        if status not in MANUAL_STATUSES:
            continue
        rows.append(
            [
                clean_text(detail.get("original_ingredient")),
                STATUS_LABELS.get(status, status),
                normalized_reason(detail, target),
                action_for_detail(detail, target),
            ]
        )
    return rows


def _halal_link_rows(halal_result: dict, target: str) -> List[List[str]]:
    relation = halal_market_relation(target)
    rows: List[List[str]] = []
    for detail in attention_details(halal_result)[:12]:
        rows.append(
            [
                clean_text(detail.get("original_ingredient")),
                halal_status_label(detail),
                relation["impact"],
                clean_text(
                    detail.get("required_evidence"),
                    "공급사 원료 기원·구성·공정·인증자료",
                ),
            ]
        )
    return rows


def _common_glossary_rows() -> List[Tuple[str, str]]:
    return [
        ("통과", "현재 규제 DB에서 일치 항목이 확인되지 않은 상태"),
        ("금지", "현재 처방 기준 해당 시장 수입·유통·판매 진행이 불가능한 상태"),
        ("제한", "농도·제품유형·사용대상 등 조건 충족이 필요한 상태"),
        ("경고·표시 필요", "라벨·고지·경고문구·통지 요건 검토가 필요한 상태"),
        ("규제 대상", "관련 규제는 있으나 실제 적용범위를 추가 확인해야 하는 상태"),
        ("수동 검토", "공식 원문이나 공급사 자료를 사람이 확인해야 하는 상태"),
        ("명칭 검증 필요", "공식 INCI가 확정되지 않아 판정을 완료할 수 없는 상태"),
        ("확인 필요", "추가 자료 없이는 최종 결론을 내릴 수 없는 상태"),
        ("부적합", "현재 처방 또는 표시 상태로 규제요건을 충족하지 못하는 상태"),
        ("조건 충족 후 진행 가능", "필수 제한·표시·자료 요건을 충족하면 진행 가능한 상태"),
        ("CAS 번호", "Chemical Abstracts Service가 화학물질에 부여한 고유 식별번호"),
    ]


def _halal_glossary_rows() -> List[Tuple[str, str]]:
    return [
        ("할랄", "이슬람 기준에 따라 허용되는 원료·공정·제품"),
        ("하람", "이슬람 기준상 허용되지 않는 원료·공정·제품"),
        ("원료 기원", "식물성·합성·동물성 등 원료가 유래한 근원"),
        ("동물 유래", "동물 조직·지방·분비물 등에서 유래한 원료"),
        ("가공보조제", "최종 성분표에 남지 않더라도 제조 중 사용되는 물질"),
        ("교차오염", "비할랄 원료·설비·보관·운송 과정에서 혼입될 가능성"),
        ("할랄 인증서", "인증기관이 원료 또는 제품의 적합성을 확인해 발급한 문서"),
        ("위험 후보 미확인", "현재 분석에서 우려 후보가 발견되지 않았으나 인증 완료는 아닌 상태"),
        ("원료 기원 확인 필요", "공급사 자료로 원료 유래와 제조공정을 확인해야 하는 상태"),
        ("CAS 번호", "Chemical Abstracts Service가 화학물질에 부여한 고유 식별번호"),
    ]


def _glossary_table(document: Document, rows: Sequence[Tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _doc_borders(table)
    for term, definition in rows:
        cells = table.add_row().cells
        _doc_shade(cells[0], GRAY)
        _doc_set_cell(cells[0], term, True, 7.9)
        _doc_set_cell(cells[1], definition, False, 7.9)
        _doc_cell_margins(cells[0], top=35, bottom=35, start=70, end=70)
        _doc_cell_margins(cells[1], top=35, bottom=35, start=70, end=70)


def create_product_report_bytes(
    product_name: str,
    source_file: str,
    market_results_map: Dict[str, dict],
    selected_markets: Sequence[str],
    halal_result: Optional[dict] = None,
    language_code: str = "ko",
    translator: Optional[Translator] = None,
) -> bytes:
    """Create a market-by-market screening report.

    English and Korean retain the complete report. Other languages use the same
    core findings but omit repeated glossary and interpretation sections so the
    localized report remains concise and less vulnerable to literal translation.
    """
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = DOC_TEXT
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(3)

    compact_language = language_code not in {"ko", "en"}
    market_labels = [
        MARKET_LABELS.get(target, target)
        for target in selected_markets
        if target in market_results_map
    ]
    scope = " · ".join(market_labels)
    if halal_result is not None:
        scope += " + HALAL 추가검토"

    _chapter_header(
        document,
        1,
        "제품 및 분석 개요",
        f"{product_name} · {scope}",
    )

    info = document.add_table(rows=5, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.LEFT
    _doc_borders(info)
    overview_rows = [
        ("제품명", product_name),
        ("입력 파일", source_file),
        ("분석 시장", " · ".join(market_labels) or "없음"),
        ("HALAL 추가검토", "선택함" if halal_result is not None else "선택하지 않음"),
        (
            "분석 성분 수",
            max(
                [safe_int(data.get("total_checked")) for data in market_results_map.values()]
                + ([safe_int(halal_result.get("total_checked"))] if halal_result else [0])
            ),
        ),
    ]
    for row_index, (label, value) in enumerate(overview_rows):
        _doc_shade(info.cell(row_index, 0), GRAY)
        _doc_set_cell(info.cell(row_index, 0), label, True, 9.5)
        _doc_set_cell(info.cell(row_index, 1), value, False, 9.5)

    first_result = next(iter(market_results_map.values()), halal_result or {})
    ingredients = [
        [
            clean_text(item.get("original_ingredient")),
            clean_text(item.get("inci_name")),
            clean_text(item.get("cas_number"), "N/A"),
        ]
        for item in result_details(first_result)
    ]
    _section_title(document, "1. 제품 성분")
    if len(ingredients) <= 24:
        _add_simple_table(
            document,
            ["입력 성분명", "INCI 명칭", "CAS 번호"],
            ingredients,
            NAVY,
            8.8,
        )
    else:
        paragraph = document.add_paragraph()
        paragraph.add_run(
            f"전체 {len(ingredients)}개 성분을 분석했습니다. 전체 성분 목록과 성분별 결과는 함께 제공되는 Excel 파일에서 확인할 수 있습니다."
        )

    _section_title(document, "2. 시장별 결과 요약")
    summary_rows: List[List[str]] = []
    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue
        decision = market_decision(result_data)
        summary_rows.append(
            [
                "화장품 성분 규제",
                target,
                decision["overall"],
                decision["possibility"],
            ]
        )
    if halal_result is not None:
        decision = halal_decision(halal_result)
        summary_rows.append(
            [
                "HALAL 추가검토",
                "HALAL",
                decision["overall"],
                decision["possibility"],
            ]
        )
    _add_simple_table(
        document,
        ["검토 구분", "시장", "분석 결과", "의미"],
        summary_rows,
        BLUE,
        8.8,
        2,
    )

    _section_title(document, "3. 보고서 사용 범위")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        "이 보고서는 성분 규제와 HALAL 위험 후보를 사전에 확인하기 위한 자료입니다. 제품 등록, 통관 승인, 판매 허가 또는 HALAL 인증을 보장하지 않으므로 실제 진행 전 최신 공식 기준과 제품 자료를 별도로 확인해야 합니다."
    )
    run.font.size = Pt(9.2)
    run.font.color.rgb = DOC_MUTED
    if halal_result is not None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            "HALAL 검토는 국가별 화장품 성분 규제 분석과 별도로 판단합니다."
        )
        run.font.size = Pt(9.2)
        run.font.color.rgb = DOC_MUTED

    chapter_number = 2
    for target in selected_markets:
        result_data = market_results_map.get(target)
        if not result_data:
            continue

        page_anchor = document.add_paragraph()
        page_anchor.paragraph_format.page_break_before = True
        page_anchor.paragraph_format.space_after = Pt(0)

        market_label = MARKET_LABELS.get(target, target)
        _chapter_header(
            document,
            chapter_number,
            f"{market_label} 규제 분석 결과",
            "이 장에서는 해당 시장의 성분 분석 결과와 필요한 조치를 설명합니다.",
        )
        chapter_number += 1
        decision = market_decision(result_data)
        _decision_box(document, "이번 분석 결과", decision["possibility"])

        _section_title(document, "1. 성분별 분석 결과")
        rows = _market_attention_rows(result_data)
        if rows:
            _add_simple_table(
                document,
                ["성분", "판정", "설명"],
                rows,
                NAVY,
                8.6,
                1,
            )
        else:
            _bullet(document, "추가 확인이 필요한 성분은 없습니다. 전체 결과는 Excel 파일에서 확인할 수 있습니다.")

        _section_title(document, "2. 필요한 조치")
        action_rows = result_action_rows(result_data, target)
        if action_rows:
            _add_simple_table(
                document,
                ["성분", "판정", "해야 할 일", "필요한 자료", "완료 기준"],
                action_rows,
                BLUE,
                8.1,
                1,
            )
        else:
            _bullet(document, "현재 분석 결과, 성분 규제와 관련하여 즉시 수정해야 할 항목은 없습니다.")
            _bullet(document, "제품 등록, 라벨, 광고 및 통관 요건은 별도로 확인해야 합니다.")

        section_index = 3
        if halal_result is not None:
            _section_title(
                document,
                f"{section_index}. HALAL 검토가 이 시장에서 필요한 경우",
            )
            section_index += 1
            relation = halal_market_relation(target)
            relation_table = [
                ["일반 화장품 판매에 직접 필요한가?", relation["direct_required"]],
                ["HALAL 자료가 필요할 수 있는 경우", relation["required_when"]],
                ["이번 성분 규제 결과가 달라지는가?", relation["screening_effect"]],
                ["이번 HALAL 검토 결과", halal_decision(halal_result)["overall"]],
            ]
            _add_simple_table(
                document,
                ["질문", "답변"],
                relation_table,
                BLUE,
                8.7,
            )
            link_rows = _halal_link_rows(halal_result, target)
            if link_rows:
                _add_simple_table(
                    document,
                    ["HALAL 확인 성분", "검토 상태", "시장에 미치는 영향", "필요한 자료"],
                    link_rows,
                    NAVY,
                    8.2,
                    1,
                )
            # When no HALAL verification item is linked, the concise answer table
            # above is sufficient; do not add a repetitive paragraph.

        manual = _manual_report_rows(result_data, target)
        if manual:
            _section_title(document, f"{section_index}. 추가로 확인할 항목")
            section_index += 1
            _add_simple_table(
                document,
                ["성분", "현재 판정", "확인 이유", "필요한 조치"],
                manual,
                BLUE,
                8.4,
                1,
            )

        if action_rows:
            _section_title(document, f"{section_index}. 조치 완료 후 다음 단계")
            _bullet(document, "위 조치를 완료하고, 완료를 입증할 자료를 보관하십시오.")
            for step in general_next_steps(target)[:2]:
                _bullet(document, step)
            if halal_result is not None and attention_details(halal_result):
                _bullet(document, "HALAL 인증 또는 관련 유통채널을 추진하는 경우 HALAL 장의 확인사항도 완료하십시오.")
        else:
            _section_title(document, f"{section_index}. 성분 규제 외 별도 절차")
            _bullet(document, "성분 규제상 추가 조치나 재분석은 필요하지 않습니다.")
            _bullet(document, "제품 등록·라벨·표시광고·통관은 성분 규제와 별도 절차이므로 필요한 경우 확인하십시오.")

    if halal_result is not None:
        page_anchor = document.add_paragraph()
        page_anchor.paragraph_format.page_break_before = True
        page_anchor.paragraph_format.space_after = Pt(0)
        _chapter_header(
            document,
            chapter_number,
            "HALAL 추가 성분 검토",
            "국가별 화장품 성분 규제와 별도로 원료 기원, 제조공정 및 인증자료를 확인합니다.",
        )
        decision = halal_decision(halal_result)
        _decision_box(document, "이번 HALAL 검토 결과", decision["possibility"])

        _section_title(document, "1. 성분별 검토 결과")
        rows: List[List[str]] = []
        for detail in result_details(halal_result):
            rows.append(
                [
                    clean_text(detail.get("original_ingredient")),
                    halal_status_label(detail),
                    "-" if normalize_status(detail.get("compliance_status") or detail.get("restriction_type")) == "PASS" else clean_text(detail.get("verification_type"), manual_type(detail, "HALAL")),
                    required_evidence_for_detail(
                        detail,
                        normalize_status(detail.get("compliance_status") or detail.get("restriction_type")),
                        "HALAL",
                    ),
                ]
            )
        if len(rows) > 22:
            attention_names = {
                clean_text(item.get("original_ingredient"))
                for item in attention_details(halal_result)
            }
            rows = [row for row in rows if row[0] in attention_names]
        _add_simple_table(
            document,
            ["성분", "검토 결과", "확인 이유", "필요한 자료"],
            rows,
            NAVY,
            8.4,
            1,
        )

        _section_title(document, "2. 필요한 조치")
        halal_action_rows = result_action_rows(halal_result, "HALAL")
        if halal_action_rows:
            _add_simple_table(
                document,
                ["성분", "판정", "해야 할 일", "필요한 자료", "완료 기준"],
                halal_action_rows,
                BLUE,
                8.1,
                1,
            )
        else:
            _bullet(document, "현재 성분명 기준으로 별도 확인이 필요한 HALAL 위험 후보는 발견되지 않았습니다.")
            _bullet(document, "최종 인증 여부는 공급사 자료, 제조공정 및 인증기관 심사를 통해 결정됩니다.")

        section_index = 3
        manual = [
            [row[1], row[4], row[7], row[8]]
            for row in halal_manual_rows(halal_result)
        ]
        if manual:
            _section_title(document, f"{section_index}. 추가로 확인할 항목")
            section_index += 1
            _add_simple_table(
                document,
                ["성분", "확인 유형", "확인할 사항", "필요한 자료"],
                manual,
                BLUE,
                8.4,
                1,
            )

        _section_title(document, f"{section_index}. 다음 단계")
        if halal_action_rows:
            _bullet(document, "필요한 자료를 확보하고 확인 결과를 기록하십시오.")
        halal_steps = [
            "적용할 HALAL 인증기관과 인증 범위를 정하십시오.",
            "원료 기원, 완제품 제조공정, 가공보조제 및 교차오염 관리자료를 준비하십시오.",
            "인증기관이 요구하는 자료를 제출하고 최종 심사를 진행하십시오.",
        ]
        halal_steps = halal_steps[:2]
        for step in halal_steps:
            _bullet(document, step)

        _section_title(document, f"{section_index + 1}. 중요 안내")
        _bullet(document, "이 검토 결과는 HALAL 인증 완료를 의미하지 않습니다.")
        _bullet(document, "최종 판단은 공급사 증빙, 제조공정, 교차오염 관리 및 인증기관 심사를 통해 이루어집니다.")

    for document_section in document.sections:
        footer = document_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(
            "K-Beauty Global Compliance | 규제 스크리닝 및 HALAL 추가검토"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = DOC_MUTED

    output = io.BytesIO()
    document.save(output)
    data = output.getvalue()
    protected_terms = _result_protected_terms(
        product_name,
        source_file,
        market_results_map,
        halal_result,
    )
    return _localize_docx_bytes(
        data,
        language_code,
        translator,
        protected_terms,
    )


# ============================================================
# 다국어 출력 후처리
# ============================================================
VALIDATION_OPTIONS: Dict[str, str] = {
    "ko": "확인 중,적합,조건부 적합,부적합,재검토 필요",
    "en": "Under review,Compliant,Conditionally compliant,Non-compliant,Re-review required",
    "zh-CN": "核验中,符合,有条件符合,不符合,需重新核验",
    "ar": "قيد المراجعة,مطابق,مطابق بشروط,غير مطابق,تستلزم إعادة المراجعة",
    "fr": "En cours,Conforme,Conforme sous conditions,Non conforme,À réexaminer",
    "id": "Dalam pemeriksaan,Sesuai,Sesuai bersyarat,Tidak sesuai,Perlu ditinjau ulang",
    "ms": "Dalam semakan,Patuh,Patuh bersyarat,Tidak patuh,Perlu semakan semula",
    "sw": "Inaangaliwa,Inakubalika,Inakubalika kwa masharti,Haikubaliki,Ukaguzi upya unahitajika",
}

FILE_LABELS: Dict[str, Dict[str, str]] = {
    "ko": {"excel": "다중시장_규제분석결과", "report": "다중시장_규제스크리닝_보고서", "zip": "다중제품_다중시장_규제분석결과", "failures": "분석실패목록"},
    "en": {"excel": "Multi_Market_Regulatory_Screening_Results", "report": "Multi_Market_Regulatory_Screening_Report", "zip": "Multi_Product_Regulatory_Screening_Results", "failures": "Analysis_Failures"},
    "zh-CN": {"excel": "多市场_法规筛查结果", "report": "多市场_法规筛查报告", "zip": "多产品_法规筛查结果", "failures": "分析失败清单"},
    "ar": {"excel": "نتائج_الفحص_التنظيمي_لأسواق_متعددة", "report": "تقرير_الفحص_التنظيمي_لأسواق_متعددة", "zip": "نتائج_الفحص_التنظيمي_لمنتجات_متعددة", "failures": "قائمة_إخفاقات_التحليل"},
    "fr": {"excel": "Resultats_Filtrage_Reglementaire_Multimarches", "report": "Rapport_Filtrage_Reglementaire_Multimarches", "zip": "Resultats_Filtrage_Reglementaire_Multiproduits", "failures": "Echecs_Analyse"},
    "id": {"excel": "Hasil_Penyaringan_Regulasi_Multi_Pasar", "report": "Laporan_Penyaringan_Regulasi_Multi_Pasar", "zip": "Hasil_Penyaringan_Regulasi_Multi_Produk", "failures": "Daftar_Kegagalan_Analisis"},
    "ms": {"excel": "Keputusan_Saringan_Peraturan_Pelbagai_Pasaran", "report": "Laporan_Saringan_Peraturan_Pelbagai_Pasaran", "zip": "Keputusan_Saringan_Peraturan_Pelbagai_Produk", "failures": "Senarai_Kegagalan_Analisis"},
    "sw": {"excel": "Matokeo_Uchunguzi_Kanuni_Masoko_Mengi", "report": "Ripoti_Uchunguzi_Kanuni_Masoko_Mengi", "zip": "Matokeo_Uchunguzi_Kanuni_Bidhaa_Nyingi", "failures": "Orodha_Hitilafu_Uchambuzi"},
}


def output_file_name(product_name: str, kind: str, language_code: str, extension: str) -> str:
    lang = language_code if language_code in OUTPUT_LANGUAGES else "ko"
    label = FILE_LABELS[lang][kind]
    suffix = LANGUAGE_FILE_SUFFIXES[lang]
    return f"{safe_filename(product_name)}_{label}_{suffix}.{extension}"


def output_zip_name(language_code: str) -> str:
    lang = language_code if language_code in OUTPUT_LANGUAGES else "ko"
    return f"{FILE_LABELS[lang]['zip']}_{LANGUAGE_FILE_SUFFIXES[lang]}.zip"


def _result_protected_terms(
    product_name: str,
    source_file: str,
    market_results_map: Dict[str, dict],
    halal_result: Optional[dict],
) -> List[str]:
    values: Set[str] = {
        clean_text(product_name),
        clean_text(source_file),
        "INCI",
        "CAS",
        "HALAL",
        "K-Beauty",
        "Global Compliance",
        "US",
        "EU",
        "UK",
        "CN",
        "ASEAN",
        "SFDA",
        "EAC",
        "EAEU",
        "FDA",
        "NMPA",
    }
    for result_data in list(market_results_map.values()) + ([halal_result] if halal_result else []):
        if not result_data:
            continue
        for key in ("database_file", "review_database_file", "database_last_updated", "report_generated_at"):
            values.add(clean_text(result_data.get(key)))
        for detail in result_details(result_data):
            for key in ("original_ingredient", "inci_name", "cas_number"):
                values.add(clean_text(detail.get(key)))
            for value in detail.values():
                text = clean_text(value)
                if text.startswith("http://") or text.startswith("https://"):
                    values.add(text)
    return sorted((value for value in values if value), key=len, reverse=True)


def _protect_strings(texts: Sequence[str], protected_terms: Sequence[str]) -> Tuple[List[str], Dict[str, str]]:
    token_to_term: Dict[str, str] = {}
    term_to_token: Dict[str, str] = {}
    for index, term in enumerate(protected_terms):
        token = f"[[P{index:04d}]]"
        term_to_token[term] = token
        token_to_term[token] = term
    protected_texts: List[str] = []
    for original in texts:
        value = original
        for term in protected_terms:
            if term and term in value:
                value = value.replace(term, term_to_token[term])
        protected_texts.append(value)
    return protected_texts, token_to_term


def _restore_tokens(text: str, token_to_term: Dict[str, str]) -> str:
    value = text
    for token, term in token_to_term.items():
        value = value.replace(token, term)
    return value


def _placeholder_tokens(text: str) -> List[str]:
    return re.findall(r"\[\[P\d{4}\]\]", clean_text(text))


def _visible_translation_text(text: str) -> str:
    value = re.sub(r"\[\[P\d{4}\]\]", " ", clean_text(text))
    return " ".join(value.split())


def _sentence_like_english(text: str) -> bool:
    value = _visible_translation_text(text)
    words = re.findall(r"[A-Za-z]{2,}", value)
    return len(words) >= 4


ENGLISH_MARKET_NAME_REPLACEMENTS: Tuple[Tuple[str, str], ...] = (
    ("미국", "United States"),
    ("유럽연합", "European Union"),
    ("영국", "United Kingdom"),
    ("중국", "China"),
    ("아세안", "ASEAN"),
    ("사우디아라비아", "Saudi Arabia"),
    ("유라시아경제연합", "Eurasian Economic Union"),
)


def _canonical_english_composite(text: str) -> Optional[str]:
    """Translate composite strings assembled from known fixed phrases."""
    value = clean_text(text)
    if not value:
        return None
    working = value
    for korean_name, english_name in ENGLISH_MARKET_NAME_REPLACEMENTS:
        working = working.replace(korean_name, english_name)
    for korean, english in sorted(
        ENGLISH_CANONICAL_EXACT.items(),
        key=lambda item: len(item[0]),
        reverse=True,
    ):
        if korean and re.search(r"[가-힣]", korean):
            working = working.replace(korean, english)
    if not re.search(r"[가-힣]", working):
        return working
    return None


def _canonical_english_pattern(text: str) -> Optional[str]:
    """Translate deterministic mixed labels that contain protected identifiers."""
    value = clean_text(text)
    if not value:
        return None
    working = value
    for korean_name, english_name in ENGLISH_MARKET_NAME_REPLACEMENTS:
        working = working.replace(korean_name, english_name)

    patterns: Tuple[Tuple[str, str], ...] = (
        (r"^(.+?) 규제 분석 결과$", r"\1 Regulatory Screening Results"),
        (r"^(.+?) 종합요약$", r"\1 Summary"),
        (r"^(.+?) 수동확인 체크리스트$", r"\1 Manual Regulatory Review Checklist"),
        (r"^(.+?) 추가검토 여부$", r"\1 Additional Review Selected"),
        (r"^(.+?) 시장 영향$", r"\1 Market Impact"),
        (r"^(.+?) 종합판정$", r"\1 Overall Review Decision"),
        (r"^연결된 (.+?) 확인 성분$", r"Linked \1 Verification Ingredients"),
        (r"^(.+?) 필수 조치$", r"Required \1 Actions"),
        (r"^(.+?) 확인 대상$", r"\1 Verification Ingredient"),
        (r"^(.+?) 추가검토의 해당 시장 영향$", r"Impact of the Additional \1 Review on This Market"),
        (r"^Eurasian Economic Union\(((?:EAC|\[\[P\d{4}\]\]))\)의 국가 화장품 규제판정을 변경하지는 않지만, 회원국별 수입자·바이어·유통채널 또는 할랄 표시 전략에 따라 원료 기원과 인증자료가 추가로 요구될 수 있습니다\.$",
         r"This does not change the national cosmetics regulatory screening decision for the Eurasian Economic Union(\1). However, ingredient-origin and certification documentation may be additionally required depending on the member state, importer, buyer, distribution channel, or HALAL labeling strategy."),
        (r"^(.+?) 인증 또는 관련 유통채널을 추진하는 경우 (.+?) 독립 장의 확인사항을 별도로 완료하십시오\\.$",
         r"When pursuing \1 certification or related distribution channels, separately complete the verification items in the \2 chapter."),
        (r"^(.+?):$", r"\1:"),
    )
    for pattern, replacement in patterns:
        if re.fullmatch(pattern, working):
            return re.sub(pattern, replacement, working)
    return None


def _canonicalize_english_text(text: str) -> str:
    """Normalize critical regulatory terminology in English output.

    Exact fixed labels are translated deterministically. Free-form translations
    are post-processed only for high-risk terminology that must not vary.
    """
    value = clean_text(text)
    if not value:
        return value
    if value in ENGLISH_CANONICAL_EXACT:
        return ENGLISH_CANONICAL_EXACT[value]

    replacements = (
        (r"\bBanned\b", "Prohibited"),
        (r"\bbanned\b", "prohibited"),
        (r"\bName Verification Required\b", "Ingredient Identity Verification Required"),
        (r"\bname verification required\b", "ingredient identity verification required"),
        (r"\bManual Review Required\b", "Manual Regulatory Review Required"),
        (r"\bmanual review required\b", "manual regulatory review required"),
        (r"\bWarning or Labeling Required\b", "Warning/Labeling Required"),
        (r"\bwarning or labeling required\b", "warning/labeling required"),
        (r"\bHalal\b", "HALAL"),
        (r"\bhalal\b", "HALAL"),
    )
    for pattern, replacement in replacements:
        value = re.sub(pattern, replacement, value)

    # Normalize market-label spacing and the order of the additional-review label.
    value = re.sub(
        r"\b(United States|European Union|United Kingdom|China|Saudi Arabia|Eurasian Economic Union)"
        r"\((US|EU|UK|CN|SFDA|EAC)\)",
        r"\1 (\2)",
        value,
    )
    value = value.replace("ASEAN (ASEAN)", "ASEAN")
    value = value.replace("ASEAN(ASEAN)", "ASEAN")
    value = value.replace("HALAL Additional Review", "Additional HALAL Review")
    return value


def _validate_english_terminology(text: str) -> None:
    value = _visible_translation_text(text)
    if not value:
        return
    if re.search(r"[가-힣]", value):
        raise ValueError(f"English output contains Korean text: {value[:160]}")
    for pattern, guidance in ENGLISH_DISALLOWED_PATTERNS:
        if re.search(pattern, value, flags=re.IGNORECASE):
            raise ValueError(
                "Non-canonical English regulatory terminology detected: "
                f"{value[:160]} ({guidance})"
            )



def _canonicalize_additional_language_text(text: str, language_code: str) -> str:
    value = clean_text(text)
    for pattern, replacement in ADDITIONAL_LANGUAGE_NORMALIZATION_PATTERNS.get(
        language_code, ()
    ):
        value = re.sub(pattern, replacement, value, flags=re.IGNORECASE)
    return value


def _validate_additional_language_terminology(text: str, language_code: str) -> None:
    value = _visible_translation_text(text)
    if not value:
        return
    for english_term in ADDITIONAL_LANGUAGE_FORBIDDEN_ENGLISH_TERMS:
        if english_term.casefold() in value.casefold():
            raise ValueError(
                f"{OUTPUT_LANGUAGES[language_code]} 출력에 영문 핵심 규제용어가 남아 있습니다: "
                f"{english_term}"
            )
    for pattern, guidance in ADDITIONAL_LANGUAGE_DISALLOWED_PATTERNS.get(
        language_code, ()
    ):
        if re.search(pattern, value, flags=re.IGNORECASE):
            raise ValueError(
                f"{OUTPUT_LANGUAGES[language_code]} 출력에 과장된 적합·승인 표현이 있습니다: "
                f"{value[:160]} ({guidance})"
            )


def _validate_translated_item(
    protected_source: str,
    translated_text: str,
    language_code: str,
) -> None:
    source = clean_text(protected_source)
    translated = clean_text(translated_text)
    if not translated:
        raise ValueError("빈 번역 결과")

    source_tokens = _placeholder_tokens(source)
    translated_tokens = _placeholder_tokens(translated)
    if sorted(source_tokens) != sorted(translated_tokens):
        raise ValueError(
            "보호 토큰이 누락되거나 변경되었습니다: "
            f"source={source_tokens}, translated={translated_tokens}"
        )

    source_visible = _visible_translation_text(source)
    translated_visible = _visible_translation_text(translated)

    if language_code == "ko":
        if _needs_korean_translation(source_visible):
            if _needs_korean_translation(translated_visible):
                raise ValueError(
                    "한국어 출력에 문장형 영문이 남아 있습니다: "
                    f"{translated_visible[:120]}"
                )
        return

    if language_code == "en":
        _validate_english_terminology(translated)
        return

    # 비한국어 출력에서는 보호 토큰 밖의 한국어가 남으면 혼합 언어로 본다.
    if re.search(r"[가-힣]", translated_visible):
        raise ValueError(
            f"{OUTPUT_LANGUAGES[language_code]} 출력에 한국어가 남아 있습니다: "
            f"{translated_visible[:120]}"
        )

    # 영어 이외의 추가 언어는 문장형 영문을 그대로 통과시키지 않는다.
    if (
        language_code != "en"
        and _sentence_like_english(source_visible)
        and translated_visible.casefold() == source_visible.casefold()
    ):
        raise ValueError(
            f"{OUTPUT_LANGUAGES[language_code]} 출력에 미번역 영문 문장이 남아 있습니다: "
            f"{translated_visible[:120]}"
        )

    if language_code in ADDITIONAL_LANGUAGE_CANONICAL_EXACT:
        _validate_additional_language_terminology(translated, language_code)


def _needs_korean_translation(text: str) -> bool:
    value = clean_text(text)
    if not value or re.search(r"[가-힣]", value):
        return False
    # 짧은 코드·파일명·수치가 아니라 문장형 영문인 경우만 번역한다.
    words = re.findall(r"[A-Za-z]{2,}", value)
    return len(words) >= 4 and not re.fullmatch(r"[A-Za-z0-9_.:/+\- ]+\.(csv|xlsx|xls|docx|pdf)", value, re.I)


def _translate_map(
    texts: Sequence[str],
    language_code: str,
    translator: Optional[Translator],
    protected_terms: Sequence[str],
) -> Dict[str, str]:
    unique = dedupe_preserve(texts)
    if not unique:
        return {}
    if language_code not in OUTPUT_LANGUAGES:
        raise ValueError(f"지원하지 않는 출력 언어입니다: {language_code}")

    if language_code == "ko":
        candidates = [text for text in unique if _needs_korean_translation(text)]
        result = {text: text for text in unique}
        if not candidates:
            return result
        if translator is None:
            # 번역 기능이 없더라도 이미 표준화된 한국어 문구는 유지하고,
            # 남은 원문은 삭제하거나 임의 번역하지 않는다.
            return result
        protected_texts, token_to_term = _protect_strings(candidates, protected_terms)
        translated = translator(protected_texts, language_code, list(token_to_term))
        for original, protected in zip(candidates, protected_texts):
            localized = clean_text(translated.get(protected), protected)
            _validate_translated_item(protected, localized, language_code)
            result[original] = _restore_tokens(localized, token_to_term)
        return result

    if language_code == "en":
        result: Dict[str, str] = {}
        candidates: List[str] = []
        for original in unique:
            exact = ENGLISH_CANONICAL_EXACT.get(clean_text(original))
            composite = _canonical_english_composite(original)
            patterned = _canonical_english_pattern(original)
            if exact is not None:
                canonical = _canonicalize_english_text(exact)
                _validate_english_terminology(canonical)
                result[original] = canonical
            elif composite is not None:
                canonical = _canonicalize_english_text(composite)
                _validate_english_terminology(canonical)
                result[original] = canonical
            elif patterned is not None:
                canonical = _canonicalize_english_text(patterned)
                _validate_english_terminology(canonical)
                result[original] = canonical
            elif not re.search(r"[가-힣]", original):
                canonical = _canonicalize_english_text(original)
                _validate_english_terminology(canonical)
                result[original] = canonical
            else:
                candidates.append(original)

        if not candidates:
            return result
        if translator is None:
            raise RuntimeError(
                "영문 기본본 생성을 위해 번역 기능이 필요합니다. "
                "고정 규제용어는 내장 용어집을 사용하며, 자유문장만 번역합니다."
            )

        protected_texts, token_to_term = _protect_strings(candidates, protected_terms)
        translated = translator(protected_texts, language_code, list(token_to_term))
        for original, protected in zip(candidates, protected_texts):
            localized = clean_text(translated.get(protected), protected)
            canonical_protected = _canonicalize_english_text(localized)
            _validate_translated_item(protected, canonical_protected, language_code)
            result[original] = _restore_tokens(canonical_protected, token_to_term)
        return result

    localized_exact = ADDITIONAL_LANGUAGE_CANONICAL_EXACT.get(language_code, {})
    result: Dict[str, str] = {}
    candidates: List[str] = []
    protected_set = {clean_text(value) for value in protected_terms if clean_text(value)}
    for original in unique:
        cleaned = clean_text(original)
        exact = localized_exact.get(cleaned)
        if exact is not None:
            canonical = _canonicalize_additional_language_text(exact, language_code)
            _validate_additional_language_terminology(canonical, language_code)
            result[original] = canonical
        elif cleaned in protected_set:
            # Product names, source files, ingredient names, INCI, CAS, URLs,
            # market codes, and official identifiers must remain unchanged.
            result[original] = original
        else:
            candidates.append(original)

    if not candidates:
        return result
    if translator is None:
        raise RuntimeError("선택 언어 출력을 생성하려면 번역 기능이 필요합니다.")

    # Additional languages are translated from the polished English version,
    # not directly from Korean. This avoids carrying Korean word order and
    # internal system terminology into the localized report.
    english_map = _translate_map(candidates, "en", translator, protected_terms)
    english_sources = [english_map[original] for original in candidates]
    protected_texts, token_to_term = _protect_strings(english_sources, protected_terms)

    translatable: List[Tuple[str, str]] = []
    translatable_originals: List[str] = []
    for original, protected in zip(candidates, protected_texts):
        if not _visible_translation_text(protected):
            # A string made only of protected identifiers must not be sent to a
            # translator. Preserve it exactly.
            result[original] = original
        else:
            translatable_originals.append(original)
            translatable.append((original, protected))

    if not translatable:
        return result

    translated = translator(
        [protected for _, protected in translatable],
        language_code,
        list(token_to_term),
    )
    for original, protected in translatable:
        localized = clean_text(translated.get(protected), protected)
        canonical_protected = _canonicalize_additional_language_text(
            localized, language_code
        )
        _validate_translated_item(
            protected, canonical_protected, language_code
        )
        result[original] = _restore_tokens(canonical_protected, token_to_term)
    return result

def _localized_sheet_name(original_name: str, language_code: str) -> str:
    if language_code == "ko":
        return original_name[:31]
    prefix = original_name.split("_", 1)[0]
    if "HALAL" in original_name:
        if "수동확인" in original_name:
            key = "halal_manual"
        elif "종합요약" in original_name:
            key = "halal_summary"
        else:
            key = "halal_result"
        return f"{prefix}_{SHEET_SUFFIXES[language_code][key]}"[:31]
    parts = original_name.split("_")
    target = parts[1] if len(parts) > 1 else "MARKET"
    if "수동확인" in original_name:
        key = "manual"
    elif "종합요약" in original_name:
        key = "summary"
    else:
        key = "result"
    return f"{prefix}_{target}_{SHEET_SUFFIXES[language_code][key]}"[:31]


def _localize_excel_bytes(
    data: bytes,
    language_code: str,
    translator: Optional[Translator],
    protected_terms: Sequence[str],
) -> bytes:
    if language_code == "ko" and translator is None:
        return data
    workbook = load_workbook(io.BytesIO(data))
    strings: List[str] = []
    for ws in workbook.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and clean_text(cell.value):
                    strings.append(cell.value)
    mapping = _translate_map(strings, language_code, translator, protected_terms)
    for ws in workbook.worksheets:
        if language_code != "ko":
            ws.title = _localized_sheet_name(ws.title, language_code)
        if language_code == "ar":
            ws.sheet_view.rightToLeft = True
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value in mapping:
                    cell.value = mapping[cell.value]
                    if language_code == "ar":
                        cell.alignment = Alignment(
                            horizontal="right" if cell.alignment.horizontal not in {"center", "centerContinuous"} else cell.alignment.horizontal,
                            vertical=cell.alignment.vertical or "top",
                            wrap_text=True,
                        )
        for validation in ws.data_validations.dataValidation:
            if validation.type == "list" and validation.formula1 and "확인 중" in validation.formula1:
                validation.formula1 = f'"{VALIDATION_OPTIONS[language_code]}"'
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _all_document_paragraphs(document: Document) -> List[Any]:
    paragraphs: List[Any] = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
    return paragraphs


def _set_paragraph_bidi(paragraph: Any) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _all_document_tables(document: Document) -> List[Any]:
    tables: List[Any] = list(document.tables)
    for section in document.sections:
        tables.extend(section.header.tables)
        tables.extend(section.footer.tables)
    return tables


def _set_table_bidi_visual(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    bidi_visual = tbl_pr.find(qn("w:bidiVisual"))
    if bidi_visual is None:
        bidi_visual = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi_visual)
    bidi_visual.set(qn("w:val"), "1")
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT


def _localized_run_text(original: str, mapping: Dict[str, str]) -> str:
    if original in mapping:
        return mapping[original]
    key = clean_text(original)
    if key not in mapping:
        return original
    leading = original[: len(original) - len(original.lstrip())]
    trailing = original[len(original.rstrip()):]
    return f"{leading}{mapping[key]}{trailing}"


def _localize_docx_bytes(
    data: bytes,
    language_code: str,
    translator: Optional[Translator],
    protected_terms: Sequence[str],
) -> bytes:
    if language_code == "ko" and translator is None:
        return data
    document = Document(io.BytesIO(data))
    paragraphs = _all_document_paragraphs(document)

    # Chapter labels are deterministic UI elements, not free-form translation.
    # Keeping them out of the translation API prevents overflow and inconsistent
    # forms such as a long sentence inside the compact chapter badge.
    chapter_runs: List[Tuple[Any, int]] = []
    excluded_run_ids: Set[int] = set()
    for paragraph in paragraphs:
        for run in paragraph.runs:
            match = re.fullmatch(r"제(\d+)장", clean_text(run.text))
            if match:
                chapter_runs.append((run, int(match.group(1))))
                excluded_run_ids.add(id(run))

    strings = [
        run.text
        for paragraph in paragraphs
        for run in paragraph.runs
        if id(run) not in excluded_run_ids and clean_text(run.text)
    ]
    mapping = _translate_map(strings, language_code, translator, protected_terms)
    for paragraph in paragraphs:
        if language_code == "ar":
            _set_paragraph_bidi(paragraph)
            if paragraph.alignment in {None, WD_ALIGN_PARAGRAPH.LEFT}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            if id(run) not in excluded_run_ids:
                run.text = _localized_run_text(run.text, mapping)
            if language_code == "ar":
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:cs"), "Arial")

    chapter_template = CHAPTER_LABELS[language_code]
    for run, number in chapter_runs:
        run.text = chapter_template.format(number=number)

    chapter_width = CHAPTER_BOX_WIDTHS[language_code]
    total_width = 6.65
    for table in _all_document_tables(document):
        if language_code == "ar":
            _set_table_bidi_visual(table)
        if len(table.rows) == 1 and len(table.columns) == 2:
            first_text = clean_text(table.cell(0, 0).text)
            if re.fullmatch(
                r"(?:제\d+장|Chapter \d+|第\d+章|الفصل \d+|Chapitre \d+|Bab \d+|Sura ya \d+)",
                first_text,
            ):
                table.autofit = False
                table.columns[0].width = Inches(chapter_width)
                table.columns[1].width = Inches(total_width - chapter_width)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()

# ============================================================
# ZIP 일괄 다운로드
# ============================================================
def create_bundle_zip_bytes(
    product_outputs: Sequence[dict],
    failures: Sequence[dict],
    language_code: str = "ko",
) -> bytes:
    lang = language_code if language_code in OUTPUT_LANGUAGES else "ko"
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for product in product_outputs:
            product_name = clean_text(product.get("product_name"), "제품")
            excel_data = product.get("excel_data")
            report_data = product.get("report_data")
            if excel_data:
                archive.writestr(
                    output_file_name(product_name, "excel", lang, "xlsx"),
                    excel_data,
                )
            if report_data:
                archive.writestr(
                    output_file_name(product_name, "report", lang, "docx"),
                    report_data,
                )

        if failures:
            text_stream = io.StringIO()
            if lang == "ko":
                fieldnames = ["제품명", "시장", "오류"]
            else:
                fieldnames = ["Product", "Market", "Error"]
            writer = csv.DictWriter(text_stream, fieldnames=fieldnames)
            writer.writeheader()
            for failure in failures:
                if lang == "ko":
                    row = {
                        "제품명": clean_text(failure.get("product_name")),
                        "시장": clean_text(failure.get("target")),
                        "오류": clean_text(failure.get("error")),
                    }
                else:
                    row = {
                        "Product": clean_text(failure.get("product_name")),
                        "Market": clean_text(failure.get("target")),
                        "Error": clean_text(failure.get("error")),
                    }
                writer.writerow(row)
            archive.writestr(
                f"{FILE_LABELS[lang]['failures']}_{LANGUAGE_FILE_SUFFIXES[lang]}.csv",
                "\ufeff" + text_stream.getvalue(),
            )

    return output.getvalue()

